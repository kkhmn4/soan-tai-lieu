import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
import docx

doc = docx.Document(r'scripts\ws_v2_copy.docx')

print("=== PARAGRAPHS ===")
for i, para in enumerate(doc.paragraphs):
    if para.text.strip():
        print(f'[P{i}] STYLE={para.style.name} | {para.text}')

print("\n=== TABLES ===")
for ti, table in enumerate(doc.tables):
    print(f'\n--- TABLE {ti} ---')
    for ri, row in enumerate(table.rows):
        cells = []
        for cell in row.cells:
            # Get full text including nested paragraphs
            cell_text = '\n'.join(p.text for p in cell.paragraphs if p.text.strip())
            cells.append(cell_text.replace('\n', ' | '))
        print(f'  ROW {ri}: ' + ' || '.join(cells))
