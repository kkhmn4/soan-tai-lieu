# -*- coding: utf-8 -*-
"""
viet_pht_exporter.py — Standardized Worksheet Exporter (PHT) → DOCX
=================================================================
Implementation of Vietnamese Educational Document Formatting Standard § 5.1
Exports study worksheets from Markdown to structured DOCX.
"""

import re
import os
from viet_styles import (
    Document, setup_viet_margins, set_viet_default_style,
    VietColors, VietFonts,
    add_viet_title, add_viet_heading, add_viet_body, add_viet_bullet,
    make_viet_table, add_viet_dot_line, add_viet_separator,
    add_viet_page_number_footer, add_formatted_runs,
    smart_typography, clean_latex, Cm, Pt
)
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def add_viet_student_info_table(doc, lesson_title):
    """Create a structured 2x3 grid border table for student metadata block."""
    table = doc.add_table(rows=3, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT = 1 # Center
    
    # Set table width (17.0 cm total width)
    # Col 0: 9.5cm, Col 1: 7.5cm
    col_widths = [Cm(9.5), Cm(7.5)]
    
    metadata = [
        [f"Họ và tên học sinh: .......................................", "Lớp: .................  Nhóm: ................."],
        ["Môn học: Vật lí 12 (CTGDPT 2018)", f"Bài học: {lesson_title}"],
        ["Tiết học: ............ ngày ......../......../202...", "Điểm số: .........  Chữ ký GV: ........."]
    ]
    
    for r_idx, row_data in enumerate(metadata):
        row = table.rows[r_idx]
        # set cantSplit
        trPr = row._tr.get_or_add_trPr()
        trPr.append(OxmlElement('w:cantSplit'))
        
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.width = col_widths[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.left_indent = Cm(0.2)
            
            run = p.add_run(val)
            run.font.name = VietFonts.BODY
            run.font.size = VietFonts.SIZE_INFO
            run.font.color.rgb = VietColors.DARK
            
            # Apply borders and margins
            # Margins
            tcPr = cell._tc.get_or_add_tcPr()
            tcMar = OxmlElement('w:tcMar')
            for m in ['top', 'bottom', 'left', 'right']:
                n = OxmlElement(f'w:{m}')
                n.set(qn('w:w'), '120') # 6pt
                n.set(qn('w:type'), 'dxa')
                tcMar.append(n)
            tcPr.append(tcMar)
            
            # Border: thin light gray
            tcBorders = OxmlElement('w:tcBorders')
            for side in ['top', 'left', 'bottom', 'right']:
                b = OxmlElement(f'w:{side}')
                b.set(qn('w:val'), 'single')
                b.set(qn('w:sz'), '4')  # 0.5 pt
                b.set(qn('w:space'), '0')
                b.set(qn('w:color'), VietColors.BORDER_GRAY_HEX)
                tcBorders.append(b)
            tcPr.append(tcBorders)
            
    doc.add_paragraph() # spacer

def add_viet_self_evaluation_table(doc):
    """Add a self-evaluation table at the end of the worksheet."""
    add_viet_heading(doc, "III. TỰ ĐÁNH GIÁ KẾT QUẢ", level=1)
    
    headers = ["Tiêu chí học tập", "Đạt", "Chưa Đạt"]
    rows = [
        ["1. Nhận biết và phát biểu được các khái niệm cốt lõi của bài học.", "   [  ]", "   [  ]"],
        ["2. Áp dụng đúng công thức tính toán và giải quyết bài tập vận dụng thực tế.", "   [  ]", "   [  ]"],
        ["3. Tích cực thảo luận, hoàn thành đầy đủ nhiệm vụ học tập nhóm/cá nhân.", "   [  ]", "   [  ]"]
    ]
    make_viet_table(doc, headers, rows, col_widths=[12.0, 2.5, 2.5], total_width_cm=17.0)

    # Chữ ký học sinh
    p = doc.add_paragraph()
    p.alignment = 2 # Right
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.right_indent = Cm(1.0)
    
    run_date = p.add_run("Ngày ........ tháng ........ năm 202...\n")
    run_date.font.name = VietFonts.BODY
    run_date.font.size = VietFonts.SIZE_INFO
    run_date.font.color.rgb = VietColors.DARK
    run_date.italic = True
    
    run_sig = p.add_run("Học sinh ký tên\n\n\n..........................................")
    run_sig.bold = True
    run_sig.font.name = VietFonts.BODY
    run_sig.font.size = VietFonts.SIZE_INFO
    run_sig.font.color.rgb = VietColors.DARK

def export_pht(md_path, output_path, lesson_label):
    doc = Document()
    setup_viet_margins(doc, doc_type="pht")
    set_viet_default_style(doc)
    add_viet_page_number_footer(doc)
    
    # 1. School name (header)
    p_header = doc.add_paragraph()
    p_header.alignment = 0 # Left
    p_header.paragraph_format.space_after = Pt(2)
    run_sch = p_header.add_run("TRƯỜNG THPT: ................................................\n")
    run_sch.font.name = VietFonts.BODY
    run_sch.font.size = VietFonts.SIZE_INFO
    run_sch.italic = True
    run_sch.font.color.rgb = VietColors.GRAY
    
    # 2. Document Title
    add_viet_title(doc, "PHIẾU HỌC TẬP")
    
    # Subtitle
    p_sub = doc.add_paragraph()
    p_sub.alignment = 1 # Center
    p_sub.paragraph_format.space_before = Pt(0)
    p_sub.paragraph_format.space_after = Pt(10)
    run_sub = p_sub.add_run("(Mẫu tài liệu học tập theo CTGDPT 2018)")
    run_sub.font.name = VietFonts.BODY
    run_sub.font.size = VietFonts.SIZE_TINY
    run_sub.font.color.rgb = VietColors.GRAY
    run_sub.italic = True
    
    # 3. Student info block
    add_viet_student_info_table(doc, lesson_label)
    
    # 4. Parse Markdown Body
    if not os.path.exists(md_path):
        doc.add_paragraph(f"Lỗi: Không tìm thấy file Markdown tại {md_path}")
        save_doc(doc, output_path)
        return
        
    with open(md_path, "r", encoding="utf-8") as f:
        raw_lines = f.readlines()
        
    lines = [line.replace('\ufeff', '') for line in raw_lines]
    
    # Lọc bỏ header metadata trùng lặp (Họ tên học sinh...)
    cleaned_lines = []
    skip_header = True
    for line in lines:
        stripped = line.strip()
        # Bắt đầu nhận diện nội dung thật từ dấu ## hoặc từ mục I. YÊU CẦU CẦN ĐẠT hoặc 📍
        if stripped.startswith("## ") or stripped.startswith("##") or "📍" in stripped:
            skip_header = False
        if not skip_header:
            cleaned_lines.append(line)
            
    # Xử lý gộp dòng trống
    processed_lines = []
    for line in cleaned_lines:
        stripped = line.strip()
        # Chuyển đổi dấu H4 (####) thành in đậm
        if stripped.startswith("####"):
            stripped = stripped.lstrip("#").strip()
            processed_lines.append(f"**{stripped}**")
        else:
            processed_lines.append(line)
            
    # Tiến hành duyệt từng dòng và vẽ DOCX
    i = 0
    while i < len(processed_lines):
        line = processed_lines[i]
        stripped = line.strip()
        
        if not stripped:
            i += 1
            continue
            
        # Kiểm tra blockquote
        is_blockquote = False
        if stripped.startswith('> '):
            stripped = stripped[2:]
            is_blockquote = True
        elif stripped.startswith('>'):
            stripped = stripped[1:]
            is_blockquote = True
            
        # Kiểm tra hình ảnh
        img_match = re.search(r'([\(\s]*)(ready[^\(\s]+\.(?:png|jpg|jpeg))([\s\)]*)', stripped)
        embedded_img_path = None
        if img_match:
            entire_match = img_match.group(0)
            img_path = img_match.group(2)
            resolved = img_path
            if not os.path.exists(resolved):
                lesson_root = os.path.dirname(os.path.dirname(md_path))
                resolved = os.path.join(lesson_root, img_path)
            if not os.path.exists(resolved):
                lesson_root = os.path.dirname(os.path.dirname(md_path))
                resolved = os.path.join(lesson_root, "ready", "hinh_anh", os.path.basename(img_path))
            if os.path.exists(resolved):
                embedded_img_path = resolved
                stripped = stripped.replace(entire_match, " ")
                stripped = re.sub(r'\s+', ' ', stripped).strip()

        # Render headings
        if stripped.startswith('## '):
            title_text = stripped[3:].strip()
            add_viet_heading(doc, title_text, level=1)
        elif stripped.startswith('### '):
            title_text = stripped[4:].strip()
            add_viet_heading(doc, title_text, level=2)
        elif stripped.startswith('#### '):
            title_text = stripped[5:].strip()
            add_viet_heading(doc, title_text, level=3)
        elif stripped.startswith('**') and stripped.endswith('**') and len(stripped) < 100:
            # heading dạng in đậm
            title_text = stripped[2:-2].strip()
            add_viet_heading(doc, title_text, level=3)
        elif stripped == '[DOT_LINE_90]':
            add_viet_dot_line(doc, count=1)
        elif stripped.startswith('- ') or stripped.startswith('* '):
            add_viet_bullet(doc, stripped, is_blockquote=is_blockquote, level=1)
            if embedded_img_path:
                p_img = doc.add_paragraph()
                p_img.alignment = 1 # center
                p_img.add_run().add_picture(embedded_img_path, width=Cm(12.0))
        elif stripped.startswith('+ '):
            add_viet_bullet(doc, stripped, is_blockquote=is_blockquote, level=2)
            if embedded_img_path:
                p_img = doc.add_paragraph()
                p_img.alignment = 1 # center
                p_img.add_run().add_picture(embedded_img_path, width=Cm(12.0))
        else:
            add_viet_body(doc, stripped, is_blockquote=is_blockquote)
            if embedded_img_path:
                p_img = doc.add_paragraph()
                p_img.alignment = 1 # center
                p_img.add_run().add_picture(embedded_img_path, width=Cm(12.0))
                
        i += 1
        
    # Thêm bảng tự đánh giá
    add_viet_self_evaluation_table(doc)
    
    # Save document
    doc.save(output_path)
