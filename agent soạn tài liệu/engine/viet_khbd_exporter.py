# -*- coding: utf-8 -*-
"""
viet_khbd_exporter.py — Lesson Plan Exporter (KHBD / Giáo án 5512) → DOCX
=======================================================================
Implementation of Vietnamese Educational Document Formatting Standard § 5.2 & § 5.3
Exports lesson plans under Official Dispatch 5512/BGDĐT-GDTrH.
"""

import re
import os
from viet_styles import (
    Document, setup_viet_margins, set_viet_default_style,
    VietColors, VietFonts,
    add_viet_title, add_viet_heading, add_viet_body, add_viet_bullet,
    make_viet_table, add_viet_dot_line, add_viet_separator,
    add_viet_page_number_footer, add_formatted_runs,
    smart_typography, clean_latex, Cm, Pt,
    set_cell_shading, set_cell_margins, set_cell_borders
)
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def add_viet_khbd_admin_header(doc, lesson_title, duration_str):
    """Add administrative department headers and metadata table for KHBD."""
    # 1. Left aligned agency details
    p = doc.add_paragraph()
    p.alignment = 0 # Left
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    
    run_dept = p.add_run("SỞ GIÁO DỤC VÀ ĐÀO TẠO: ................................................\n")
    run_dept.font.name = VietFonts.BODY
    run_dept.font.size = VietFonts.SIZE_INFO
    run_dept.font.color.rgb = VietColors.DARK
    
    run_school = p.add_run("TRƯỜNG THPT: ................................................................\n")
    run_school.bold = True
    run_school.font.name = VietFonts.BODY
    run_school.font.size = VietFonts.SIZE_INFO
    run_school.font.color.rgb = VietColors.DARK
    
    run_dept2 = p.add_run("Tổ bộ môn: Vật lí - Công nghệ")
    run_dept2.font.name = VietFonts.BODY
    run_dept2.font.size = VietFonts.SIZE_INFO
    run_dept2.italic = True
    run_dept2.font.color.rgb = VietColors.GRAY
    
    # 2. Main Title
    add_viet_title(doc, "KẾ HOẠCH BÀI DẠY")
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = 1 # Center
    p_sub.paragraph_format.space_before = Pt(0)
    p_sub.paragraph_format.space_after = Pt(8)
    run_sub = p_sub.add_run("(Mẫu chuẩn theo Công văn 5512/BGDĐT-GDTrH)")
    run_sub.font.name = VietFonts.BODY
    run_sub.font.size = VietFonts.SIZE_INFO
    run_sub.font.color.rgb = VietColors.GRAY
    run_sub.italic = True
    
    # 3. Metadata Table
    table = doc.add_table(rows=2, cols=4)
    table.alignment = 1 # center
    set_table_width_dxa_custom(table, 16.0) # custom 16.0 cm width for KHBD
    
    col_widths = [Cm(4.0), Cm(4.0), Cm(4.0), Cm(4.0)]
    metadata = [
        ["Tên môn học:", "Vật lí", "Lớp học:", "12"],
        ["Tên bài học:", lesson_title, "Thời lượng:", duration_str]
    ]
    
    for r_idx, row_data in enumerate(metadata):
        row = table.rows[r_idx]
        trPr = row._tr.get_or_add_trPr()
        trPr.append(OxmlElement('w:cantSplit'))
        
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.width = col_widths[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.left_indent = Cm(0.2)
            
            # Label in bold
            is_label = (c_idx % 2 == 0)
            run = p.add_run(val)
            run.bold = is_label
            run.font.name = VietFonts.BODY
            run.font.size = VietFonts.SIZE_INFO
            run.font.color.rgb = VietColors.DARK
            
            # Formatting
            tcPr = cell._tc.get_or_add_tcPr()
            tcMar = OxmlElement('w:tcMar')
            for m in ['top', 'bottom', 'left', 'right']:
                n = OxmlElement(f'w:{m}')
                n.set(qn('w:w'), '100')
                n.set(qn('w:type'), 'dxa')
                tcMar.append(n)
            tcPr.append(tcMar)
            
            tcBorders = OxmlElement('w:tcBorders')
            for side in ['top', 'left', 'bottom', 'right']:
                b = OxmlElement(f'w:{side}')
                b.set(qn('w:val'), 'single')
                b.set(qn('w:sz'), '4')
                b.set(qn('w:space'), '0')
                b.set(qn('w:color'), VietColors.BORDER_GRAY_HEX)
                tcBorders.append(b)
            tcPr.append(tcBorders)
            
    doc.add_paragraph() # spacer

def set_table_width_dxa_custom(table, width_cm):
    width_dxa = int(width_cm * 567)
    tblPr = table._tbl.tblPr
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(width_dxa))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)

def add_viet_signature_block(doc):
    """Add the 3-party approval signature block at the end of the lesson plan."""
    p_date = doc.add_paragraph()
    p_date.alignment = 2 # Right
    p_date.paragraph_format.space_before = Pt(14)
    p_date.paragraph_format.right_indent = Cm(1.0)
    
    run_date = p_date.add_run("........., ngày ....... tháng ....... năm 202...\n")
    run_date.font.name = VietFonts.BODY
    run_date.font.size = VietFonts.SIZE_INFO
    run_date.italic = True
    run_date.font.color.rgb = VietColors.DARK

    table = doc.add_table(rows=2, cols=3)
    table.alignment = 1 # Center
    set_table_width_dxa_custom(table, 16.0)
    col_widths = [Cm(5.3), Cm(5.3), Cm(5.4)]
    
    titles = [
        "XÁC NHẬN CỦA BAN GIÁM HIỆU\n(Ký, đóng dấu)",
        "XÁC NHẬN CỦA TỔ TRƯỞNG\n(Ký và ghi rõ họ tên)",
        "GIÁO VIÊN SOẠN BÀI\n(Ký và ghi rõ họ tên)"
    ]
    
    # Row 0: Titles
    row0 = table.rows[0]
    trPr0 = row0._tr.get_or_add_trPr()
    trPr0.append(OxmlElement('w:cantSplit'))
    for idx, text in enumerate(titles):
        cell = row0.cells[idx]
        cell.width = col_widths[idx]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = 1 # Center
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        run.bold = True
        run.font.name = VietFonts.BODY
        run.font.size = VietFonts.SIZE_INFO
        run.font.color.rgb = VietColors.DARK
        
    # Row 1: Blank space for signature
    row1 = table.rows[1]
    trPr1 = row1._tr.get_or_add_trPr()
    trPr1.append(OxmlElement('w:cantSplit'))
    for idx in range(3):
        cell = row1.cells[idx]
        cell.width = col_widths[idx]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = 1 # Center
        p.paragraph_format.space_before = Pt(45) # blank height for signature
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run("..................................................")
        run.font.name = VietFonts.BODY
        run.font.size = VietFonts.SIZE_INFO
        run.font.color.rgb = VietColors.GRAY
        
    # Remove borders from signature block table
    for row in table.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for side in ['top', 'left', 'bottom', 'right']:
                b = OxmlElement(f'w:{side}')
                b.set(qn('w:val'), 'none')
                tcBorders.append(b)
            tcPr.append(tcBorders)

def add_viet_adjustments_section(doc, advantages, limitations, solutions):
    """Add post-teaching reflection adjustment box with extra handwritten note space."""
    add_viet_heading(doc, "IV. ĐIỀU CHỈNH BÀI DẠY SAU TIẾT GIẢNG", level=1)
    
    p = doc.add_paragraph()
    p.alignment = 0 # Left
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.3
    
    p.add_run("1. Ưu điểm nổi bật:\n").bold = True
    p.add_run(advantages or "............................................................................................................................................\n")
    p.add_run("   ............................................................................................................................................\n").font.color.rgb = VietColors.GRAY
    
    p.add_run("2. Hạn chế, khó khăn:\n").bold = True
    p.add_run(limitations or "............................................................................................................................................\n")
    p.add_run("   ............................................................................................................................................\n").font.color.rgb = VietColors.GRAY
    
    p.add_run("3. Hướng điều chỉnh sư phạm:\n").bold = True
    p.add_run(solutions or "............................................................................................................................................\n")
    p.add_run("   ............................................................................................................................................\n").font.color.rgb = VietColors.GRAY
    
    # Set formatting for all runs
    for run in p.runs:
        run.font.name = VietFonts.BODY
        run.font.size = VietFonts.SIZE_BODY
        if run.text.startswith(" ") or run.text.startswith("   "):
            run.font.color.rgb = VietColors.GRAY

def split_cell_into_blocks(text):
    text = re.sub(r'<br\s*/?>', '\n', text)
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    
    blocks = []
    current_block = []
    
    new_block_triggers = [
        r'^\*\*chuyển giao nhiệm vụ',
        r'^\*\*tiếp nhận nhiệm vụ',
        r'^\*\*thực hiện nhiệm vụ',
        r'^\*\*theo dõi,\s*hỗ trợ',
        r'^\*\*tổ chức báo cáo',
        r'^\*\*báo cáo,\s*thảo luận',
        r'^\*\*kết luận,\s*nhận định',
        r'^\*\*kết luận',
        r'^\*\*ghi nhận kiến thức',
        r'^\*\*báo cáo',
        r'^\*\*hướng dẫn',
        r'^\*\*đánh giá',
        r'^\*\*thực hiện',
        r'^\*\*thảo luận',
        r'^\*\*nội dung',
        r'^\*\*bước',
        r'^- gv đặt câu hỏi',
        r'^- gv hỏi',
        r'^- thảo luận cặp đôi',
        r'^- thảo luận nhóm',
        r'^- thảo luận nhanh'
    ]
    
    for line in lines:
        is_trigger = False
        line_lower = line.lower()
        for trig in new_block_triggers:
            if re.match(trig, line_lower):
                is_trigger = True
                break
                
        if is_trigger:
            if current_block:
                blocks.append('\n'.join(current_block))
                current_block = [line]
            else:
                current_block.append(line)
        else:
            current_block.append(line)
            
    if current_block:
        blocks.append('\n'.join(current_block))
        
    return blocks

def make_viet_gv_hs_table(doc, headers, rows):
    """
    Create a 2-column GV/HS activity table.
    Header background: light gray #D9D9D9, font 11pt, padding 0.1cm.
    """
    processed_rows = []
    for row_data in rows:
        if len(row_data) >= 2:
            gv_blocks = split_cell_into_blocks(row_data[0])
            hs_blocks = split_cell_into_blocks(row_data[1])
            max_blocks = max(len(gv_blocks), len(hs_blocks))
            for idx in range(max_blocks):
                gv_val = gv_blocks[idx] if idx < len(gv_blocks) else ""
                hs_val = hs_blocks[idx] if idx < len(hs_blocks) else ""
                processed_rows.append([gv_val, hs_val])
        else:
            processed_rows.append(row_data)
            
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.alignment = 1 # Center
    
    set_table_width_dxa_custom(table, 16.0) # 16cm total width for KHBD
    col_widths = [Cm(8.0), Cm(8.0)]
    
    # Header Row
    hdr_cells = table.rows[0].cells
    for i, text in enumerate(headers):
        cell = hdr_cells[i]
        cell.width = col_widths[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = 1 # Center
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        
        run = p.add_run(smart_typography(text))
        run.bold = True
        run.font.name = VietFonts.BODY
        run.font.size = VietFonts.SIZE_INFO # 11pt for GV/HS table
        run.font.color.rgb = VietColors.DARK
        
        set_cell_shading(cell, VietColors.LIGHT_GRAY_HEX)
        set_cell_margins(cell)
        set_cell_borders(cell)
        
    # Repeat Header Row
    trPr_hdr = table.rows[0]._tr.get_or_add_trPr()
    trPr_hdr.append(OxmlElement('w:tblHeader'))
    
    # Data Rows
    for r_idx, row_data in enumerate(processed_rows):
        row = table.add_row()
        trPr = row._tr.get_or_add_trPr()
        trPr.append(OxmlElement('w:cantSplit'))
        
        for c_idx in range(2):
            val = row_data[c_idx] if c_idx < len(row_data) else ""
            cell = row.cells[c_idx]
            cell.width = col_widths[c_idx]
            cell.text = ""
            
            lines = str(val).replace("<br>", "\n").replace("<br/>", "\n").split("\n")
            first_p = True
            for line in lines:
                line_text = line.strip()
                if not line_text:
                    continue
                    
                if first_p:
                    p = cell.paragraphs[0]
                    first_p = False
                else:
                    p = cell.add_paragraph()
                    
                p.alignment = 0 # Left
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.15
                
                # Bullet list inside table cell
                bullet_prefix = ""
                if line_text.startswith('- ') or line_text.startswith('* '):
                    bullet_prefix = "- "
                    line_text = line_text[2:]
                    p.paragraph_format.left_indent = Cm(0.3)
                    p.paragraph_format.first_line_indent = Cm(-0.2)
                elif line_text.startswith('+ '):
                    bullet_prefix = "+ "
                    line_text = line_text[2:]
                    p.paragraph_format.left_indent = Cm(0.5)
                    p.paragraph_format.first_line_indent = Cm(-0.2)
                    
                cell_parts = re.split(r'(\*\*.*?\*\*|\$.*?\$)', line_text)
                is_first_run = True
                for part in cell_parts:
                    if not part:
                        continue
                    if part.startswith('**') and part.endswith('**'):
                        r_text = smart_typography(part[2:-2])
                        run = p.add_run(r_text)
                        run.bold = True
                        
                        # Auto bold primary keywords inside KHBD tables (e.g. Chuyển giao nhiệm vụ:)
                        _BOLD_KWS = [
                            "Chuyển giao nhiệm vụ:", "Tiếp nhận nhiệm vụ:", "Thực hiện nhiệm vụ:", 
                            "Theo dõi, hỗ trợ:", "Tổ chức báo cáo:", "Báo cáo, thảo luận:", 
                            "Kết luận:", "Kết luận, nhận định:", "Ghi nhận kiến thức:", 
                            "Sản phẩm:", "Hỗ trợ:", "Mục tiêu:", "Nội dung:"
                        ]
                        if any(kw in r_text for kw in _BOLD_KWS):
                            run.font.color.rgb = VietColors.PRIMARY
                    elif part.startswith('$') and part.endswith('$'):
                        formula = clean_latex(part[1:-1])
                        run = p.add_run(formula)
                        run.italic = True
                    else:
                        r_text = smart_typography(part)
                        if bullet_prefix and is_first_run:
                            r_bullet = p.add_run(bullet_prefix)
                            r_bullet.font.name = VietFonts.BODY
                            r_bullet.font.size = VietFonts.SIZE_INFO
                            r_bullet.font.color.rgb = VietColors.DARK
                            is_first_run = False
                        run = p.add_run(r_text)
                        
                    run.font.name = VietFonts.BODY
                    run.font.size = VietFonts.SIZE_INFO # 11pt
                    run.font.color.rgb = VietColors.DARK
                    
            set_cell_margins(cell)
            set_cell_borders(cell)
            
    doc.add_paragraph() # spacer

def export_khbd(md_path, output_path):
    doc = Document()
    setup_viet_margins(doc, doc_type="khbd")
    set_viet_default_style(doc)
    add_viet_page_number_footer(doc)
    
    if not os.path.exists(md_path):
        doc.add_paragraph(f"Lỗi: Không tìm thấy file Markdown tại {md_path}")
        doc.save(output_path)
        return
        
    with open(md_path, "r", encoding="utf-8") as f:
        raw_lines = f.readlines()
        
    lines = [line.replace('\ufeff', '') for line in raw_lines]
    
    # 1. Trích xuất tên bài học & thời lượng từ metadata
    lesson_title = "Bài học"
    duration_str = "02 tiết"
    for line in lines[:10]:
        if "KẾ HOẠCH BÀI DẠY:" in line or "KẾ HOẠCH BÀI DẠY —" in line or "KẾ HOẠCH BÀI DẠY" in line:
            lesson_title = line.replace("KẾ HOẠCH BÀI DẠY:", "").replace("#", "").strip()
        if "Thời lượng thực hiện:" in line or "Thời lượng:" in line:
            duration_str = line.replace("Thời lượng thực hiện:", "").replace("Thời lượng:", "").replace("*", "").strip()
            
    # 2. Tạo Header hành chính
    add_viet_khbd_admin_header(doc, lesson_title, duration_str)
    
    # Clean lines to extract actual body content starting from ## I. YÊU CẦU CẦN ĐẠT
    cleaned_lines = []
    skip_header = True
    for line in lines:
        stripped = line.strip()
        if stripped.startswith("## I.") or stripped.startswith("## I ") or "## I. YÊU CẦU CẦN ĐẠT" in stripped.upper():
            skip_header = False
        if not skip_header:
            cleaned_lines.append(line)
            
    # Parse body
    i = 0
    advantages, limitations, solutions = "", "", ""
    in_reflection = False
    
    while i < len(cleaned_lines):
        line = cleaned_lines[i]
        stripped = line.strip()
        
        if not stripped:
            i += 1
            continue
            
        # Kiểm tra phần phản hồi rút kinh nghiệm cuối bài dạy
        if stripped.startswith("## IV. ĐIỀU CHỈNH BÀI DẠY") or stripped.startswith("## IV. RÚT KINH NGHIỆM") or stripped.startswith("## IV "):
            in_reflection = True
            i += 1
            continue
            
        if in_reflection:
            # Thu thập thông tin ưu/khuyết/điều chỉnh
            if "1. Ưu điểm:" in stripped or "Ưu điểm:" in stripped:
                advantages = stripped.replace("1. Ưu điểm:", "").replace("Ưu điểm:", "").strip()
            elif "2. Hạn chế:" in stripped or "Hạn chế:" in stripped:
                limitations = stripped.replace("2. Hạn chế:", "").replace("Hạn chế:", "").strip()
            elif "3. Hướng điều chỉnh:" in stripped or "Hướng điều chỉnh:" in stripped:
                solutions = stripped.replace("3. Hướng điều chỉnh:", "").replace("Hướng điều chỉnh:", "").strip()
            i += 1
            continue
            
        # Kiểm tra blockquote
        is_blockquote = False
        if stripped.startswith('> '):
            stripped = stripped[2:]
            is_blockquote = True
        elif stripped.startswith('>'):
            stripped = stripped[1:]
            is_blockquote = True

        # Nhận diện hình ảnh
        img_match = re.search(r'([\(\s]*)(ready[^\(\s]+\.(?:png|jpg|jpeg))([\s\)]*)', stripped)
        embedded_img_path = None
        if img_match:
            entire_match = img_match.group(0)
            img_path = img_match.group(2)
            resolved = img_path
            if not os.path.exists(resolved):
                lesson_root = os.path.dirname(os.path.dirname(md_path))
                resolved = os.path.join(lesson_root, img_path)
            if not os.path.exists(resolved):
                lesson_root = os.path.dirname(os.path.dirname(md_path))
                resolved = os.path.join(lesson_root, "ready", "hinh_anh", os.path.basename(img_path))
            if os.path.exists(resolved):
                embedded_img_path = resolved
                stripped = stripped.replace(entire_match, " ")
                stripped = re.sub(r'\s+', ' ', stripped).strip()

        # Kiểm tra bảng hoạt động GV/HS (phát hiện bảng Markdown 2 cột)
        if stripped.startswith('|'):
            # Đọc headers của bảng
            header_line = stripped
            i += 1
            separator_line = cleaned_lines[i] if i < len(cleaned_lines) else ""
            i += 1
            
            # Đọc các hàng dữ liệu
            rows_data = []
            while i < len(cleaned_lines) and cleaned_lines[i].strip().startswith('|'):
                rows_data.append(cleaned_lines[i].strip())
                i += 1
                
            # Parse các ô dữ liệu
            def parse_row(row_str):
                return [cell.strip() for cell in row_str.split('|')[1:-1]]
                
            tbl_headers = parse_row(header_line)
            tbl_rows = [parse_row(row) for row in rows_data]
            
            # Kiểm tra xem có phải bảng hoạt động GV/HS 2 cột
            if len(tbl_headers) == 2:
                make_viet_gv_hs_table(doc, tbl_headers, tbl_rows)
            else:
                make_viet_table(doc, tbl_headers, tbl_rows, col_widths=None, total_width_cm=16.0)
            continue

        # Render headings thông thường
        if stripped.startswith('## '):
            add_viet_heading(doc, stripped[3:].strip(), level=1)
        elif stripped.startswith('### '):
            add_viet_heading(doc, stripped[4:].strip(), level=2)
        elif stripped.startswith('#### '):
            add_viet_heading(doc, stripped[5:].strip(), level=3)
        elif stripped.startswith('- ') or stripped.startswith('* '):
            add_viet_bullet(doc, stripped, is_blockquote=is_blockquote, level=1)
            if embedded_img_path:
                p_img = doc.add_paragraph()
                p_img.alignment = 1
                p_img.add_run().add_picture(embedded_img_path, width=Cm(12.0))
        elif stripped.startswith('+ '):
            add_viet_bullet(doc, stripped, is_blockquote=is_blockquote, level=2)
            if embedded_img_path:
                p_img = doc.add_paragraph()
                p_img.alignment = 1
                p_img.add_run().add_picture(embedded_img_path, width=Cm(12.0))
        else:
            add_viet_body(doc, stripped, is_blockquote=is_blockquote)
            if embedded_img_path:
                p_img = doc.add_paragraph()
                p_img.alignment = 1
                p_img.add_run().add_picture(embedded_img_path, width=Cm(12.0))
                
        i += 1
        
    # Thêm phần điều chỉnh bài dạy hậu tiết giảng
    add_viet_adjustments_section(doc, advantages, limitations, solutions)
    
    # Thêm bảng chữ ký phê duyệt 3 bên
    add_viet_signature_block(doc)
    
    doc.save(output_path)
