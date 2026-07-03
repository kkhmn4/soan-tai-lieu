# -*- coding: utf-8 -*-
"""
gems_styles.py — GEMS v8.0 Design System
==========================================
Shared styles, colors, fonts, and helpers for all DOCX exports.
USAGE: from gems_styles import *

Design Tokens:
  - NAVY (#1E3A5F): Primary headers, table header backgrounds
  - MINT (#E8F5E9): Table row alternating shading
  - Times New Roman throughout (12pt body, 18pt H1, 15pt H2, 13pt H3)
  - A4, 2cm margins
  - Gray (#BFBFBF) dotted answer lines and separators
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
import os
import re



# ============================================================
#  COLOUR PALETTE
# ============================================================
class GEMSColors:
    NAVY = RGBColor(0x1E, 0x3A, 0x5F)          # #1E3A5F
    WHITE = RGBColor(0xFF, 0xFF, 0xFF)
    DARK = RGBColor(0x21, 0x21, 0x21)           # #212121 body text
    DARK = RGBColor(0x21, 0x21, 0x21)           # #212121 body text
    GRAY = RGBColor(0x59, 0x59, 0x59)           # #595959
    LIGHT_GRAY = RGBColor(0xBF, 0xBF, 0xBF)     # #BFBFBF
    ORANGE = RGBColor(0xFF, 0x98, 0x00)          # accent
    GREEN = RGBColor(0x4C, 0xAF, 0x50)           # secondary accent
    MINT = RGBColor(0xE8, 0xF5, 0xE9)            # #E8F5E9

    # Hex strings for XML shading attributes
    NAVY_HEX = "1E3A5F"
    MINT_HEX = "E8F5E9"
    WHITE_HEX = "FFFFFF"
    LIGHT_GRAY_HEX = "BFBFBF"


# ============================================================
#  FONTS
# ============================================================
class GEMSFonts:
    BODY = "Times New Roman"
    HEADER = "Times New Roman"

    SIZE_BODY = Pt(13)
    SIZE_SMALL = Pt(11)
    SIZE_TABLE = Pt(11)
    SIZE_H1 = Pt(16)   # CV5512: tiêu đề bài 14–16pt
    SIZE_H2 = Pt(14)   # CV5512: mục lớn I, II, III 13–14pt
    SIZE_H3 = Pt(13)
    SIZE_INFO = Pt(11)
    SIZE_TINY = Pt(10)


# ============================================================
#  PAGE SETUP
# ============================================================
def setup_page_margins(doc, doc_type="general"):
    """Setup page dimensions (A4) and margins based on document type (Rule 1)."""
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        # Strictly Left: 3.0cm, Right: 1.5cm, Top/Bottom: 2.0cm for all documents (Rule 1)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(1.5)


def setup_a4_page(doc):
    """Fallback function for backward compatibility."""
    setup_page_margins(doc, doc_type="general")


def set_default_style(doc):
    """Default font: Times New Roman 13 pt, 1.15 spacing, Justify alignment, space 4pt after (Rule 1)."""
    style = doc.styles['Normal']
    font = style.font
    font.name = GEMSFonts.BODY
    font.size = GEMSFonts.SIZE_BODY
    font.color.rgb = GEMSColors.DARK
    pf = style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(4)
    pf.line_spacing = 1.15
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY   # CV5512: căn đều hai bên


# ============================================================
#  TABLE HELPERS
# ============================================================
def set_cell_shading(cell, hex_color):
    """Background colour for a table cell (hex str e.g. '1E3A5F')."""
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_margins(cell, top=120, bottom=120, left=160, right=160):
    """Padding inside a table cell (in dxa = 1/20 pt). Top/bottom: 6pt (120 dxa), left/right: 8pt (160 dxa) per Rule 5."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('w:top', top), ('w:bottom', bottom),
                       ('w:left', left), ('w:right', right)]:
        node = OxmlElement(side)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)


def set_cell_border(cell, color="BFBFBF", sz="4"):
    """Light-gray border on all cell sides."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['w:top', 'w:left', 'w:bottom', 'w:right']:
        el = OxmlElement(side)
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), sz)
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def set_table_width(table, width_cm):
    """Set explicit table width in DXA (per docx skill: always use DXA, never percentage)."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    # Remove any existing tblW element
    for old in tblPr.findall(qn('w:tblW')):
        tblPr.remove(old)
    tblW = OxmlElement('w:tblW')
    dxa = int(width_cm / 2.54 * 1440)  # cm -> DXA (1440 DXA = 1 inch = 2.54 cm)
    tblW.set(qn('w:w'), str(dxa))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)


def split_cell_into_blocks(text):
    text = re.sub(r'<br\s*/?>', '\n', text)
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    
    blocks = []
    current_block = []
    
    new_block_triggers = [
        r'^\*\*chuyển giao nhiệm vụ',
        r'^\*\*thực hiện nhiệm vụ',
        r'^\*\*kết luận',
        r'^\*\*báo cáo',
        r'^\*\*hướng dẫn',
        r'^\*\*đánh giá',
        r'^\*\*thực hiện',
        r'^\*\*thảo luận',
        r'^\*\*nội dung',
        r'^\*\*bước',
        r'^- gv đặt câu hỏi',
        r'^- gv hỏi',
        r'^- thảo luận cặp đôi',
        r'^- thảo luận nhóm',
        r'^- thảo luận nhanh',
        r'^- học sinh thảo luận',
        r'^- suy nghĩ',
        r'^- học sinh suy nghĩ',
        r'^- trả lời',
        r'^- học sinh trả lời',
        r'^- báo cáo',
        r'^- học sinh báo cáo'
    ]
    
    for line in lines:
        is_trigger = False
        line_lower = line.lower()
        for trig in new_block_triggers:
            if re.match(trig, line_lower):
                is_trigger = True
                break
                
        if is_trigger:
            if current_block:
                only_headers = all(l.startswith('**') for l in current_block)
                is_header = line.startswith('**')
                if only_headers and is_header:
                    current_block.append(line)
                else:
                    blocks.append('\n'.join(current_block))
                    current_block = [line]
            else:
                current_block.append(line)
        else:
            current_block.append(line)
            
    if current_block:
        blocks.append('\n'.join(current_block))
        
    return blocks


# ============================================================
#  TEXT & TABLE HELPERS (Dọn dẹp code trùng lặp và nâng cấp)
# ============================================================

def _fill_cell_markdown(cell, val):
    """
    Điền văn bản có định dạng Markdown và LaTeX vào ô bảng.
    Hỗ trợ in đậm, in nghiêng, công thức latex, bullet list, blockquote.
    Áp dụng smart_typography và tránh dùng bullet thô (•).
    """
    cell.text = ""
    val_str = str(val)
    val_str = val_str.replace("<br>", "\n").replace("<br/>", "\n").replace("</br>", "\n")
    lines = val_str.split("\n")
    
    first_p = True
    for line in lines:
        line_text = line.strip()
        if not line_text:
            continue
            
        if first_p:
            p = cell.paragraphs[0]
            first_p = False
        else:
            p = cell.add_paragraph()
            
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.15
        
        # Check blockquote (dấu >)
        is_blockquote = False
        if line_text.startswith('> '):
            line_text = line_text[2:]
            is_blockquote = True
        elif line_text.startswith('>'):
            line_text = line_text[1:]
            is_blockquote = True
            
        if is_blockquote:
            p.paragraph_format.left_indent = Cm(0.5)
            
        # Check bullet
        bullet = ""
        if line_text.startswith('- ') or line_text.startswith('* '):
            bullet = "- "
            line_text = line_text[2:]
            p.paragraph_format.left_indent = Cm(0.3)
            p.paragraph_format.first_line_indent = Cm(-0.2)
        elif line_text.startswith('+ '):
            bullet = "+ "
            line_text = line_text[2:]
            p.paragraph_format.left_indent = Cm(0.5)
            p.paragraph_format.first_line_indent = Cm(-0.2)
            
        # Parse **bold** và $latex$
        parts = re.split(r'(\*\*.*?\*\*|\$.*?\$)', line_text)
        first_run = True
        for part in parts:
            if not part:
                continue
            if part.startswith('**') and part.endswith('**'):
                run_text = smart_typography(part[2:-2])
                run = p.add_run(run_text)
                run.bold = True
                # CV5512: Auto-bold Navy color for activity keywords inside table
                _BOLD_KEYWORDS = [
                    "Chuyển giao nhiệm vụ:", "Thực hiện nhiệm vụ:", "Báo cáo, thảo luận:", 
                    "Kết luận:", "Kết luận, nhận định:", "Sản phẩm:", "Hỗ trợ:", 
                    "Mục tiêu:", "Nội dung:"
                ]
                if any(kw in run_text for kw in _BOLD_KEYWORDS):
                    run.font.color.rgb = GEMSColors.NAVY
            elif part.startswith('$') and part.endswith('$'):
                formula = clean_latex(part[1:-1])
                run = p.add_run(formula)
                run.italic = True
            else:
                run_text = smart_typography(part)
                if bullet and first_run:
                    r_bullet = p.add_run(bullet)
                    r_bullet.font.name = GEMSFonts.BODY
                    r_bullet.font.size = GEMSFonts.SIZE_TABLE
                    r_bullet.font.color.rgb = GEMSColors.DARK
                    first_run = False
                run = p.add_run(run_text)
                if is_blockquote:
                    run.italic = True
            run.font.name = GEMSFonts.BODY
            run.font.size = GEMSFonts.SIZE_TABLE
            run.font.color.rgb = GEMSColors.DARK


def make_navy_table(doc, headers, rows, col_widths=None):
    """
    Tạo bảng phong cách GEMS:
      - Dòng tiêu đề: Nền Navy (#1E3A5F), chữ trắng in đậm, căn giữa.
      - Các dòng dữ liệu: Màu trắng và Mint (#E8F5E9) xen kẽ.
      - Viền xám nhạt, cell padding chuẩn.
      - Tự động split block GV/HS cho bảng 2 cột Hoạt động dạy học.
      - Set explicit table width và cell width (Dual Widths Rule 11).
      - Ngăn dòng bị cắt đôi khi sang trang (cantSplit).
      - Lặp lại dòng tiêu đề khi sang trang mới (tblHeader).
    """
    is_gv_hs_table = False
    if len(headers) == 2:
        h0, h1 = headers[0].lower(), headers[1].lower()
        if ('giáo viên' in h0 or 'gv' in h0) and ('học sinh' in h1 or 'hs' in h1):
            is_gv_hs_table = True
            
    processed_rows = []
    if is_gv_hs_table:
        for row_data in rows:
            if len(row_data) >= 2:
                gv_blocks = split_cell_into_blocks(row_data[0])
                hs_blocks = split_cell_into_blocks(row_data[1])
                max_blocks = max(len(gv_blocks), len(hs_blocks))
                for idx in range(max_blocks):
                    gv_val = gv_blocks[idx] if idx < len(gv_blocks) else ""
                    hs_val = hs_blocks[idx] if idx < len(hs_blocks) else ""
                    processed_rows.append([gv_val, hs_val])
            else:
                processed_rows.append(row_data)
    else:
        processed_rows = rows

    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = False  # Bắt buộc phải là False để áp dụng width cụ thể
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Thiết lập độ rộng bảng (A4 nội dung: 16.5cm)
    set_table_width(table, 16.5)

    # Tính toán độ rộng từng cột (cm)
    if col_widths:
        widths_in_cm = [col_widths[i] if i < len(col_widths) else 16.5 / len(headers) for i in range(len(headers))]
    else:
        widths_in_cm = [16.5 / len(headers)] * len(headers)

    # -- header row --
    hdr = table.rows[0].cells
    for i, text in enumerate(headers):
        cell = hdr[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(smart_typography(str(text)))
        run.bold = True
        run.font.name = GEMSFonts.BODY
        run.font.size = GEMSFonts.SIZE_TABLE
        run.font.color.rgb = GEMSColors.WHITE
        
        # Dual width: set cell.width
        cell.width = Cm(widths_in_cm[i])
        
        set_cell_shading(cell, GEMSColors.NAVY_HEX)
        set_cell_margins(cell)
        set_cell_border(cell)

    # Set tblHeader to repeat header on new page (Rule 6)
    trPr_hdr = table.rows[0]._tr.get_or_add_trPr()
    trPr_hdr.append(OxmlElement('w:tblHeader'))

    # -- data rows --
    for r_idx, row_data in enumerate(processed_rows):
        row = table.add_row()
        # Set cantSplit to prevent row breaking across pages (Rule 6)
        trPr = row._tr.get_or_add_trPr()
        trPr.append(OxmlElement('w:cantSplit'))
        
        while len(row_data) < len(headers):
            row_data.append("")
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            
            # Dual width: set cell.width
            cell.width = Cm(widths_in_cm[c_idx])
            
            _fill_cell_markdown(cell, val)
            set_cell_margins(cell)
            set_cell_border(cell)
            if r_idx % 2 == 1:
                set_cell_shading(cell, GEMSColors.MINT_HEX)

    doc.add_paragraph()  # spacer
    return table


# ============================================================
#  TEXT HELPERS
# ============================================================
def add_heading(doc, text, level=1):
    """
    Add a navy (H1, H2) or gray (H3) heading.
    level 1 = centered 16 pt  - used for main title (CV5512: Auto-UPPERCASE)
    level 2 = left 14 pt      - used for sections
    level 3 = left 13 pt      - used for sub-sections
    """
    p = doc.add_paragraph()
    sizes = {1: GEMSFonts.SIZE_H1, 2: GEMSFonts.SIZE_H2, 3: GEMSFonts.SIZE_H3}
    spaces_before = {1: Pt(18), 2: Pt(14), 3: Pt(10)}
    spaces_after = {1: Pt(8), 2: Pt(6), 3: Pt(4)}

    p.paragraph_format.space_before = spaces_before.get(level, Pt(10))
    p.paragraph_format.space_after = spaces_after.get(level, Pt(4))
    
    display_text = text.upper() if level == 1 else text
    run = p.add_run(display_text)
    run.bold = True
    run.font.size = sizes.get(level, GEMSFonts.SIZE_BODY)
    run.font.name = GEMSFonts.HEADER

    if level == 1:
        run.font.color.rgb = GEMSColors.NAVY
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 2:
        run.font.color.rgb = GEMSColors.NAVY
    else:
        run.font.color.rgb = GEMSColors.GRAY

    return p


def add_body(doc, text, bold=False, italic=False, size=None):
    """Standard body paragraph, Times New Roman, Justify, first line indent 1cm, line spacing 1.15, space after 4pt (Rule 1)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY   # CV5512: căn đều hai bên
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.first_line_indent = Cm(1.0)
    run = p.add_run(text)
    run.font.name = GEMSFonts.BODY
    run.font.size = size or GEMSFonts.SIZE_BODY
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = GEMSColors.DARK
    return p


def add_dot_line(doc, length=90):
    """Dotted answer line, gray."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("." * length)
    run.font.name = GEMSFonts.BODY
    run.font.size = GEMSFonts.SIZE_SMALL
    run.font.color.rgb = GEMSColors.LIGHT_GRAY
    return p


def add_separator(doc):
    """Horizontal gray rule."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("—" * 70)
    run.font.color.rgb = GEMSColors.LIGHT_GRAY
    run.font.size = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_student_info_block(doc, lesson_label):
    """
    Top-of-worksheet block: lesson title + name / class / date fields + Subtitle row (Rule 3).
    Returns the created table.
    """
    # Main title
    add_heading(doc, lesson_label, level=1)

    # Info row (3 cells, borderless) (Rule 3)
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    # Set explicit column widths to prevent text wrapping (Total: 16.5cm)
    widths = [Cm(8.5), Cm(3.2), Cm(4.8)]
    labels = [
        "Họ và tên: .......................................",
        "Lớp: ........................",
        "Ngày: ...... / ...... / 2026"
    ]
    
    for i, label in enumerate(labels):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(label)
        r.font.name = GEMSFonts.BODY
        r.font.size = GEMSFonts.SIZE_INFO
        r.font.color.rgb = GEMSColors.DARK
        
        # Apply width
        cell.width = widths[i]
        
        # Make borderless
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for side in ['w:top', 'w:left', 'w:bottom', 'w:right']:
            el = OxmlElement(side)
            el.set(qn('w:val'), 'none')
            tcBorders.append(el)
        tcPr.append(tcBorders)

    # Subtitle line (Rule 3)
    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_before = Pt(4)
    p_sub.paragraph_format.space_after = Pt(4)
    p_sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    r_sub = p_sub.add_run("Thời gian: 45 phút\t\tĐiểm: ....................\t\tÝ kiến của Giáo viên: ................................................................")
    r_sub.font.name = GEMSFonts.BODY
    r_sub.font.size = GEMSFonts.SIZE_INFO
    r_sub.font.color.rgb = GEMSColors.GRAY
    
    # Add separator rule
    add_separator(doc)

    return table


# ============================================================
#  SAVE
# ============================================================
def smart_typography(text):
    """
    Convert straight quotes to smart/curly quotes for professional typography
    and normalize Vietnamese punctuation spaces and physical units.
    """
    import re
    # Convert straight quotes to smart quotes
    text = re.sub(r'"(\S)', '\u201c\\1', text)  # Left double quote
    text = re.sub(r'(\S)"', '\\1\u201d', text)  # Right double quote
    text = re.sub(r'"', '\u201d', text)           # Remaining double quotes
    text = text.replace("'", '\u2019')
    text = re.sub(r"'(\S)", '\u2018\\1', text)
    text = re.sub(r"(\S)'", '\\1\u2019', text)
    
    # 1. Chuẩn hóa khoảng trắng sau dấu câu: .,;:!? (nếu đứng sát chữ cái)
    # Tránh làm hỏng số thập phân (như 0,5 hoặc 1.2)
    # Dấu phẩy sát chữ: ví dụ "đó,và" -> "đó, và"
    text = re.sub(r'([,;:!?])([A-Za-zĂăÂâĐđÊêÔôƠơƯưÁáÀàẢảÃãẠạẤấẦầẨẩẪẫẬậẮắẰằẲẳẴẵẶặÉéÈèẺẻẼẽẸẹẾếỀềỂểỄễỆệÍíÌìỈỉĨĩỊịÓóÒòỎỏÕõỌọỐốỒồỔổỖỗỘộỚớỜờỞởỠỡỢợÚúÙùỦủŨũỤụỨứỪừỬửỮữỰựÝýỲỳỶỷỸỹỴỵ])', r'\1 \2', text)
    # Dấu chấm sát chữ (nhưng không phải số thập phân, ví dụ "vật.Thế" -> "vật. Thế")
    text = re.sub(r'(\.[^\d\s\)])([A-Za-zĂăÂâĐđÊêÔôƠơƯưÁáÀàẢảÃãẠạẤấẦầẨẩẪẫẬậẮắẰằẲẳẴẵẶặÉéÈèẺẻẼẽẸẹẾếỀềỂểỄễỆệÍíÌìỈỉĨĩỊịÓóÒòỎỏÕõỌọỐốỒồỔổỖỗỘộỚớỜờỞởỠỡỢợÚúÙùỦủŨũỤụỨứỪừỬửỮữỰựÝýỲỳỶỷỸỹỴỵ])', r'\1 \2', text)
    
    # 2. Xóa khoảng trắng thừa trước dấu câu (như "động ." -> "động.")
    text = re.sub(r'\s+([,.;:!?])', r'\1', text)

    # 3. Định dạng khoảng trắng giữa số và đơn vị vật lý (ví dụ: "5cm" -> "5 cm")
    text = re.sub(r'(\d)\s*(cm|dm|mm|kg|Hz|Pa|rad/s|rad|m/s²|m/s|m|s|kg|g|kJ|J|W|kW|V|A|N|°C|K)\b', r'\1 \2', text)
    
    return text


def save_doc(doc, output_path):
    """Create parent directory if needed and save."""
    d = os.path.dirname(output_path)
    if d and not os.path.exists(d):
        os.makedirs(d, exist_ok=True)
    doc.save(output_path)
    return output_path


def clean_latex(text):
    """Clean LaTeX math tokens to Unicode for professional Word outputs."""
    import re
    # Remove HTML <br> tags
    text = re.sub(r'(?i)<br\s*/?>', ' ', text)
    
    # 1. Convert LaTeX superscripts (e.g. 10^6, 10^{-3}, 10^-2) to Unicode superscript characters
    def to_superscript(num_str):
        sups = {
            '0': '⁰', '1': '¹', '2': '²', '3': '³', '4': '⁴',
            '5': '⁵', '6': '⁶', '7': '⁷', '8': '⁸', '9': '⁹',
            '-': '⁻', '+': '⁺'
        }
        return "".join(sups.get(c, c) for c in num_str)

    def replace_sup(match):
        return to_superscript(match.group(1))

    # Pattern matches ^6 or ^{-6} or ^-6
    text = re.sub(r'\^\{?([+-]?\d+)\}?', replace_sup, text)

    # 2. Replace standard LaTeX symbols
    replacements = {
        r'\lambda': 'λ',
        r'\Delta': 'Δ',
        r'\delta': 'δ',
        r'\omega': 'ω',
        r'\varphi': 'φ',
        r'\pi': 'π',
        r'\cos': 'cos',
        r'\sin': 'sin',
        r'\times': '×',
        r'\approx': '≈',
        r'\cdot': '·',
        r'^\circ': '°',
        r'\circ': '°',
        r'^{\circ}C': '°C',
        r'^{\circ} C': '°C',
        r'^\circ C': '°C',
        r'\text{J/kg.K}': 'J/kg.K',
        r'\text{J/kg.}^\circ\text{C}': 'J/kg.°C',
        r'\text{J/kg}': 'J/kg',
        r'\text{ J/kg}': ' J/kg',
        r'\text{ W}': ' W',
        r'\text{ s}': ' s',
        r'\text{ J}': ' J',
        r'\text{ kJ}': ' kJ',
        r'\text{C}': 'C',
        r'\text{K}': 'K',
        r'\text{kg}': 'kg',
        r'_{nước}': '_nước',
        r'_{nhôm}': '_nhôm',
        r'_{đá\_tan}': '_đá_tan',
        r'_{đá}': '_đá',
        r'_{nước\_đá}': '_nước_đá',
        r'_{tn}': '_tn',
        r'\_': '_',
        r'\ ': ' ',
    }
    # Apply static replacements
    for k, v in replacements.items():
        text = text.replace(k, v)
    
    # Remove any remaining \text{...}
    text = re.sub(r'\\text\{([^}]+)\}', r'\1', text)
    # Remove any remaining backslashes
    text = text.replace('\\', '')
    # Remove all raw $ characters remaining in the final text (solves the '$0°C$' problem)
    text = text.replace('$', '')
    return text


def add_image(doc, img_path, caption):
    """
    Embed an image centered in the document with accessibility alt text (wp:docPr) 
    and an italicized center-aligned caption below it. (Shared GEMS helper)
    """
    if not os.path.exists(img_path):
        print(f"[WARNING] Image path not found: {img_path}")
        return False
    try:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run()
        run.add_picture(img_path, width=Inches(4.5))

        # Add Alt text for accessibility XML compliance (Rule 12)
        try:
            from docx.oxml.ns import qn as _qn
            inline = run._r.find('.//' + _qn('wp:inline'))
            if inline is not None:
                alt_title = caption or os.path.basename(img_path)
                inline.set('distT', '0'); inline.set('distB', '0')
                inline.set('distL', '0'); inline.set('distR', '0')
                docPr = inline.find(_qn('wp:docPr'))
                if docPr is None:
                    from docx.oxml import OxmlElement as _OxmlElement
                    docPr = _OxmlElement('wp:docPr')
                    docPr.set('id', '1')
                    docPr.set('name', alt_title)
                    docPr.set('descr', alt_title)
                    inline.insert(0, docPr)
                else:
                    docPr.set('name', alt_title)
                    docPr.set('descr', alt_title)
        except Exception:
            pass  # Alt text failure is non-critical

        # Add Caption
        if caption:
            p_cap = doc.add_paragraph()
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_cap.paragraph_format.space_after = Pt(8)
            run_cap = p_cap.add_run(f"Hình: {caption}")
            run_cap.italic = True
            run_cap.font.name = GEMSFonts.BODY
            run_cap.font.size = Pt(10)
            run_cap.font.color.rgb = GEMSColors.GRAY
        return True
    except Exception as e:
        print(f"[ERROR] Failed to insert image {img_path}: {e}")
        return False


def preprocess_markdown_lines(lines):
    """
    Clean up lines that are split due to math formulas or variables on their own lines.
    Applies standard paragraph folding combined with forced folding for formulas and connective symbols.
    """
    current_lines = [l.rstrip('\n') for l in lines]
    while True:
        changed = False
        new_lines = []
        i = 0
        while i < len(current_lines):
            line = current_lines[i]
            if i + 1 < len(current_lines):
                next_idx = i + 1
                while next_idx < len(current_lines) and not current_lines[next_idx].strip():
                    next_idx += 1
                
                if next_idx < len(current_lines):
                    nline = current_lines[next_idx].strip()
                    curr_strip = line.strip()
                    
                    if curr_strip and nline:
                        # 1. Connective and type checks
                        first_char = nline[0]
                        is_start_lower = first_char.islower()
                        is_connective = first_char in [')', ',', '.', ';', ':', ']', '%', '$']
                        is_current_open = curr_strip.endswith('(') or curr_strip.endswith('[') or curr_strip.endswith('$')
                        
                        # Short formula detection
                        is_formula = not any(c in 'áàảãạăắằẳẵặâấầẩẫậéèẻẽẹêếềểễệíìỉĩịóòỏõọôốồổỗộơớờởỡợúùủũụưứừửữựýỳỷỹỵđ' for c in nline.lower()) and len(nline) < 45 and len(nline) > 0
                        
                        # 2. Block structural elements
                        is_curr_heading_or_table = (curr_strip.startswith('#') or 
                                                    curr_strip.startswith('|') or 
                                                    curr_strip in ['---', '___'])
                                                    
                        is_next_q = re.match(r'^\*?\*?(?:Câu\s+)?\d+(?:[\.\)]|(?:\s*\([A-Z]+\))?:?\*?\*?)', nline)
                        is_next_opt = re.match(r'^\*?\*?[A-Da-d][\.\)]', nline)

                        is_next_special_all = (nline.startswith('#') or 
                                               nline.startswith('- ') or 
                                               nline.startswith('+ ') or 
                                               nline.startswith('* ') or 
                                               nline.startswith('|') or
                                               nline.startswith('![') or
                                               nline in ['---', '___'] or
                                               is_next_q or is_next_opt or
                                               (nline[0].isdigit() and '. ' in nline[:5]))
                        
                        is_curr_list = (curr_strip.startswith('- ') or 
                                        curr_strip.startswith('+ ') or 
                                        curr_strip.startswith('* ') or 
                                        re.match(r'^\*?\*?(?:Câu\s+)?\d+(?:[\.\)]|(?:\s*\([A-Z]+\))?:?\*?\*?)', curr_strip) or
                                        re.match(r'^\*?\*?[A-Da-d][\.\)]', curr_strip) or
                                        (curr_strip[0].isdigit() and '. ' in curr_strip[:5]))
                        
                        # 3. Merging logic
                        has_no_empty_between = (next_idx == i + 1)
                        
                        # Normal merge: consecutive lines, neither is a structural element
                        merge_normal = (has_no_empty_between and 
                                        not is_curr_heading_or_table and not is_curr_list and 
                                        not is_next_special_all)
                        
                        # Forced merge: even with blank lines, merge if next line continues the current sentence/formula
                        merge_forced = (not is_curr_heading_or_table and 
                                        not is_next_special_all and 
                                        (is_start_lower or is_connective or is_current_open or is_formula))
                        
                        if merge_normal or merge_forced:
                            # Merge line with next non-empty line
                            merged_line = line.rstrip() + " " + nline
                            new_lines.append(merged_line)
                            i = next_idx + 1
                            changed = True
                            continue
            new_lines.append(line)
            i += 1
        
        current_lines = new_lines
        if not changed:
            break
            
    return [l + "\n" for l in current_lines]


def add_page_number_footer(doc):
    """Add 'Trang X / Y' page numbering centered in the footer of all sections."""
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.text = ""  # Clear default
        
        # Add "Trang "
        r1 = p.add_run("Trang ")
        r1.font.name = GEMSFonts.BODY
        r1.font.size = GEMSFonts.SIZE_TINY
        r1.font.color.rgb = GEMSColors.GRAY
        
        # Add PAGE field
        fldSimple_page = OxmlElement('w:fldSimple')
        fldSimple_page.set(qn('w:instr'), 'PAGE')
        p._p.append(fldSimple_page)
        
        # Add " / "
        r2 = p.add_run(" / ")
        r2.font.name = GEMSFonts.BODY
        r2.font.size = GEMSFonts.SIZE_TINY
        r2.font.color.rgb = GEMSColors.GRAY
        
        # Add NUMPAGES field
        fldSimple_numpages = OxmlElement('w:fldSimple')
        fldSimple_numpages.set(qn('w:instr'), 'NUMPAGES')
        p._p.append(fldSimple_numpages)


# ============================================================
#  FORMATTED TEXT RUN HELPERS (LaTeX + Bold processing)
# ============================================================

def add_formatted_runs(paragraph, text, force_italic=False, font_size=None, is_bold_all=False):
    """
    Parse **bold** and $latex$ in text and add structured runs to the paragraph.
    Applies clean_latex and smart_typography individually to runs.
    """
    parts = re.split(r'(\*\*.*?\*\*|\$.*?\$)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run_text = smart_typography(clean_latex(part[2:-2]))
            run = paragraph.add_run(run_text)
            run.bold = True
            run.italic = force_italic
        elif part.startswith('$') and part.endswith('$'):
            formula = clean_latex(part[1:-1])
            run = paragraph.add_run(formula)
            run.italic = True
        else:
            run_text = smart_typography(clean_latex(part))
            run = paragraph.add_run(run_text)
            run.bold = is_bold_all
            run.italic = force_italic
            
        run.font.name = GEMSFonts.BODY
        if font_size:
            run.font.size = font_size
        else:
            run.font.size = GEMSFonts.SIZE_BODY
        run.font.color.rgb = GEMSColors.DARK


def add_body_formatted(doc, text, is_blockquote=False, size=None):
    """Add a body paragraph with proper spacing and formatted runs."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    if is_blockquote:
        p.paragraph_format.left_indent = Cm(1.0)
    else:
        p.paragraph_format.first_line_indent = Cm(1.0)
        
    add_formatted_runs(p, text, force_italic=is_blockquote, font_size=size)
    return p


def add_bullet_formatted(doc, text, is_blockquote=False, level=1):
    """Add a bullet item with clean indentation and formatted runs."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    
    indent_cm = 0.5 if level == 1 else 1.0
    p.paragraph_format.left_indent = Cm(indent_cm)
    p.paragraph_format.first_line_indent = Cm(-0.3)
    
    bullet_str = "- " if level == 1 else "+ "
    r_bullet = p.add_run(bullet_str)
    r_bullet.font.name = GEMSFonts.BODY
    r_bullet.font.size = GEMSFonts.SIZE_BODY
    r_bullet.font.color.rgb = GEMSColors.DARK
    
    # Clean prefix bullets
    if text.startswith('- ') or text.startswith('* '):
        text = text[2:]
    elif text.startswith('+ '):
        text = text[2:]
        
    add_formatted_runs(p, text, force_italic=is_blockquote)
    return p