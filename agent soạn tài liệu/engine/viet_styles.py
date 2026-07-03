# -*- coding: utf-8 -*-
"""
viet_styles.py — Vietnamese Educational Document Formatting Standard
===================================================================
Implementation of Microsoft Word styling guidelines for Vietnamese educational documents.
Provides clean fonts, standardized margins, and custom helpers for worksheets, lesson plans, and exams.
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
import os
import re

# ============================================================
#  COLOUR PALETTE (Vietnamese Edu Standards)
# ============================================================
class VietColors:
    PRIMARY = RGBColor(0x1F, 0x4E, 0x79)        # #1F4E79 Steel Blue for Title/H1
    DARK = RGBColor(0x00, 0x00, 0x00)           # #000000 Black body text
    GRAY = RGBColor(0x66, 0x66, 0x66)           # #666666
    LIGHT_GRAY = RGBColor(0xD9, 0xD9, 0xD9)     # #D9D9D9 For Table Headers
    BORDER_GRAY = RGBColor(0xBF, 0xBF, 0xBF)    # #BFBFBF
    
    PRIMARY_HEX = "1F4E79"
    LIGHT_GRAY_HEX = "D9D9D9"
    WHITE_HEX = "FFFFFF"
    BORDER_GRAY_HEX = "BFBFBF"

# ============================================================
#  FONTS
# ============================================================
class VietFonts:
    BODY = "Times New Roman"
    HEADER = "Times New Roman"
    
    SIZE_TITLE = Pt(16)
    SIZE_H1 = Pt(14)
    SIZE_H2 = Pt(13)
    SIZE_H3 = Pt(13)
    SIZE_BODY = Pt(13)
    SIZE_TABLE = Pt(12)
    SIZE_INFO = Pt(11)
    SIZE_TINY = Pt(9)

# ============================================================
#  PAGE SETUP & MARGINS
# ============================================================
def setup_viet_margins(doc, doc_type="general"):
    """
    Set page margins according to Vietnamese school guidelines:
    - Worksheet (pht): Top/Bottom: 2cm, Left/Right: 2cm
    - Lesson Plan (khbd): Top/Bottom: 2cm, Left: 3cm (for binding), Right: 2cm
    - Exams (homework): Top/Bottom: 1.5cm, Left/Right: 2cm
    """
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        
        if doc_type == "pht":
            section.top_margin = Cm(2.0)
            section.bottom_margin = Cm(2.0)
            section.left_margin = Cm(2.0)
            section.right_margin = Cm(2.0)
        elif doc_type == "khbd":
            section.top_margin = Cm(2.0)
            section.bottom_margin = Cm(2.0)
            section.left_margin = Cm(3.0)
            section.right_margin = Cm(2.0)
        elif doc_type == "homework":
            section.top_margin = Cm(1.5)
            section.bottom_margin = Cm(1.5)
            section.left_margin = Cm(2.0)
            section.right_margin = Cm(2.0)
        else:
            # General default
            section.top_margin = Cm(2.0)
            section.bottom_margin = Cm(2.0)
            section.left_margin = Cm(3.0)
            section.right_margin = Cm(2.0)

def set_viet_default_style(doc):
    """Set default font style: Times New Roman 13 pt, 1.3 lines spacing, Justified, 6pt before/after."""
    style = doc.styles['Normal']
    font = style.font
    font.name = VietFonts.BODY
    font.size = VietFonts.SIZE_BODY
    font.color.rgb = VietColors.DARK
    
    pf = style.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.3
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# ============================================================
#  TABLE FORMATTING UTILITIES
# ============================================================
def set_cell_shading(cell, color_hex):
    """Set cell background shading."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top_pt=5.7, bottom_pt=5.7, left_pt=8.5, right_pt=8.5):
    """Set padding (margins) inside cells (Top/Bottom ≈ 0.1cm, Left/Right ≈ 0.15cm)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    
    # conversion factor: pt to dxa (1 pt = 20 dxa)
    for m_type, pt_val in [('top', top_pt), ('bottom', bottom_pt), ('left', left_pt), ('right', right_pt)]:
        node = OxmlElement(f'w:{m_type}')
        node.set(qn('w:w'), str(int(pt_val * 20)))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_borders(cell):
    """Apply thin borders (0.5 pt solid gray) on all sides of the cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    
    for side in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')  # 4/8 pt = 0.5 pt
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), VietColors.BORDER_GRAY_HEX)
        tcBorders.append(border)
    tcPr.append(tcBorders)

def set_table_width_dxa(table, width_cm):
    """Set table width using dxa units (never use autofit/percentages)."""
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # CM to Twips/DXA: 1 cm = 567 twips
    width_dxa = int(width_cm * 567)
    tblPr = table._tbl.tblPr
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(width_dxa))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)

# ============================================================
#  TEXT CLEANUP HELPER FUNCTIONS
# ============================================================
def clean_latex(text):
    """Translate LaTeX markup and math variables to clean Unicode equivalent glyphs."""
    if not text:
        return ""
        
    replacements = {
        r'\Delta': 'Δ', r'\delta': 'δ',
        r'\omega': 'ω', r'\Omega': 'Ω',
        r'\varphi': 'φ', r'\phi': 'φ',
        r'\pi': 'π',
        r'\alpha': 'α', r'\beta': 'β', r'\gamma': 'γ',
        r'\lambda': 'λ', r'\mu': 'μ', r'\theta': 'θ',
        r'\rho': 'ρ', r'\tau': 'τ',
        r'\approx': '≈', r'\neq': '≠', r'\le': '≤', r'\ge': '≥',
        r'\times': '×', r'\cdot': '·', r'\pm': '±',
        r'\circ': '°', r'^{\circ}': '°', r'^\circ': '°',
        r'\infty': '∞', r'\to': '→',
        r'\text{ J}': ' J', r'\text{ J/kg.K}': ' J/kg.K', r'\text{ J/kg.}^{\circ}C}': ' J/kg.°C',
        r'\text{ kg}': ' kg', r'\text{ g}': ' g', r'\text{ m}': ' m', r'\text{ s}': ' s',
        r'\text{ W}': ' W', r'\text{ K}': ' K', r'\text{ kJ}': ' kJ', r'\text{ MJ}': ' MJ',
        r'\text{}^{\circ}C': '°C', r'\text{}^{\circ}': '°',
        r'^\circ C': '°C', r'^\circ\text{C}': '°C',
        r'\Delta t': 'Δt', r'\Delta U': 'ΔU'
    }
    
    # Thay thế các ký hiệu LaTeX
    for key, val in replacements.items():
        text = text.replace(key, val)
        
    # Clean superscript superscript LaTeX patterns (e.g. 10^5 -> 10⁵, 10^{-3} -> 10⁻³)
    superscript_map = {
        '0': '⁰', '1': '¹', '2': '²', '3': '³', '4': '⁴',
        '5': '⁵', '6': '⁶', '7': '⁷', '8': '⁸', '9': '⁹',
        '-': '⁻', '+': '⁺', '=': '⁼', '(': '⁽', ')': '⁾',
        'n': 'ⁿ', 'x': 'ˣ', 'i': 'ⁱ'
    }
    
    # 10^{-x} or 10^-x pattern
    def repl_super(match):
        chars = match.group(1) or match.group(2)
        return "".join(superscript_map.get(c, c) for c in chars)
        
    text = re.sub(r'\^\{([^}]+)\}', repl_super, text)
    text = re.sub(r'\^([0-9a-zA-Z\-\+\=]+)', repl_super, text)
    
    # Xóa sạch các dấu $ còn lại
    text = text.replace('$', '')
    return text

def smart_typography(text):
    """Convert straight quotes to curly/smart quotes."""
    if not text:
        return ""
    # Thay thế nháy kép thẳng thành nháy kép cong
    text = re.sub(r'"([^"]*)"', r'“\1”', text)
    # Thay thế nháy đơn thẳng thành nháy đơn cong
    text = re.sub(r"'([^']*)'", r'‘\1’', text)
    # Thay thế dấu apostrophe cho các từ tiếng Anh (nếu có)
    text = re.sub(r'(\w)\'(\w)', r'\1’\2', text)
    return text

# ============================================================
#  FORMATTED RUN PARSER & EXPORTER
# ============================================================
def add_formatted_runs(paragraph, text, force_italic=False, font_size=None, is_bold_all=False):
    """
    Parse **bold** and $latex$ chunks in text and append runs to the paragraph.
    Clean LaTeX symbols and straight quotes on individual chunks.
    """
    parts = re.split(r'(\*\*.*?\*\*|\$.*?\$)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run_text = smart_typography(clean_latex(part[2:-2]))
            run = paragraph.add_run(run_text)
            run.bold = True
            run.italic = force_italic
        elif part.startswith('$') and part.endswith('$'):
            formula = clean_latex(part[1:-1])
            run = paragraph.add_run(formula)
            run.italic = True
        else:
            run_text = smart_typography(clean_latex(part))
            run = paragraph.add_run(run_text)
            run.bold = is_bold_all
            run.italic = force_italic
            
        run.font.name = VietFonts.BODY
        if font_size:
            run.font.size = font_size
        else:
            run.font.size = VietFonts.SIZE_BODY
        run.font.color.rgb = VietColors.DARK

# ============================================================
#  STANDARDIZED PARAGRAPH CREATORS
# ============================================================
def add_viet_title(doc, text):
    """Add a bold title centered, 16pt, primary color #1F4E79."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.keep_with_next = True
    
    run = p.add_run(smart_typography(text).upper())
    run.bold = True
    run.font.name = VietFonts.HEADER
    run.font.size = VietFonts.SIZE_TITLE
    run.font.color.rgb = VietColors.PRIMARY
    return p

def add_viet_heading(doc, text, level=1):
    """
    Add standard Heading 1, 2, or 3.
    Level 1 = 14pt Bold, color #1F4E79, before: 12pt, after: 6pt (used for I, II, III)
    Level 2 = 13pt Bold, color #000000, before: 8pt, after: 4pt  (used for 1, 2, 3)
    Level 3 = 13pt Italic, color #000000, before: 6pt, after: 3pt  (used for a, b, c)
    """
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    
    sizes = {1: VietFonts.SIZE_H1, 2: VietFonts.SIZE_H2, 3: VietFonts.SIZE_H3}
    spaces_before = {1: Pt(12), 2: Pt(8), 3: Pt(6)}
    spaces_after = {1: Pt(6), 2: Pt(4), 3: Pt(3)}
    
    p.paragraph_format.space_before = spaces_before.get(level, Pt(6))
    p.paragraph_format.space_after = spaces_after.get(level, Pt(4))
    
    # Assign paragraph heading styles (Heading 1, 2, 3)
    p.style = doc.styles[f'Heading {level}']
    
    p.paragraph_format.left_indent = Cm(0)
    p.paragraph_format.first_line_indent = Cm(0)
    
    display_text = clean_latex(text)
    if level == 1:
        display_text = display_text.upper()
        
    run = p.add_run(display_text)
    run.bold = (level in [1, 2])
    run.italic = (level == 3)
    run.font.name = VietFonts.HEADER
    run.font.size = sizes.get(level, VietFonts.SIZE_BODY)
    
    if level == 1:
        run.font.color.rgb = VietColors.PRIMARY
    else:
        run.font.color.rgb = VietColors.DARK
        
    return p

def add_viet_body(doc, text, is_blockquote=False, size=None):
    """Add standard body paragraph, Justified, 1.3 spacing, first line indent 1cm, 6pt before/after."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.3
    
    if is_blockquote:
        p.paragraph_format.left_indent = Cm(1.0)
        p.paragraph_format.first_line_indent = Cm(0)
    else:
        p.paragraph_format.left_indent = Cm(0)
        p.paragraph_format.first_line_indent = Cm(1.0)
        
    add_formatted_runs(p, text, force_italic=is_blockquote, font_size=size)
    return p

def add_viet_bullet(doc, text, is_blockquote=False, level=1):
    """Add a bullet list paragraph, no first line indent, left indent 0.5cm (lvl 1) or 1.0cm (lvl 2)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.3
    
    indent_cm = 0.5 if level == 1 else 1.0
    p.paragraph_format.left_indent = Cm(indent_cm)
    p.paragraph_format.first_line_indent = Cm(-0.3)
    
    bullet_symbol = "- " if level == 1 else "+ "
    run_bullet = p.add_run(bullet_symbol)
    run_bullet.font.name = VietFonts.BODY
    run_bullet.font.size = VietFonts.SIZE_BODY
    run_bullet.font.color.rgb = VietColors.DARK
    
    # Strip existing bullet markers in text
    if text.startswith('- ') or text.startswith('* '):
        text = text[2:]
    elif text.startswith('+ '):
        text = text[2:]
        
    add_formatted_runs(p, text, force_italic=is_blockquote)
    return p

# ============================================================
#  TABLES AND REPEATING HEADERS (Vietnamese standards)
# ============================================================
def make_viet_table(doc, headers, rows, col_widths=None, total_width_cm=17.0):
    """
    Create a professional table with:
      - Header Row: Background #D9D9D9 (gray), bold center text, 12pt.
      - Borders: solid gray 0.5pt.
      - Cell padding: 0.1cm top/bottom, 0.15cm left/right.
      - cantSplit on rows & tblHeader on row 0.
    """
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Set total table width
    set_table_width_dxa(table, total_width_cm)
    
    # Calculate column widths
    if col_widths:
        widths_cm = [col_widths[i] if i < len(col_widths) else total_width_cm / len(headers) for i in range(len(headers))]
    else:
        widths_cm = [total_width_cm / len(headers)] * len(headers)
        
    # Format Header Row
    hdr_cells = table.rows[0].cells
    for i, text in enumerate(headers):
        cell = hdr_cells[i]
        cell.width = Cm(widths_cm[i])
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        
        run = p.add_run(smart_typography(str(text)))
        run.bold = True
        run.font.name = VietFonts.BODY
        run.font.size = VietFonts.SIZE_TABLE
        run.font.color.rgb = VietColors.DARK
        
        set_cell_shading(cell, VietColors.LIGHT_GRAY_HEX)
        set_cell_margins(cell)
        set_cell_borders(cell)
        
    # Repeat header row on new page
    trPr_hdr = table.rows[0]._tr.get_or_add_trPr()
    trPr_hdr.append(OxmlElement('w:tblHeader'))
    
    # Add Data Rows
    for r_idx, row_data in enumerate(rows):
        row = table.add_row()
        # Prevent row break across pages
        trPr = row._tr.get_or_add_trPr()
        trPr.append(OxmlElement('w:cantSplit'))
        
        # Ensure row data length matches columns
        padded_row_data = list(row_data) + [""] * (len(headers) - len(row_data))
        
        for c_idx, val in enumerate(padded_row_data):
            cell = row.cells[c_idx]
            cell.width = Cm(widths_cm[c_idx])
            cell.text = ""
            
            # Populate text
            lines = str(val).replace("<br>", "\n").replace("<br/>", "\n").split("\n")
            first_p = True
            for line in lines:
                line_text = line.strip()
                if not line_text:
                    continue
                if first_p:
                    p = cell.paragraphs[0]
                    first_p = False
                else:
                    p = cell.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.15
                
                # Check list/bullet inside cell
                bullet_prefix = ""
                if line_text.startswith('- ') or line_text.startswith('* '):
                    bullet_prefix = "- "
                    line_text = line_text[2:]
                    p.paragraph_format.left_indent = Cm(0.3)
                    p.paragraph_format.first_line_indent = Cm(-0.2)
                elif line_text.startswith('+ '):
                    bullet_prefix = "+ "
                    line_text = line_text[2:]
                    p.paragraph_format.left_indent = Cm(0.5)
                    p.paragraph_format.first_line_indent = Cm(-0.2)
                
                # Split runs for cell text
                cell_parts = re.split(r'(\*\*.*?\*\*|\$.*?\$)', line_text)
                is_first_run = True
                for part in cell_parts:
                    if not part:
                        continue
                    if part.startswith('**') and part.endswith('**'):
                        r_text = smart_typography(part[2:-2])
                        run = p.add_run(r_text)
                        run.bold = True
                    elif part.startswith('$') and part.endswith('$'):
                        formula = clean_latex(part[1:-1])
                        run = p.add_run(formula)
                        run.italic = True
                    else:
                        r_text = smart_typography(part)
                        if bullet_prefix and is_first_run:
                            r_bullet = p.add_run(bullet_prefix)
                            r_bullet.font.name = VietFonts.BODY
                            r_bullet.font.size = VietFonts.SIZE_INFO
                            r_bullet.font.color.rgb = VietColors.DARK
                            is_first_run = False
                        run = p.add_run(r_text)
                        
                    run.font.name = VietFonts.BODY
                    run.font.size = VietFonts.SIZE_INFO
                    run.font.color.rgb = VietColors.DARK
                    
            set_cell_margins(cell)
            set_cell_borders(cell)
            
    doc.add_paragraph()  # spacer paragraph after table
    return table

# ============================================================
#  FOOTER AND PAGE NUMBERING
# ============================================================
def add_viet_page_number_footer(doc):
    """
    Insert a dynamic 'Trang X/Y' footer at the center bottom of each page.
    Uses MS Word dynamic XML field codes PAGE and NUMPAGES.
    """
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        
        p.text = "" # Clear default
        
        # Add "Trang "
        run1 = p.add_run("Trang ")
        run1.font.name = VietFonts.BODY
        run1.font.size = VietFonts.SIZE_TINY
        run1.font.color.rgb = VietColors.GRAY
        
        # Add dynamic PAGE XML field code
        fld_page = OxmlElement('w:fldSimple')
        fld_page.set(qn('w:instr'), 'PAGE')
        p._p.append(fld_page)
        
        # Add "/"
        run2 = p.add_run("/")
        run2.font.name = VietFonts.BODY
        run2.font.size = VietFonts.SIZE_TINY
        run2.font.color.rgb = VietColors.GRAY
        
        # Add dynamic NUMPAGES XML field code
        fld_numpages = OxmlElement('w:fldSimple')
        fld_numpages.set(qn('w:instr'), 'NUMPAGES')
        p._p.append(fld_numpages)

# ============================================================
#  MISCELLANEOUS HELPERS
# ============================================================
def add_viet_dot_line(doc, count=1):
    """Add a grey dotted line for handwritten answers."""
    for _ in range(count):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Cm(1.0)
        
        # Generate dots fitting page width (~90 dots for 2cm margins)
        run = p.add_run("." * 90)
        run.font.name = VietFonts.BODY
        run.font.size = VietFonts.SIZE_INFO
        run.font.color.rgb = VietColors.GRAY

def add_viet_separator(doc):
    """Add a light gray separator rule between major parts."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    
    run = p.add_run("—" * 35)
    run.font.name = VietFonts.BODY
    run.font.size = VietFonts.SIZE_TINY
    run.font.color.rgb = VietColors.BORDER_GRAY
