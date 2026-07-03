# -*- coding: utf-8 -*-
"""
viet_homework_exporter.py — Standardized Exam Exporter (Homework) → DOCX
========================================================================
Implementation of Vietnamese Educational Document Formatting Standard § 5.4
Exports high school graduation style homework (18 MCQ + 4 True/False + 6 Short Answer).
"""

import re
import os
from viet_styles import (
    Document, setup_viet_margins, set_viet_default_style,
    VietColors, VietFonts,
    add_viet_title, add_viet_heading, add_viet_body, add_viet_bullet,
    make_viet_table, add_viet_dot_line, add_viet_separator,
    add_viet_page_number_footer, add_formatted_runs,
    smart_typography, clean_latex, Cm, Pt, OxmlElement, qn
)
from docx.oxml import parse_xml

def add_viet_exam_header(doc):
    """Add standardized examination header and student details block (CTGDPT 2018)."""
    # Create 2-column layout for administrative titles
    table = doc.add_table(rows=1, cols=2)
    table.alignment = 1 # Center
    
    # 17.0cm width (Col 0: 9.5cm, Col 1: 7.5cm)
    col_widths = [Cm(9.5), Cm(7.5)]
    
    # Col 0: Agency details
    cell_lh = table.rows[0].cells[0]
    cell_lh.width = col_widths[0]
    cell_lh.text = ""
    p_lh = cell_lh.paragraphs[0]
    p_lh.paragraph_format.space_before = Pt(0)
    p_lh.paragraph_format.space_after = Pt(2)
    p_lh.paragraph_format.line_spacing = 1.15
    
    r1 = p_lh.add_run("SỞ GIÁO DỤC VÀ ĐÀO TẠO: ................................\n")
    r1.font.name = VietFonts.BODY
    r1.font.size = VietFonts.SIZE_INFO
    r1.font.color.rgb = VietColors.GRAY
    
    r2 = p_lh.add_run("TRƯỜNG THPT: ................................................\n")
    r2.bold = True
    r2.font.name = VietFonts.BODY
    r2.font.size = VietFonts.SIZE_INFO
    
    r3 = p_lh.add_run("Tổ Vật lí - Công nghệ")
    r3.font.name = VietFonts.BODY
    r3.font.size = VietFonts.SIZE_TINY
    r3.italic = True
    r3.font.color.rgb = VietColors.GRAY
    
    # Col 1: Exam info
    cell_rh = table.rows[0].cells[1]
    cell_rh.width = col_widths[1]
    cell_rh.text = ""
    p_rh = cell_rh.paragraphs[0]
    p_rh.alignment = 1 # Center
    p_rh.paragraph_format.space_before = Pt(0)
    p_rh.paragraph_format.space_after = Pt(2)
    p_rh.paragraph_format.line_spacing = 1.15
    
    r_ex = p_rh.add_run("ĐỀ BÀI TẬP VỀ NHÀ\n")
    r_ex.bold = True
    r_ex.font.name = VietFonts.BODY
    r_ex.font.size = VietFonts.SIZE_H1
    r_ex.font.color.rgb = VietColors.PRIMARY
    
    r_sub = p_rh.add_run("Môn: Vật lí 12 (CTGDPT 2018)\nThời gian làm bài: 50 phút")
    r_sub.font.name = VietFonts.BODY
    r_sub.font.size = VietFonts.SIZE_INFO
    r_sub.font.color.rgb = VietColors.DARK
    
    # Remove borders
    for cell in table.rows[0].cells:
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for side in ['top', 'left', 'bottom', 'right']:
            b = OxmlElement(f'w:{side}')
            b.set(qn('w:val'), 'none')
            tcBorders.append(b)
        tcPr.append(tcBorders)
        
    doc.add_paragraph() # spacing
    
    # Student metadata block (single line for exam style)
    p_stud = doc.add_paragraph()
    p_stud.paragraph_format.space_before = Pt(2)
    p_stud.paragraph_format.space_after = Pt(8)
    p_stud.paragraph_format.keep_with_next = True
    
    r_name = p_stud.add_run("Họ và tên thí sinh: .......................................................................... ")
    r_name.font.name = VietFonts.BODY
    r_name.font.size = VietFonts.SIZE_INFO
    
    r_sbd = p_stud.add_run("Số báo danh: .............................")
    r_sbd.font.name = VietFonts.BODY
    r_sbd.font.size = VietFonts.SIZE_INFO
    
    # Bottom border line
    add_viet_separator(doc)

def _add_viet_question_stem(doc, num, text, image_path=None):
    """Add question stem with keepNext constraints."""
    p = doc.add_paragraph()
    p.alignment = 3 # Justify
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.3
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.keep_together = True
    
    # Bullet-less formatting
    run_num = p.add_run(f"Câu {num}. ")
    run_num.bold = True
    run_num.font.name = VietFonts.BODY
    run_num.font.size = VietFonts.SIZE_BODY
    run_num.font.color.rgb = VietColors.DARK
    
    add_formatted_runs(p, text)
    
    if image_path:
        p_img = doc.add_paragraph()
        p_img.alignment = 1 # Center
        p_img.paragraph_format.space_before = Pt(4)
        p_img.paragraph_format.space_after = Pt(4)
        p_img.paragraph_format.keep_with_next = True
        p_img.add_run().add_picture(image_path, width=Cm(10.0))
        
    return p

def _add_viet_mcq_options(doc, options):
    """
    Render MCQ options (A, B, C, D) using a borderless table layout.
    Chooses 4-column, 2-column, or 1-column layout dynamically.
    """
    max_len = max(len(opt) for opt in options) if options else 0
    option_labels = ['A', 'B', 'C', 'D']
    
    if max_len < 18:
        cols_count = 4
        col_widths = [Cm(4.25)] * 4
        rows_count = 1
        layout = [[0, 1, 2, 3]]
    elif max_len < 38:
        cols_count = 2
        col_widths = [Cm(8.5)] * 2
        rows_count = 2
        layout = [[0, 1], [2, 3]]
    else:
        cols_count = 1
        col_widths = [Cm(17.0)]
        rows_count = 4
        layout = [[0], [1], [2], [3]]
        
    table = doc.add_table(rows=rows_count, cols=cols_count)
    table.alignment = 1 # Center
    
    # Set custom table width
    tblPr = table._tbl.tblPr
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(int(17.0 * 567)))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)
    
    for r_idx in range(rows_count):
        row = table.rows[r_idx]
        trPr = row._tr.get_or_add_trPr()
        trPr.append(OxmlElement('w:cantSplit'))
        trPr.append(OxmlElement('w:keepNext'))
        
        row_indices = layout[r_idx]
        for c_idx, opt_idx in enumerate(row_indices):
            cell = row.cells[c_idx]
            cell.width = col_widths[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = 0 # Left
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.15
            
            run_lbl = p.add_run(f"{option_labels[opt_idx]}. ")
            run_lbl.bold = True
            run_lbl.font.name = VietFonts.BODY
            run_lbl.font.size = VietFonts.SIZE_INFO
            run_lbl.font.color.rgb = VietColors.DARK
            
            opt_text = options[opt_idx] if opt_idx < len(options) else ""
            add_formatted_runs(p, opt_text, font_size=VietFonts.SIZE_INFO)
            
            # Remove cell border
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for side in ['top', 'left', 'bottom', 'right']:
                b = OxmlElement(f'w:{side}')
                b.set(qn('w:val'), 'none')
                tcBorders.append(b)
            tcPr.append(tcBorders)
            
            # cell margins (padding)
            tcMar = OxmlElement('w:tcMar')
            for m in ['top', 'bottom', 'left', 'right']:
                n = OxmlElement(f'w:{m}')
                n.set(qn('w:w'), '60')
                n.set(qn('w:type'), 'dxa')
                tcMar.append(n)
            tcPr.append(tcMar)
            
    doc.add_paragraph() # spacer

def _add_viet_tf_table(doc, statements):
    """
    Render Part II True/False statement block in a clean grid box with checkboxes.
    Columns: Statement (12.0cm) | Đúng (2.5cm) | Sai (2.5cm)
    """
    headers = ["Mệnh đề phát biểu nhận định", "Đúng", "Sai"]
    table = doc.add_table(rows=1, cols=3)
    table.alignment = 1 # Center
    
    # 17cm width
    set_table_width_dxa_custom(table, 17.0)
    col_widths = [Cm(12.0), Cm(2.5), Cm(2.5)]
    
    # Header
    hdr_cells = table.rows[0].cells
    for i, text in enumerate(headers):
        cell = hdr_cells[i]
        cell.width = col_widths[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = 1 if i > 0 else 0
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        
        run = p.add_run(text)
        run.bold = True
        run.font.name = VietFonts.BODY
        run.font.size = VietFonts.SIZE_INFO
        run.font.color.rgb = VietColors.DARK
        
        set_cell_shading_custom(cell, VietColors.LIGHT_GRAY_HEX)
        set_cell_margins_custom(cell)
        set_cell_borders_custom(cell)
        
    labels = ['a', 'b', 'c', 'd']
    for idx, label in enumerate(labels):
        stmt_text = statements.get(label, "")
        row = table.add_row()
        trPr = row._tr.get_or_add_trPr()
        trPr.append(OxmlElement('w:cantSplit'))
        
        # Col 0: Statement text
        cell_stmt = row.cells[0]
        cell_stmt.width = col_widths[0]
        cell_stmt.text = ""
        p_stmt = cell_stmt.paragraphs[0]
        p_stmt.alignment = 3 # justify
        p_stmt.paragraph_format.space_before = Pt(3)
        p_stmt.paragraph_format.space_after = Pt(3)
        p_stmt.paragraph_format.line_spacing = 1.15
        
        run_lbl = p_stmt.add_run(f"{label}) ")
        run_lbl.bold = True
        run_lbl.font.name = VietFonts.BODY
        run_lbl.font.size = VietFonts.SIZE_INFO
        
        add_formatted_runs(p_stmt, stmt_text, font_size=VietFonts.SIZE_INFO)
        set_cell_margins_custom(cell_stmt)
        set_cell_borders_custom(cell_stmt)
        
        # Col 1: True checkbox
        cell_t = row.cells[1]
        cell_t.width = col_widths[1]
        cell_t.text = ""
        p_t = cell_t.paragraphs[0]
        p_t.alignment = 1 # center
        p_t.paragraph_format.space_before = Pt(3)
        p_t.paragraph_format.space_after = Pt(3)
        r_t = p_t.add_run("☐ Đúng")
        r_t.font.name = VietFonts.BODY
        r_t.font.size = VietFonts.SIZE_INFO
        r_t.font.color.rgb = VietColors.DARK
        set_cell_margins_custom(cell_t)
        set_cell_borders_custom(cell_t)
        
        # Col 2: False checkbox
        cell_f = row.cells[2]
        cell_f.width = col_widths[2]
        cell_f.text = ""
        p_f = cell_f.paragraphs[0]
        p_f.alignment = 1 # center
        p_f.paragraph_format.space_before = Pt(3)
        p_f.paragraph_format.space_after = Pt(3)
        r_f = p_f.add_run("☐ Sai")
        r_f.font.name = VietFonts.BODY
        r_f.font.size = VietFonts.SIZE_INFO
        r_f.font.color.rgb = VietColors.DARK
        set_cell_margins_custom(cell_f)
        set_cell_borders_custom(cell_f)
        
    doc.add_paragraph() # spacer

def _add_viet_short_answer_box(doc, num, text, image_path=None):
    """Add short answer question with answer box."""
    _add_viet_question_stem(doc, num, text, image_path=image_path)
    
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.left_indent = Cm(1.0)
    
    run_ans = p.add_run("Đáp số của học sinh: ..............................................................")
    run_ans.font.name = VietFonts.BODY
    run_ans.font.size = VietFonts.SIZE_INFO
    run_ans.font.color.rgb = VietColors.GRAY
    
    doc.add_paragraph() # spacing

def set_table_width_dxa_custom(table, width_cm):
    width_dxa = int(width_cm * 567)
    tblPr = table._tbl.tblPr
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(width_dxa))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)

def set_cell_shading_custom(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls_custom("w")} w:fill="{color_hex}"/>')
    tcPr.append(shd)

def nsdecls_custom(prefix):
    return 'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'

def set_cell_margins_custom(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m in ['top', 'bottom', 'left', 'right']:
        n = OxmlElement(f'w:{m}')
        n.set(qn('w:w'), '120')
        n.set(qn('w:type'), 'dxa')
        tcMar.append(n)
    tcPr.append(tcMar)

def set_cell_borders_custom(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), VietColors.BORDER_GRAY_HEX)
        tcBorders.append(b)
    tcPr.append(tcBorders)

def export_homework(md_path, output_path):
    doc = Document()
    setup_viet_margins(doc, doc_type="homework")
    set_viet_default_style(doc)
    add_viet_page_number_footer(doc)
    
    # 1. Examination header
    add_viet_exam_header(doc)
    
    if not os.path.exists(md_path):
        doc.add_paragraph(f"Lỗi: Không tìm thấy file Markdown tại {md_path}")
        doc.save(output_path)
        return
        
    with open(md_path, "r", encoding="utf-8") as f:
        raw_lines = f.readlines()
        
    lines = [line.replace('\ufeff', '') for line in raw_lines]
    
    # Parse Markdown blocks
    current_part = None
    mcq_count = 0
    short_count = 0
    
    buffered_stem = None
    buffered_options = []
    buffered_image = None
    
    tf_stem = None
    tf_statements = {}
    tf_image = None
    
    def flush_mcq():
        nonlocal buffered_stem, buffered_options, buffered_image, mcq_count
        if buffered_stem and len(buffered_options) >= 4:
            mcq_count += 1
            _add_viet_question_stem(doc, mcq_count, buffered_stem, image_path=buffered_image)
            _add_viet_mcq_options(doc, buffered_options)
        buffered_stem = None
        buffered_options = []
        buffered_image = None
        
    def flush_tf():
        nonlocal tf_stem, tf_statements, tf_image, mcq_count
        if tf_stem and tf_statements:
            mcq_count += 1
            _add_viet_question_stem(doc, mcq_count, tf_stem, image_path=tf_image)
            _add_viet_tf_table(doc, tf_statements)
        tf_stem = None
        tf_statements = {}
        tf_image = None

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        
        if not stripped:
            i += 1
            continue
            
        # Part classification
        if "PHẦN I:" in stripped or "PHẦN I." in stripped:
            flush_mcq()
            flush_tf()
            current_part = 'I'
            add_viet_heading(doc, stripped.lstrip('#').strip(), level=2)
            i += 1
            continue
        elif "PHẦN II:" in stripped or "PHẦN II." in stripped:
            flush_mcq()
            flush_tf()
            current_part = 'II'
            add_viet_heading(doc, stripped.lstrip('#').strip(), level=2)
            i += 1
            continue
        elif "PHẦN III:" in stripped or "PHẦN III." in stripped:
            flush_mcq()
            flush_tf()
            current_part = 'III'
            add_viet_heading(doc, stripped.lstrip('#').strip(), level=2)
            i += 1
            continue
        elif "HƯỚNG DẪN GIẢI CHI TIẾT" in stripped or "ĐÁP ÁN CHO GIÁO VIÊN" in stripped:
            flush_mcq()
            flush_tf()
            break # Stop main parsing and delegate to answer key section
            
        # Blockquote Check
        is_blockquote = False
        if stripped.startswith('> '):
            stripped = stripped[2:]
            is_blockquote = True
        elif stripped.startswith('>'):
            stripped = stripped[1:]
            is_blockquote = True
            
        # Image Check
        img_match = re.search(r'([\(\s]*)(ready[^\(\s]+\.(?:png|jpg|jpeg))([\s\)]*)', stripped)
        embedded_img_path = None
        if img_match:
            entire_match = img_match.group(0)
            img_path = img_match.group(2)
            resolved = img_path
            if not os.path.exists(resolved):
                lesson_root = os.path.dirname(os.path.dirname(md_path))
                resolved = os.path.join(lesson_root, img_path)
            if not os.path.exists(resolved):
                lesson_root = os.path.dirname(os.path.dirname(md_path))
                resolved = os.path.join(lesson_root, "ready", "hinh_anh", os.path.basename(img_path))
            if os.path.exists(resolved):
                embedded_img_path = resolved
                stripped = stripped.replace(entire_match, " ")
                stripped = re.sub(r'\s+', ' ', stripped).strip()

        # Regex Match Question Stem
        # Matches "Câu 1. ..." or "**Câu 1:** ..."
        mcq_qmatch = re.match(r'^(?:\*?\*?)?(?:Câu\s+)?(\d+)(?:[\.\)]|(?:\s*\([A-Z]+\))?:?\*?\*?)\s+(.*)', stripped)
        
        # --- MCQ Parsing ---
        if current_part == 'I':
            if mcq_qmatch:
                flush_mcq()
                # Parse next lines to check if they are options A-D
                opts_found = []
                j = i + 1
                while j < min(i + 5, len(lines)):
                    opt_line = lines[j].strip()
                    optm = re.match(r'^(?:\*?\*?)?([A-Da-d])[\.\)]\s*(?:\*?\*?)?\s*(.*)', opt_line)
                    if optm:
                        opts_found.append(optm.group(2).strip())
                    else:
                        break
                    j += 1
                    
                if len(opts_found) >= 3:
                    buffered_stem = mcq_qmatch.group(2).strip()
                    buffered_options = opts_found[:4]
                    buffered_image = embedded_img_path
                    i = j
                    continue
            
            # Standard body text inside Part I
            add_viet_body(doc, stripped, is_blockquote=is_blockquote, size=VietFonts.SIZE_INFO)
            if embedded_img_path:
                p_img = doc.add_paragraph()
                p_img.alignment = 1
                p_img.add_run().add_picture(embedded_img_path, width=Cm(10.0))
                
        # --- True/False Parsing ---
        elif current_part == 'II':
            if mcq_qmatch:
                flush_tf()
                tf_stem = mcq_qmatch.group(2).strip()
                tf_image = embedded_img_path
                
                # Parse 4 sub-statements a, b, c, d
                j = i + 1
                while j < len(lines):
                    sub_line = lines[j].strip()
                    subm = re.match(r'^(?:\*?\*?)?([a-d])[\)\.]\s*(?:\*?\*?)?\s*(.*)', sub_line)
                    if subm:
                        tf_statements[subm.group(1)] = subm.group(2).strip()
                    else:
                        break
                    j += 1
                if tf_statements:
                    i = j
                    continue
            
        # --- Short Answer Parsing ---
        elif current_part == 'III':
            if mcq_qmatch:
                short_count += 1
                _add_viet_short_answer_box(doc, short_count, mcq_qmatch.group(2).strip(), image_path=embedded_img_path)
                i += 1
                continue
                
        i += 1
        
    # Flush remaining
    flush_mcq()
    flush_tf()
    
    # 5. Output Teacher Answer Key (Trang riêng)
    doc.add_page_break()
    add_viet_heading(doc, "ĐÁP ÁN VÀ HƯỚNG DẪN GIẢI CHI TIẾT", level=2)
    
    # Scan for ## Hướng dẫn giải chi tiết in markdown
    ans_text_lines = []
    ans_mode = False
    for line in lines:
        sl = line.strip()
        if re.match(r'^#+\s*[Đđ]áp\s*[áa]n', sl, re.IGNORECASE) or (sl.startswith('#') and 'ĐÁP ÁN' in sl.upper()) or "HƯỚNG DẪN GIẢI CHI TIẾT" in sl.upper():
            ans_mode = True
            continue
        if ans_mode:
            ans_text_lines.append(line)
            
    if ans_text_lines:
        ans_lines = [l.strip() for l in ans_text_lines if l.strip()]
        
        table_headers = None
        table_rows = []
        
        def flush_answer_table():
            nonlocal table_headers, table_rows
            if table_headers and table_rows:
                # render answer key table
                make_viet_table(doc, table_headers, table_rows, col_widths=None, total_width_cm=17.0)
            table_headers = None
            table_rows = []

        k = 0
        while k < len(ans_lines):
            l = ans_lines[k]
            
            # Check Markdown table inside answer sheet
            if l.startswith('|'):
                # Read header
                h_line = l
                k += 1
                sep = ans_lines[k] if k < len(ans_lines) else ""
                k += 1
                
                rows_buf = []
                while k < len(ans_lines) and ans_lines[k].startswith('|'):
                    rows_buf.append(ans_lines[k])
                    k += 1
                    
                def parse_cell(row_str):
                    return [c.strip() for c in row_str.split('|')[1:-1]]
                    
                table_headers = parse_cell(h_line)
                table_rows = [parse_cell(r) for r in rows_buf]
                flush_answer_table()
                continue
                
            # Regular text rendering
            if l.startswith('## '):
                add_viet_heading(doc, l[3:].strip(), level=1)
            elif l.startswith('### '):
                add_viet_heading(doc, l[4:].strip(), level=2)
            elif l.startswith('#### '):
                add_viet_heading(doc, l[5:].strip(), level=3)
            elif l.startswith('- ') or l.startswith('* '):
                add_viet_bullet(doc, l, level=1)
            elif l.startswith('+ '):
                add_viet_bullet(doc, l, level=2)
            else:
                add_viet_body(doc, l)
                
            k += 1
            
    doc.save(output_path)
