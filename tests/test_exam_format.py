"""Kiểm tra các nâng cấp định dạng Bài tập về nhà: bám PHONG CÁCH CÂU HỎI của
đề thi tốt nghiệp THPT 2025/2026 thật (nhóm câu dùng chung dữ kiện, khối hằng
số) nhưng KHÔNG đóng giả HÌNH THỨC đề thi (không Sở GD&ĐT/Mã đề/Số báo danh —
header đơn giản như PHT/KHBD); và KHBD theo đúng Phụ lục IV/CV5512."""
from docx import Document

from gems.config.loader import load_curriculum, load_identity
from gems.docx_export.homework_exporter import export_homework
from gems.docx_export.khbd_exporter import export_khbd
from gems.generation import stages
from gems.offline import fixtures

IDENTITY = load_identity()


def _all_text(doc) -> str:
    """`doc.paragraphs` chỉ trả về đoạn văn cấp thân tài liệu, KHÔNG bao gồm
    nội dung bên trong bảng (kể cả bảng lồng trong ô, như hộp "Mã đề thi") —
    hàm này duyệt đệ quy để lấy toàn bộ văn bản thật sự xuất hiện trong file."""
    parts = [p.text for p in doc.paragraphs]

    def walk(table):
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
                for nested in cell.tables:
                    walk(nested)

    for table in doc.tables:
        walk(table)
    return "\n".join(parts)


def _build_homework_with_extras():
    curriculum = load_curriculum()
    lesson = curriculum.get("bai4")
    architect = fixtures.build_architect(lesson)
    homework = fixtures.build_homework(architect)
    homework.shared_constants = "Nhiệt dung riêng của nước c = 4200 J/(kg.K)."
    homework.part1_questions[0].shared_context = "Một khối nước đá có khối lượng 200 g ở 0°C."
    homework.part1_questions[1].shared_context = homework.part1_questions[0].shared_context
    return lesson, homework


def test_shared_context_group_renders_once(tmp_path):
    lesson, homework = _build_homework_with_extras()
    md_dir = tmp_path / "md"
    md_dir.mkdir()
    hw_md = stages.write_homework_markdown(homework, lesson, md_dir)

    md_text = hw_md.read_text(encoding="utf-8")
    assert "Nội dung câu 1 và 2:" in md_text
    assert md_text.count(homework.part1_questions[0].shared_context) == 1

    out_path = tmp_path / "homework.docx"
    export_homework(str(hw_md), str(out_path), IDENTITY, lesson.name)
    doc = Document(str(out_path))
    all_text = _all_text(doc)
    assert all_text.count(homework.part1_questions[0].shared_context) == 1


def test_shared_constants_rendered_as_cho_biet(tmp_path):
    lesson, homework = _build_homework_with_extras()
    md_dir = tmp_path / "md"
    md_dir.mkdir()
    hw_md = stages.write_homework_markdown(homework, lesson, md_dir)
    out_path = tmp_path / "homework.docx"
    export_homework(str(hw_md), str(out_path), IDENTITY, lesson.name)
    doc = Document(str(out_path))
    all_text = _all_text(doc)
    assert "Cho biết:" in all_text
    assert homework.shared_constants in all_text


def test_homework_header_is_simple_not_fake_exam(tmp_path):
    """Header đơn giản như PHT/KHBD (tiêu đề bài + Họ tên/Lớp) — không còn
    đóng giả hình thức đề thi thật (Sở GD&ĐT, Mã đề thi, Số báo danh)."""
    lesson, homework = _build_homework_with_extras()
    md_dir = tmp_path / "md"
    md_dir.mkdir()
    hw_md = stages.write_homework_markdown(homework, lesson, md_dir)
    out_path = tmp_path / "homework.docx"
    export_homework(str(hw_md), str(out_path), IDENTITY, lesson.name)
    doc = Document(str(out_path))
    all_text = _all_text(doc)
    assert "BÀI TẬP VỀ NHÀ" in all_text
    assert "Họ và tên học sinh" in all_text
    assert "Mã đề thi" not in all_text
    assert "SỞ GIÁO DỤC" not in all_text
    assert "Số báo danh" not in all_text
    assert "ĐỀ THI TỐT NGHIỆP THPT NĂM" not in all_text


def test_exam_body_matches_reference_without_checkbox_table_or_answer_box(tmp_path):
    lesson, homework = _build_homework_with_extras()
    md_dir = tmp_path / "md"
    md_dir.mkdir()
    hw_md = stages.write_homework_markdown(homework, lesson, md_dir)
    out_path = tmp_path / "homework.docx"
    export_homework(str(hw_md), str(out_path), IDENTITY, lesson.name)
    doc = Document(str(out_path))
    all_text = _all_text(doc)

    assert "PHẦN I. Thí sinh trả lời từ câu 1 đến câu 18" in all_text
    assert "PHẦN II. Thí sinh trả lời từ câu 1 đến câu 4" in all_text
    assert "PHẦN III. Thí sinh trả lời từ câu 1 đến câu 6" in all_text
    assert "☐ Đúng" not in all_text and "☐ Sai" not in all_text
    assert "Đáp số của học sinh" not in all_text
    assert not any("Mệnh đề phát biểu nhận định" in cell.text for table in doc.tables for row in table.rows for cell in row.cells)


def test_end_notice_het_present(tmp_path):
    lesson, homework = _build_homework_with_extras()
    md_dir = tmp_path / "md"
    md_dir.mkdir()
    hw_md = stages.write_homework_markdown(homework, lesson, md_dir)
    out_path = tmp_path / "homework.docx"
    export_homework(str(hw_md), str(out_path), IDENTITY, lesson.name)
    doc = Document(str(out_path))
    all_text = _all_text(doc)
    assert "HẾT" in all_text


def test_khbd_admin_header_has_teacher_name_and_process_table(tmp_path):
    curriculum = load_curriculum()
    lesson = curriculum.get("bai4")
    architect = fixtures.build_architect(lesson)
    plan = fixtures.build_lesson_plan(architect)
    md_dir = tmp_path / "md"
    md_dir.mkdir()
    khbd_md = stages.write_lesson_plan_markdown(plan, lesson, md_dir, IDENTITY)

    out_path = tmp_path / "khbd.docx"
    export_khbd(str(khbd_md), str(out_path), IDENTITY)
    doc = Document(str(out_path))
    all_text = _all_text(doc)
    assert IDENTITY.teacher_name in all_text
    assert "c) Sản phẩm:" in all_text
    assert "Kĩ thuật dạy học tích cực:" in all_text
    assert "Năng lực số (Công văn 3456/BGDĐT-GDPT)" in all_text
    assert "1.2" in all_text

    process_tables = [t for t in doc.tables if len(t.columns) == 2
                       and "Sản phẩm" in [c.text.strip() for c in t.rows[0].cells]]
    assert process_tables, "phải có ít nhất 1 bảng tiến trình 'Hoạt động của giáo viên và học sinh | Sản phẩm'"
    header_texts = [c.text.strip() for c in process_tables[0].rows[0].cells]
    assert any("giáo viên và học sinh" in h.lower() for h in header_texts)
