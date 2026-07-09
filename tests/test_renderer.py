from docx import Document

from gems.docx_export.layout import get_layout
from gems.docx_export.markdown_ir import parse_markdown_blocks
from gems.docx_export.renderer import render_blocks


def test_render_blocks_smoke():
    doc = Document()
    layout = get_layout("pht")
    lines = [
        "## Đơn vị kiến thức 1",
        "Nội dung mở đầu **quan trọng** cho $Q = mc\\Delta t$",
        "- Gạch đầu dòng 1",
        "[DOT_LINE_90]",
        "> Trích dẫn mở rộng",
        "| A | B |",
        "| --- | --- |",
        "| 1 | 2 |",
    ]
    blocks = parse_markdown_blocks(lines)
    render_blocks(doc, blocks, layout)
    paragraphs_text = [p.text for p in doc.paragraphs if p.text]
    assert any("ĐƠN VỊ KIẾN THỨC 1" in t.upper() for t in paragraphs_text)
    tables = doc.tables
    assert len(tables) == 1
    assert len(tables[0].columns) == 2
