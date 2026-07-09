from pathlib import Path

from docx import Document

from gems.docx_export.pht_exporter import export_pht

FIXTURE = Path(__file__).parent / "fixtures" / "pht_bai3.md"


def test_export_pht_produces_valid_docx(tmp_path):
    out_path = tmp_path / "pht.docx"
    export_pht(str(FIXTURE), str(out_path), "Bài 3 - Nội năng, nhiệt lượng")
    assert out_path.exists()

    doc = Document(str(out_path))
    all_text = "\n".join(p.text for p in doc.paragraphs)
    assert "PHIẾU HỌC TẬP" in all_text
    assert "TỰ ĐÁNH GIÁ" in all_text.upper()


def test_pht_margins_are_2cm_all_sides(tmp_path):
    out_path = tmp_path / "pht.docx"
    export_pht(str(FIXTURE), str(out_path), "Bài 3")
    doc = Document(str(out_path))
    section = doc.sections[0]
    for margin in (section.top_margin, section.bottom_margin, section.left_margin, section.right_margin):
        assert round(margin.cm, 2) == 2.0


def test_pht_has_student_info_and_self_eval_tables(tmp_path):
    out_path = tmp_path / "pht.docx"
    export_pht(str(FIXTURE), str(out_path), "Bài 3")
    doc = Document(str(out_path))
    assert len(doc.tables) >= 2  # bảng thông tin học sinh + bảng tự đánh giá


def test_missing_file_produces_placeholder(tmp_path):
    out_path = tmp_path / "missing.docx"
    export_pht(str(tmp_path / "khong_co.md"), str(out_path), "Bài X")
    doc = Document(str(out_path))
    assert any("Không tìm thấy" in p.text for p in doc.paragraphs)
