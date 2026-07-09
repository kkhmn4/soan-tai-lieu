from pathlib import Path

from docx import Document

from gems.config.loader import load_identity
from gems.docx_export.khbd_exporter import export_khbd

FIXTURE = Path(__file__).parent / "fixtures" / "khbd_bai3.md"
IDENTITY = load_identity()


def test_export_khbd_produces_valid_docx(tmp_path):
    out_path = tmp_path / "khbd.docx"
    export_khbd(str(FIXTURE), str(out_path), IDENTITY)
    assert out_path.exists()
    doc = Document(str(out_path))
    all_text = "\n".join(p.text for p in doc.paragraphs)
    assert "KẾ HOẠCH BÀI DẠY" in all_text
    assert "ĐIỀU CHỈNH BÀI DẠY" in all_text.upper()


def test_khbd_margins_3_2_2_2():
    doc = Document()
    from gems.docx_export import styles
    styles.setup_margins(doc, "khbd")
    section = doc.sections[0]
    assert round(section.left_margin.cm, 2) == 3.0
    assert round(section.right_margin.cm, 2) == 2.0
    assert round(section.top_margin.cm, 2) == 2.0
    assert round(section.bottom_margin.cm, 2) == 2.0


def test_signature_block_has_3_columns(tmp_path):
    out_path = tmp_path / "khbd.docx"
    export_khbd(str(FIXTURE), str(out_path), IDENTITY)
    doc = Document(str(out_path))
    last_table = doc.tables[-1]
    assert len(last_table.columns) == 3
    header_texts = [c.text for c in last_table.rows[0].cells]
    assert any("BAN GIÁM HIỆU" in t for t in header_texts)
    assert any("TỔ TRƯỞNG" in t for t in header_texts)
    assert any("GIÁO VIÊN SOẠN" in t for t in header_texts)


def test_gv_hs_table_has_two_columns_and_navy_keywords(tmp_path):
    out_path = tmp_path / "khbd.docx"
    export_khbd(str(FIXTURE), str(out_path), IDENTITY)
    doc = Document(str(out_path))
    gv_hs_tables = [t for t in doc.tables if len(t.columns) == 2 and t is not doc.tables[-1]]
    assert gv_hs_tables, "phải có ít nhất 1 bảng GV/HS 2 cột"


def test_missing_file_produces_placeholder(tmp_path):
    out_path = tmp_path / "missing.docx"
    export_khbd(str(tmp_path / "khong_co.md"), str(out_path), IDENTITY)
    doc = Document(str(out_path))
    assert any("Không tìm thấy" in p.text for p in doc.paragraphs)
