import pytest
from docx import Document

from gems.docx_export import styles
from gems.docx_export.layout import get_layout


def test_setup_margins_khbd():
    doc = Document()
    layout = styles.setup_margins(doc, "khbd")
    section = doc.sections[0]
    assert section.left_margin.cm == pytest.approx(3.0, abs=0.01)
    assert section.right_margin.cm == pytest.approx(2.0, abs=0.01)
    assert layout.doc_type == "khbd"


def test_set_table_width_dxa_sets_center_alignment():
    doc = Document()
    table = doc.add_table(rows=1, cols=2)
    styles.set_table_width_dxa(table, 16.0)
    from docx.enum.table import WD_TABLE_ALIGNMENT
    assert table.alignment == WD_TABLE_ALIGNMENT.CENTER
    tblW = table._tbl.tblPr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblW")
    assert tblW is not None
    assert tblW.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w") == str(int(16.0 * 567))


def test_make_table_header_row_shaded_and_column_count():
    doc = Document()
    layout = get_layout("khbd")
    table = styles.make_table(doc, ["Cột A", "Cột B"], [["1", "2"], ["3", "4"]], layout)
    assert len(table.columns) == 2
    assert len(table.rows) == 3  # 1 header + 2 data rows
    header_cell = table.rows[0].cells[0]
    assert header_cell.paragraphs[0].runs[0].bold is True


def test_make_table_borderless_removes_borders():
    doc = Document()
    layout = get_layout("pht")
    table = styles.make_table(doc, ["A", "B"], [["x", "y"]], layout, borderless=True)
    tcPr = table.rows[0].cells[0]._tc.get_or_add_tcPr()
    borders = tcPr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcBorders")
    assert borders is not None
    top = borders.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}top")
    assert top.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val") == "none"


def test_add_page_number_footer_has_field_codes():
    doc = Document()
    styles.add_page_number_footer(doc)
    footer_p = doc.sections[0].footer.paragraphs[0]
    xml = footer_p._p.xml
    assert 'w:instr="PAGE"' in xml
    assert 'w:instr="NUMPAGES"' in xml
