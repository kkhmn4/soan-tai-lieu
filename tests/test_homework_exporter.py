from pathlib import Path

from docx import Document

from gems.config.loader import load_identity
from gems.docx_export.homework_exporter import export_homework

FIXTURE = Path(__file__).parent / "fixtures" / "homework_bai3.md"
IDENTITY = load_identity()


def test_export_homework_produces_valid_docx(tmp_path):
    out_path = tmp_path / "homework.docx"
    export_homework(str(FIXTURE), str(out_path), IDENTITY, "Bài 3 - Nội năng, nhiệt lượng")
    assert out_path.exists()

    doc = Document(str(out_path))
    all_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Câu 1." in all_text
    assert "PHẦN I" in all_text.upper()
    assert "ĐÁP ÁN" in all_text.upper()


def test_leaked_answer_is_stripped_from_student_section(tmp_path):
    out_path = tmp_path / "homework.docx"
    export_homework(str(FIXTURE), str(out_path), IDENTITY, "Bài 3")
    doc = Document(str(out_path))

    # Nội dung trước trang ngắt (student section) không được lộ "Đ/S: Đ" hay "(Đ/S: S)"
    student_paragraphs = []
    for p in doc.paragraphs:
        student_paragraphs.append(p.text)
        if "ĐÁP ÁN" in p.text.upper():
            break
    student_text = "\n".join(student_paragraphs)
    assert "(Đ/S: Đ)" not in student_text
    assert "(Đ/S: S)" not in student_text


def test_mcq_options_rendered_as_table(tmp_path):
    out_path = tmp_path / "homework.docx"
    export_homework(str(FIXTURE), str(out_path), IDENTITY, "Bài 3")
    doc = Document(str(out_path))
    assert len(doc.tables) >= 1


def test_missing_file_produces_placeholder_docx(tmp_path):
    out_path = tmp_path / "missing.docx"
    export_homework(str(tmp_path / "khong_ton_tai.md"), str(out_path), IDENTITY, "Bài X")
    assert out_path.exists()
    doc = Document(str(out_path))
    assert any("Không tìm thấy" in p.text for p in doc.paragraphs)
