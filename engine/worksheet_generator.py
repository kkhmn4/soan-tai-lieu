            p_theory = doc.add_paragraph()
            p_theory.paragraph_format.left_indent = Inches(0.2)
            cloze_text = unit.core_theory.summary_cloze.replace("[DOT_LINE_90]", "." * 90)
            run_theory = p_theory.add_run(cloze_text)
            run_theory.font.size = Pt(10)
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from templates.prompts import SYSTEM_WORKSHEET_GENERATION_PROMPT

load_dotenv()

# Định nghĩa các Schema Pydantic cho Structured Outputs của Phiếu học tập
class TaskContent(BaseModel):
    task_id: str = Field(description="Mã nhiệm vụ (ví dụ: NV1...)")
    task_type: str = Field(description="Loại nhiệm vụ (Matching Matrix, Visual Cloze Test...)")
    task_name: str = Field(description="Tên nhiệm vụ")
    content: str = Field(description="Nội dung chi tiết câu hỏi, bài tập hoặc hoạt động thực tế. Phải cụ thể và hướng dẫn chi tiết học sinh làm gì.")

class DiscoverySection(BaseModel):
    tasks: list[TaskContent] = Field(description="Danh sách các nhiệm vụ khám phá của đơn vị kiến thức")

class CoreTheorySection(BaseModel):
    summary_cloze: str = Field(description="Đoạn văn tóm tắt lý thuyết đục lỗ điền khuyết sử dụng các số thứ tự (1), (2)...")
    key_words: list[str] = Field(description="Danh sách đáp án từ khóa tương ứng với các số thứ tự đục lỗ")

class ApplicationSection(BaseModel):
    scenario: str = Field(description="Tình huống thực tế đời sống mới và thử thách tính toán/lập luận vật lý")

class ExtensionSection(BaseModel):
    reading_content: str = Field(description="Nội dung đọc hiểu mở rộng, kết nối STEM hoặc hướng nghiệp")

class UnitWorksheetContent(BaseModel):
    unit_id: str = Field(description="Mã đơn vị kiến thức")
    unit_name: str = Field(description="Tên đơn vị kiến thức")
    discovery: DiscoverySection
    core_theory: CoreTheorySection
    application: ApplicationSection
    extension: ExtensionSection

class LessonWorksheet(BaseModel):
    lesson_name: str
    units: list[UnitWorksheetContent]

class WorksheetGenerator:
    def __init__(self, api_key: str = None):
        # Ưu tiên API Key truyền trực tiếp, sau đó đến env, nếu không có sẽ tự động dùng credentials mặc định của Antigravity
        self.api_key = api_key or os.getenv("GEMINI_API_KEY")
        if self.api_key:
            self.client = genai.Client(api_key=self.api_key)
        else:
            self.client = genai.Client()
        self.model_name = "gemini-2.5-flash"

    def generate_content(self, matrix_data: dict) -> LessonWorksheet:
        """Gọi Gemini API để sinh nội dung chi tiết cho từng phần của phiếu học tập dựa vào ma trận."""
        print(f"[PROCESS] Dang sinh noi dung chi tiet cho phieu hoc tap: '{matrix_data.get('lesson_name')}'...")
        
        matrix_json = json.dumps(matrix_data, ensure_ascii=False)


        )
        
        response = self.client.models.generate_content(
            model=self.model_name,
            contents=prompt,
            config=types.GenerateContentConfig(
                response_mime_type="application/json",
                response_schema=LessonWorksheet,
                system_instruction=SYSTEM_WORKSHEET_GENERATION_PROMPT,
                temperature=0.3,
            )
        )
        
        data = json.loads(response.text)
        return LessonWorksheet(**data)

    def write_dots(self, doc, num_lines=3):
        """Thêm các dòng chấm chấm để học sinh ghi câu trả lời."""
        for _ in range(num_lines):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run("." * 105)
            run.font.name = "Arial"
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(180, 180, 180) # Màu xám nhạt

    def export_to_docx(self, worksheet: LessonWorksheet, output_path: str):
        """Xuất phiếu học tập sang file MS Word (.docx) định dạng chuyên nghiệp."""
        print(f"[PROCESS] Dang xuat file Word tai: {output_path}...")
        doc = Document()
        
        # Thiết lập lề trang in ấn (lề gọn gàng để tối ưu không gian viết bài)
        for section in doc.sections:
            section.top_margin = Inches(0.8)
            section.bottom_margin = Inches(0.8)
            section.left_margin = Inches(0.8)
            section.right_margin = Inches(0.8)

        # 1. Khung thông tin chung (Họ tên học sinh, lớp, ngày)
        header_table = doc.add_table(rows=2, cols=2)
        header_table.autofit = False
        
        # Dòng 1: Tên bài học và thông tin quản lý lớp học
        cell_title = header_table.cell(0, 0)
        p_title = cell_title.paragraphs[0]
        run_title = p_title.add_run(f"PHIẾU HỌC TẬP: {worksheet.lesson_name.upper()}")
        run_title.bold = True
        run_title.font.size = Pt(14)
        run_title.font.name = "Arial"
        
        cell_info = header_table.cell(0, 1)
        p_info = cell_info.paragraphs[0]
        run_info = p_info.add_run("Họ và tên: ....................................................\nLớp: .................. Ngày: ......./......./.......")
        run_info.font.size = Pt(10)
        run_info.font.name = "Arial"
        
        doc.add_paragraph().paragraph_format.space_after = Pt(12)

        # 2. Chi tiết từng Đơn vị kiến thức (Tuyến tính)
        for unit in worksheet.units:
            # Tiêu đề Đơn vị kiến thức
            p_unit = doc.add_paragraph()
            p_unit.paragraph_format.space_before = Pt(12)
            p_unit.paragraph_format.space_after = Pt(6)
            run_unit = p_unit.add_run(f"📍 {unit.unit_id}: {unit.unit_name.upper()}")
            run_u




















                run_task_name.font.color.rgb = RGBColor(0, 102, 204)
                
                p_desc = doc.add_paragraph()
                p_desc.paragraph_format.left_indent = Inches(0.4)
                run_desc = p_desc.add_run(task.content)
                run_desc.font.size = Pt(10)
                
                # Chừa 4 dòng trống cho học sinh viết câu trả lời khám phá
                self.write_dots(doc, num_lines=4)

            # --- PHẦN 2: KIẾN THỨC TRỌNG TÂM ---
            p_sec2 = doc.add_paragraph()
            run_sec2 = p_sec2.add_run("[Phần 2] KIẾN THỨC TRỌNG TÂM")
            run_sec2.bold = True
            run_sec2.font.size = Pt(11)
            run_sec2.font.color.rgb = RGBColor(0, 102, 51) # Màu xanh lá
            
            p_theory = doc.add_paragraph()
            p_theory.paragraph_format.left_indent = Inches(0.2)
            cloze_text = unit.core_theory.summary_cloze.replace("[DOT_LINE_90]", "." * 90)
            run_theory = p_theory.add_run(cloze_text)
            run_theory.font.size = Pt(10)
            
            # Khung danh sách từ khóa gợi ý bên dưới (để học sinh điền)
            p_keywords = doc.add_paragraph()
            p_keywords.paragraph_format.left_indent = Inches(0.2)
            run_kw_label = p_keywords.add_run("🔑 Gợi ý từ khóa cần điền: ")
            run_kw_l






































        # Lưu tài liệu Word
        doc.save(output_path)
        print(f"✅ Da xuat thanh cong: {output_path}")

    def save_markdown(self, worksheet: LessonWorksheet, output_path: str):

    def save_markdown(self, worksheet: LessonWorksheet, output_path: str):
        """Lưu bản nháp Markdown của phiếu học tập để xem trước."""
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(f"# PHIẾU HỌC TẬP: {worksheet.lesson_name}\n\n")
            f.write("Họ và tên học sinh: ....................................................\n")
            f.write("Lớp: .................. Ngày: ......./......./.......\n\n")
            f.write("---\n\n")
            
            for unit in worksheet.units:
                f.write(f"## 📍 {unit.unit_id}: {unit.unit_name}\n\n")
                
                f.write("### [Phần 1] NHIỆM VỤ KHÁM PHÁ\n")
                for task in unit.discovery.tasks:
                    f.write(f"#### 👉 {task.task_id} ({task.task_type}): {task.task_name}\n")
                    f.write(f"{task.content}\n\n")
                    f.write("*Trả lời:*\n\n\n\n")
                    
                f.write("### [Phần 2] KIẾN THỨC TRỌNG TÂM\n")
                f.write(f"{unit.core_theory.summary_cloze}\n\n")
                f.write(f"*Từ khóa gợi ý:* {', '.join(unit.core_theory.key_words)}\n\n")
                
                f.write("### [Phần 3] THỬ THÁCH VẬN DỤNG\n")
                f.write(f"{unit.application.scenario}\n\n")
                f.write("*Bài giải:*\n\n\n\n\n")
                
                f.write("### [Phần 4] MỞ RỘNG KIẾN THỨC\n")
                f.write(f"> {unit.extension.reading_content}\n\n")
                f.write("---\n\n")
        print(f"✅ Da luu file Markdown: {output_path}")