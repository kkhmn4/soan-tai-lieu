"""Bộ helper Word duy nhất cho lề trang, bảng biểu, ô, tiêu đề, thân bài...

Trước đây mỗi exporter (PHT/KHBD/Bài tập) tự viết lại một bản riêng của
set_table_width_dxa/set_cell_margins/set_cell_shading/set_cell_borders (ít
nhất 4 bản độ-rộng-bảng và 5 kiểu đệm ô khác nhau trong cùng 1 dự án). Từ
giờ CHỈ có các hàm trong module này — không exporter nào được tự viết XML
bảng/ô riêng nữa.
"""
from __future__ import annotations

from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Cm, Pt

from gems.docx_export.layout import DocLayout, PAGE_HEIGHT_CM, PAGE_WIDTH_CM, get_layout
from gems.docx_export.latex_clean import clean_latex, smart_typography
from gems.docx_export.palette import VietColors, VietFonts
from gems.docx_export.run_formatter import add_formatted_runs

CM_TO_DXA = 567  # 1cm ~= 567 twips (dxa)


# ============================================================
#  PAGE SETUP
# ============================================================
def setup_margins(doc, doc_type: str) -> DocLayout:
    layout = get_layout(doc_type)
    for section in doc.sections:
        section.page_width = Cm(PAGE_WIDTH_CM)
        section.page_height = Cm(PAGE_HEIGHT_CM)
        section.top_margin = Cm(layout.top_margin_cm)
        section.bottom_margin = Cm(layout.bottom_margin_cm)
        section.left_margin = Cm(layout.left_margin_cm)
        section.right_margin = Cm(layout.right_margin_cm)
    return layout


def set_default_style(doc) -> None:
    style = doc.styles["Normal"]
    style.font.name = VietFonts.BODY
    style.font.size = VietFonts.SIZE_BODY
    style.font.color.rgb = VietColors.DARK
    pf = style.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.3
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


# ============================================================
#  CELL / TABLE HELPERS — CHỈ 1 bản duy nhất trong toàn bộ codebase
# ============================================================
def set_cell_shading(cell, color_hex: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    tcPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'))


def set_cell_margins(cell, top_pt=5.7, bottom_pt=5.7, left_pt=8.5, right_pt=8.5) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, pt_val in [("top", top_pt), ("bottom", bottom_pt), ("left", left_pt), ("right", right_pt)]:
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(int(pt_val * 20)))
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)


def set_cell_borders(cell, color_hex: str = VietColors.BORDER_GRAY_HEX) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ["top", "left", "bottom", "right"]:
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color_hex)
        tcBorders.append(border)
    tcPr.append(tcBorders)


def remove_cell_borders(cell) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ["top", "left", "bottom", "right"]:
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "none")
        tcBorders.append(border)
    tcPr.append(tcBorders)


def set_table_width_dxa(table, width_cm: float) -> None:
    """Đặt độ rộng bảng bằng đơn vị dxa (KHÔNG BAO GIỜ dùng autofit/percentage).

    python-docx tự tạo sẵn 1 phần tử `w:tblW` mặc định (w:w="0" w:type="auto")
    trong `tblPr` của mọi bảng mới — nếu chỉ append thêm phần tử mới mà
    không xoá cái cũ, XML sẽ có 2 `w:tblW` xung đột nhau (bug này từng có
    trong bản gốc). Ở đây xoá phần tử cũ trước khi thêm phần tử đúng.
    """
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    width_dxa = int(width_cm * CM_TO_DXA)
    tblPr = table._tbl.tblPr
    for existing in tblPr.findall(qn("w:tblW")):
        tblPr.remove(existing)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(width_dxa))
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)


def mark_row_no_split(row, repeat_header: bool = False) -> None:
    trPr = row._tr.get_or_add_trPr()
    trPr.append(OxmlElement("w:cantSplit"))
    if repeat_header:
        trPr.append(OxmlElement("w:tblHeader"))


def make_table(doc, headers: list[str], rows: list[list[str]], layout: DocLayout,
               col_widths_cm: list[float] | None = None, borderless: bool = False):
    """Dựng bảng chuẩn GEMS: header xám #D9D9D9 in đậm căn giữa, viền xám
    mảnh (trừ khi borderless=True), không vỡ dòng giữa trang, lặp header
    khi sang trang, độ rộng luôn tính từ `layout.content_width_cm`.
    """
    total_width_cm = layout.content_width_cm
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = False
    set_table_width_dxa(table, total_width_cm)

    if col_widths_cm:
        widths_cm = [col_widths_cm[i] if i < len(col_widths_cm) else total_width_cm / len(headers)
                     for i in range(len(headers))]
    else:
        widths_cm = [total_width_cm / len(headers)] * len(headers)

    hdr_cells = table.rows[0].cells
    for i, text in enumerate(headers):
        cell = hdr_cells[i]
        cell.width = Cm(widths_cm[i])
        # Không gọi cell.text = "" — 1 ô bảng mới đã có sẵn 1 paragraph
        # rỗng KHÔNG run nào; gán cell.text = "" trước khi add_run sẽ tạo
        # dư 1 run rỗng đứng trước run thật (bug từng có ở bản gốc, khiến
        # paragraphs[0].runs[0] không phải là run chứa nội dung thật).
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(smart_typography(str(text)))
        run.bold = True
        run.font.name = VietFonts.BODY
        run.font.size = VietFonts.SIZE_TABLE
        run.font.color.rgb = VietColors.DARK
        if not borderless:
            set_cell_shading(cell, VietColors.LIGHT_GRAY_HEX)
        set_cell_margins(cell)
        if borderless:
            remove_cell_borders(cell)
        else:
            set_cell_borders(cell)
    mark_row_no_split(table.rows[0], repeat_header=True)

    for row_data in rows:
        row = table.add_row()
        mark_row_no_split(row)
        padded = list(row_data) + [""] * (len(headers) - len(row_data))
        for c_idx, val in enumerate(padded):
            cell = row.cells[c_idx]
            cell.width = Cm(widths_cm[c_idx])
            lines = str(val).replace("<br>", "\n").replace("<br/>", "\n").split("\n")
            first = True
            for line in lines:
                line_text = line.strip()
                if not line_text:
                    continue
                p = cell.paragraphs[0] if first else cell.add_paragraph()
                first = False
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.15

                if line_text.startswith("- ") or line_text.startswith("* "):
                    line_text = line_text[2:]
                    p.paragraph_format.left_indent = Cm(0.3)
                    p.paragraph_format.first_line_indent = Cm(-0.2)
                    run_bullet = p.add_run("- ")
                    run_bullet.font.name = VietFonts.BODY
                    run_bullet.font.size = VietFonts.SIZE_INFO
                    run_bullet.font.color.rgb = VietColors.DARK
                elif line_text.startswith("+ "):
                    line_text = line_text[2:]
                    p.paragraph_format.left_indent = Cm(0.5)
                    p.paragraph_format.first_line_indent = Cm(-0.2)
                    run_bullet = p.add_run("+ ")
                    run_bullet.font.name = VietFonts.BODY
                    run_bullet.font.size = VietFonts.SIZE_INFO
                    run_bullet.font.color.rgb = VietColors.DARK

                add_formatted_runs(p, line_text, font_size=VietFonts.SIZE_INFO)

            set_cell_margins(cell)
            if borderless:
                remove_cell_borders(cell)
            else:
                set_cell_borders(cell)

    doc.add_paragraph()
    return table


# ============================================================
#  PARAGRAPH CREATORS
# ============================================================
def add_title(doc, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(smart_typography(text).upper())
    run.bold = True
    run.font.name = VietFonts.HEADER
    run.font.size = VietFonts.SIZE_TITLE
    run.font.color.rgb = VietColors.PRIMARY
    return p


_HEADING_SIZES = {1: VietFonts.SIZE_H1, 2: VietFonts.SIZE_H2, 3: VietFonts.SIZE_H3}
_HEADING_SPACE_BEFORE = {1: Pt(12), 2: Pt(8), 3: Pt(6)}
_HEADING_SPACE_AFTER = {1: Pt(6), 2: Pt(4), 3: Pt(3)}


def add_heading(doc, text: str, level: int = 1):
    """Heading cấp 1 (14pt đậm, primary, dùng cho I/II/III), cấp 2 (13pt đậm,
    đen, dùng cho 1/2/3), cấp 3 (13pt nghiêng, đen, dùng cho a/b/c)."""
    p = doc.add_paragraph()
    p.style = doc.styles[f"Heading {level}"]
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = _HEADING_SPACE_BEFORE.get(level, Pt(6))
    p.paragraph_format.space_after = _HEADING_SPACE_AFTER.get(level, Pt(4))
    p.paragraph_format.left_indent = Cm(0)
    p.paragraph_format.first_line_indent = Cm(0)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    display_text = clean_latex(text)
    if level == 1:
        display_text = display_text.upper()

    run = p.add_run(display_text)
    run.bold = level in (1, 2)
    run.italic = level == 3
    run.font.name = VietFonts.HEADER
    run.font.size = _HEADING_SIZES.get(level, VietFonts.SIZE_BODY)
    run.font.color.rgb = VietColors.PRIMARY if level == 1 else VietColors.DARK
    return p


def add_body(doc, text: str, *, is_blockquote: bool = False, size=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.3
    if is_blockquote:
        p.paragraph_format.left_indent = Cm(1.0)
        p.paragraph_format.first_line_indent = Cm(0)
    else:
        p.paragraph_format.left_indent = Cm(0)
        p.paragraph_format.first_line_indent = Cm(1.0)
    add_formatted_runs(p, text, force_italic=is_blockquote, font_size=size)
    return p


def add_bullet(doc, text: str, *, is_blockquote: bool = False, level: int = 1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.3
    p.paragraph_format.left_indent = Cm(0.5 if level == 1 else 1.0)
    p.paragraph_format.first_line_indent = Cm(-0.3)

    bullet_symbol = "- " if level == 1 else "+ "
    run_bullet = p.add_run(bullet_symbol)
    run_bullet.font.name = VietFonts.BODY
    run_bullet.font.size = VietFonts.SIZE_BODY
    run_bullet.font.color.rgb = VietColors.DARK

    if text.startswith(("- ", "* ", "+ ")):
        text = text[2:]

    add_formatted_runs(p, text, force_italic=is_blockquote)
    return p


def add_dot_line(doc, count: int = 1):
    for _ in range(count):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Cm(1.0)
        run = p.add_run("." * 90)
        run.font.name = VietFonts.BODY
        run.font.size = VietFonts.SIZE_INFO
        run.font.color.rgb = VietColors.GRAY


def add_separator(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    run = p.add_run("—" * 35)
    run.font.name = VietFonts.BODY
    run.font.size = VietFonts.SIZE_TINY
    run.font.color.rgb = VietColors.BORDER_GRAY


def add_simple_field(paragraph, instr: str) -> None:
    """Chèn 1 field code Word đơn giản (vd. PAGE, NUMPAGES) — dùng chung cho
    footer số trang và dòng "(Đề có N trang)" ở đầu đề bài kiểu đề thi."""
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), instr)
    paragraph._p.append(fld)


def add_page_number_footer(doc) -> None:
    """Chèn footer 'Trang X/Y' bằng field code động PAGE/NUMPAGES của Word."""
    for section in doc.sections:
        p = section.footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        p.text = ""

        run1 = p.add_run("Trang ")
        run1.font.name = VietFonts.BODY
        run1.font.size = VietFonts.SIZE_TINY
        run1.font.color.rgb = VietColors.GRAY

        add_simple_field(p, "PAGE")

        run2 = p.add_run("/")
        run2.font.name = VietFonts.BODY
        run2.font.size = VietFonts.SIZE_TINY
        run2.font.color.rgb = VietColors.GRAY

        add_simple_field(p, "NUMPAGES")
