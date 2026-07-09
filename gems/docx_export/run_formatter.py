"""Tách **đậm**, *nghiêng*, $công thức$, ==highlight== trong 1 chuỗi text
thành các text run Word được định dạng đúng — dùng chung cho MỌI nơi ghi
văn bản (thân bài, tiêu đề, VÀ ô bảng). Bản cũ có tới 2 bản sao chép yếu hơn
của logic này bên trong các hàm dựng bảng, thiếu clean_latex trên đoạn văn
bản thường và không hỗ trợ ==highlight== — hợp nhất về đây để không còn lệch.
"""
from __future__ import annotations

import re

from docx.enum.text import WD_COLOR_INDEX

from gems.docx_export.latex_clean import clean_latex, smart_typography
from gems.docx_export.palette import VietColors, VietFonts

_SPLIT_RE = re.compile(r"(\*\*.*?\*\*|\$.*?\$|==.*?==)")
_ITALIC_RE = re.compile(r"(\*.*?\*)")


def _set_run_style(run, *, font_size=None, bold=False, italic=False, highlight=False):
    run.font.name = VietFonts.BODY
    run.font.size = font_size or VietFonts.SIZE_BODY
    run.font.color.rgb = VietColors.DARK
    run.bold = bold
    run.italic = italic
    if highlight:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def add_formatted_runs(paragraph, text: str, *, force_italic=False, font_size=None, bold_all=False) -> None:
    """Ghi `text` thành các run trong `paragraph`, nhận diện:
    - `**đậm**`
    - `*nghiêng*`
    - `$công thức latex$`  -> luôn nghiêng, chạy qua clean_latex
    - `==highlight==`      -> tô vàng (WD_COLOR_INDEX.YELLOW)
    """
    if not text:
        return

    for part in _SPLIT_RE.split(text):
        if not part:
            continue

        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(smart_typography(clean_latex(part[2:-2])))
            _set_run_style(run, font_size=font_size, bold=True, italic=force_italic)
        elif part.startswith("$") and part.endswith("$"):
            run = paragraph.add_run(clean_latex(part[1:-1]))
            _set_run_style(run, font_size=font_size, bold=bold_all, italic=True)
        elif part.startswith("==") and part.endswith("=="):
            run = paragraph.add_run(smart_typography(clean_latex(part[2:-2])))
            _set_run_style(run, font_size=font_size, bold=bold_all, italic=force_italic, highlight=True)
        else:
            for sub_part in _ITALIC_RE.split(part):
                if not sub_part:
                    continue
                if sub_part.startswith("*") and sub_part.endswith("*") and len(sub_part) > 1:
                    run = paragraph.add_run(smart_typography(clean_latex(sub_part[1:-1])))
                    _set_run_style(run, font_size=font_size, bold=bold_all, italic=True)
                else:
                    run = paragraph.add_run(smart_typography(clean_latex(sub_part)))
                    _set_run_style(run, font_size=font_size, bold=bold_all, italic=force_italic)
