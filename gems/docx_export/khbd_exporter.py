"""Xuất Kế hoạch bài dạy (KHBD, chuẩn Công văn 5512) từ Markdown sang DOCX."""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from gems.config.loader import Identity
from gems.docx_export import styles
from gems.docx_export.layout import get_layout
from gems.docx_export.markdown_ir import parse_markdown_blocks
from gems.docx_export.palette import VietColors, VietFonts
from gems.docx_export.renderer import embed_image

_CV5512_KEYWORDS = [
    "Bước 1: Chuyển giao nhiệm vụ:", "Bước 2: Thực hiện nhiệm vụ:",
    "Bước 3: Báo cáo, thảo luận:", "Bước 4: Kết luận, nhận định:",
]


def _recolor_cv5512_keywords(table) -> None:
    """Sau khi bảng đã dựng xong (bold theo `**...**` chuẩn), tô Navy cho
    các run đậm khớp 4 bước CV5512 — tách riêng bước này khỏi
    `styles.make_table` dùng chung để bảng generic không phải biết về danh
    sách từ khoá đặc thù của KHBD."""
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    if run.bold and any(kw in run.text for kw in _CV5512_KEYWORDS):
                        run.font.color.rgb = VietColors.PRIMARY


def _make_process_table(doc, headers: list[str], rows: list[list[str]], layout) -> None:
    """Bảng tiến trình dạy học khớp Phụ lục IV: cột 1 "Hoạt động của giáo
    viên và học sinh" (đã gộp 4 bước thành 1 văn bản liên tục ở stages.py,
    không còn tách theo cặp GV/HS như bản trước), cột 2 "Sản phẩm"."""
    table = styles.make_table(doc, headers, rows, layout, col_widths_cm=[11.0, 5.0])
    _recolor_cv5512_keywords(table)


def _add_admin_header(doc, lesson_title: str, duration_str: str, identity: Identity) -> None:
    """Khớp mẫu chính thức Phụ lục IV (Công văn 5512/BGDĐT-GDTrH): dòng
    "Trường:"/"Tổ:" ở góc trái, tiêu đề "KẾ HOẠCH BÀI DẠY", rồi "Họ và tên
    giáo viên:" — trường này hoàn toàn vắng mặt ở bản trước."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15

    run_school = p.add_run("Trường: ................................................\n")
    run_school.font.name = VietFonts.BODY
    run_school.font.size = VietFonts.SIZE_INFO
    run_school.font.color.rgb = VietColors.DARK

    run_to = p.add_run(f"Tổ: {identity.department}")
    run_to.italic = True
    run_to.font.name = VietFonts.BODY
    run_to.font.size = VietFonts.SIZE_INFO
    run_to.font.color.rgb = VietColors.GRAY

    styles.add_title(doc, "KẾ HOẠCH BÀI DẠY")

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(8)
    run_sub = p_sub.add_run("(Kèm theo Công văn số 5512/BGDĐT-GDTrH ngày 18/12/2020 của Bộ GDĐT)")
    run_sub.font.name = VietFonts.BODY
    run_sub.font.size = VietFonts.SIZE_INFO
    run_sub.font.color.rgb = VietColors.GRAY
    run_sub.italic = True

    layout = get_layout("khbd")
    table = doc.add_table(rows=3, cols=4)
    styles.set_table_width_dxa(table, layout.content_width_cm)
    col_widths = [Cm(4.0)] * 4
    metadata = [
        ["Họ và tên GV:", identity.teacher_name, "Môn học:", "Vật lí"],
        ["Tên bài dạy:", lesson_title, "Lớp:", "12"],
        ["Thời gian thực hiện:", duration_str, "", ""],
    ]
    for r_idx, row_data in enumerate(metadata):
        row = table.rows[r_idx]
        styles.mark_row_no_split(row)
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.width = col_widths[c_idx]
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.left_indent = Cm(0.2)
            run = p.add_run(val)
            run.bold = c_idx % 2 == 0
            run.font.name = VietFonts.BODY
            run.font.size = VietFonts.SIZE_INFO
            run.font.color.rgb = VietColors.DARK
            styles.set_cell_margins(cell, top_pt=5.0, bottom_pt=5.0, left_pt=5.0, right_pt=5.0)
            styles.set_cell_borders(cell)
    doc.add_paragraph()


def _add_signature_block(doc) -> None:
    layout = get_layout("khbd")
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_date.paragraph_format.space_before = Pt(14)
    p_date.paragraph_format.right_indent = Cm(1.0)
    run_date = p_date.add_run("........., ngày ....... tháng ....... năm 202...")
    run_date.font.name = VietFonts.BODY
    run_date.font.size = VietFonts.SIZE_INFO
    run_date.italic = True
    run_date.font.color.rgb = VietColors.DARK

    table = doc.add_table(rows=2, cols=3)
    styles.set_table_width_dxa(table, layout.content_width_cm)
    col_widths = [Cm(5.3), Cm(5.3), Cm(5.4)]
    titles = [
        "XÁC NHẬN CỦA BAN GIÁM HIỆU\n(Ký, đóng dấu)",
        "XÁC NHẬN CỦA TỔ TRƯỞNG\n(Ký và ghi rõ họ tên)",
        "GIÁO VIÊN SOẠN BÀI\n(Ký và ghi rõ họ tên)",
    ]
    row0 = table.rows[0]
    styles.mark_row_no_split(row0)
    for idx, text in enumerate(titles):
        cell = row0.cells[idx]
        cell.width = col_widths[idx]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        run.bold = True
        run.font.name = VietFonts.BODY
        run.font.size = VietFonts.SIZE_INFO
        run.font.color.rgb = VietColors.DARK
        styles.remove_cell_borders(cell)

    row1 = table.rows[1]
    styles.mark_row_no_split(row1)
    for idx in range(3):
        cell = row1.cells[idx]
        cell.width = col_widths[idx]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(45)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run("..................................................")
        run.font.name = VietFonts.BODY
        run.font.size = VietFonts.SIZE_INFO
        run.font.color.rgb = VietColors.GRAY
        styles.remove_cell_borders(cell)


def _add_adjustments_section(doc, advantages: str, limitations: str, solutions: str) -> None:
    styles.add_heading(doc, "IV. ĐIỀU CHỈNH BÀI DẠY SAU TIẾT GIẢNG", level=1)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.3

    placeholder = "." * 120

    def _write_field(label: str, value: str) -> None:
        r_label = p.add_run(f"{label}\n")
        r_label.bold = True
        p.add_run((value or placeholder) + "\n")
        r_extra = p.add_run(f"   {placeholder}\n")
        r_extra.font.color.rgb = VietColors.GRAY

    _write_field("1. Ưu điểm nổi bật:", advantages)
    _write_field("2. Hạn chế, khó khăn:", limitations)
    _write_field("3. Hướng điều chỉnh sư phạm:", solutions)

    for run in p.runs:
        run.font.name = VietFonts.BODY
        run.font.size = VietFonts.SIZE_BODY
        if run.text.startswith((" ", "   ")):
            run.font.color.rgb = VietColors.GRAY


def _extract_title_and_duration(lines: list[str]) -> tuple[str, str]:
    lesson_title, duration_str = "Bài học", "02 tiết"
    for line in lines[:10]:
        if "Tên bài dạy:" in line:
            lesson_title = line.replace("**Tên bài dạy:**", "").replace("*", "").strip()
        if "Thời gian thực hiện:" in line:
            duration_str = line.replace("**Thời gian thực hiện:**", "").replace("*", "").strip()
    return lesson_title, duration_str


def _strip_admin_header(lines: list[str]) -> list[str]:
    cleaned: list[str] = []
    skip_header = True
    for line in lines:
        stripped = line.strip()
        if stripped.startswith("## I.") or stripped.startswith("## I ") or "## I. YÊU CẦU CẦN ĐẠT" in stripped.upper():
            skip_header = False
        if not skip_header:
            cleaned.append(line)
    return cleaned


def _extract_reflection_section(lines: list[str]) -> tuple[list[str], str, str, str]:
    """Tách phần 'IV. ĐIỀU CHỈNH BÀI DẠY' ra khỏi phần thân — phần này được
    dựng lại bằng `_add_adjustments_section` với heading chuẩn hoá riêng,
    không dùng heading gốc trong markdown nguồn."""
    body_lines: list[str] = []
    advantages = limitations = solutions = ""
    in_reflection = False
    for line in lines:
        stripped = line.strip()
        if (stripped.startswith("## IV. ĐIỀU CHỈNH BÀI DẠY")
                or stripped.startswith("## IV. RÚT KINH NGHIỆM")
                or stripped.startswith("## IV ")):
            in_reflection = True
            continue
        if in_reflection:
            if "Ưu điểm:" in stripped:
                advantages = re.sub(r"^\d*\.?\s*Ưu điểm:", "", stripped).strip()
            elif "Hạn chế:" in stripped:
                limitations = re.sub(r"^\d*\.?\s*Hạn chế:", "", stripped).strip()
            elif "Hướng điều chỉnh:" in stripped:
                solutions = re.sub(r"^\d*\.?\s*Hướng điều chỉnh:", "", stripped).strip()
            continue
        body_lines.append(line)
    return body_lines, advantages, limitations, solutions


def export_khbd(md_path: str, output_path: str, identity: Identity) -> None:
    md_file = Path(md_path)
    doc = Document()
    styles.setup_margins(doc, "khbd")
    styles.set_default_style(doc)
    styles.add_page_number_footer(doc)

    if not md_file.exists():
        doc.add_paragraph(f"Lỗi: Không tìm thấy file Markdown tại {md_path}")
        doc.save(output_path)
        return

    raw_lines = md_file.read_text(encoding="utf-8").splitlines()
    lesson_title, duration_str = _extract_title_and_duration(raw_lines)
    _add_admin_header(doc, lesson_title, duration_str, identity)

    body_lines = _strip_admin_header(raw_lines)
    body_lines, advantages, limitations, solutions = _extract_reflection_section(body_lines)

    layout = get_layout("khbd")
    image_dirs = [md_file.parent.parent]
    blocks = parse_markdown_blocks(body_lines, image_search_dirs=image_dirs)
    for block in blocks:
        if block.kind == "heading":
            styles.add_heading(doc, block.text, level=block.level)
        elif block.kind == "bullet":
            styles.add_bullet(doc, block.text, level=block.level)
        elif block.kind == "blockquote":
            styles.add_body(doc, block.text, is_blockquote=True)
        elif block.kind == "dotline":
            styles.add_dot_line(doc, count=1)
        elif block.kind == "separator":
            styles.add_separator(doc)
        elif block.kind == "table":
            headers = block.table_headers or []
            if len(headers) == 2 and "Sản phẩm" in headers:
                _make_process_table(doc, headers, block.table_rows, layout)
            else:
                styles.make_table(doc, headers, block.table_rows, layout)
        elif block.text:
            styles.add_body(doc, block.text)
        embed_image(doc, block.image_path, 12.0)

    _add_adjustments_section(doc, advantages, limitations, solutions)
    _add_signature_block(doc)

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
