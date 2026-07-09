"""Xuất Phiếu học tập (PHT) từ Markdown sang DOCX."""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from gems.docx_export import styles
from gems.docx_export.layout import get_layout
from gems.docx_export.markdown_ir import parse_markdown_blocks
from gems.docx_export.palette import VietColors, VietFonts
from gems.docx_export.renderer import render_blocks

_BOLD_ONLY_RE = re.compile(r"^\*\*(.+)\*\*$")


def _add_student_info_table(doc, lesson_title: str) -> None:
    """Bảng thông tin học sinh 2 cột x 3 dòng, tổng 17.0cm (khớp layout PHT)."""
    layout = get_layout("pht")
    table = doc.add_table(rows=3, cols=2)
    styles.set_table_width_dxa(table, layout.content_width_cm)
    col_widths = [Cm(9.5), Cm(7.5)]

    metadata = [
        ["Họ và tên học sinh: .......................................", "Lớp: .................  Nhóm: ................."],
        ["Môn học: Vật lí 12 (CTGDPT 2018)", f"Bài học: {lesson_title}"],
        ["Tiết học: ............ ngày ......../......../202...", "Điểm số: .........  Chữ ký GV: ........."],
    ]
    for r_idx, row_data in enumerate(metadata):
        row = table.rows[r_idx]
        styles.mark_row_no_split(row)
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.width = col_widths[c_idx]
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.left_indent = Cm(0.2)
            run = p.add_run(val)
            run.font.name = VietFonts.BODY
            run.font.size = VietFonts.SIZE_INFO
            run.font.color.rgb = VietColors.DARK
            styles.set_cell_margins(cell, top_pt=6.0, bottom_pt=6.0, left_pt=6.0, right_pt=6.0)
            styles.set_cell_borders(cell)
    doc.add_paragraph()


def _add_self_evaluation_section(doc) -> None:
    styles.add_heading(doc, "III. TỰ ĐÁNH GIÁ KẾT QUẢ", level=1)
    layout = get_layout("pht")
    headers = ["Tiêu chí học tập", "Đạt", "Chưa Đạt"]
    rows = [
        ["1. Nhận biết và phát biểu được các khái niệm cốt lõi của bài học.", "   [  ]", "   [  ]"],
        ["2. Áp dụng đúng công thức tính toán và giải quyết bài tập vận dụng thực tế.", "   [  ]", "   [  ]"],
        ["3. Tích cực thảo luận, hoàn thành đầy đủ nhiệm vụ học tập nhóm/cá nhân.", "   [  ]", "   [  ]"],
    ]
    styles.make_table(doc, headers, rows, layout, col_widths_cm=[12.0, 2.5, 2.5])

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.right_indent = Cm(1.0)

    run_date = p.add_run("Ngày ........ tháng ........ năm 202...\n")
    run_date.font.name = VietFonts.BODY
    run_date.font.size = VietFonts.SIZE_INFO
    run_date.font.color.rgb = VietColors.DARK
    run_date.italic = True

    run_sig = p.add_run("Học sinh ký tên\n\n\n..........................................")
    run_sig.bold = True
    run_sig.font.name = VietFonts.BODY
    run_sig.font.size = VietFonts.SIZE_INFO
    run_sig.font.color.rgb = VietColors.DARK


def _strip_duplicated_header(lines: list[str]) -> list[str]:
    """Bỏ phần header lặp lại (Họ tên/Lớp...) do model tự sinh ở đầu file —
    nội dung thật bắt đầu từ dòng `##` đầu tiên hoặc dòng chứa 📍."""
    cleaned: list[str] = []
    skip_header = True
    for line in lines:
        stripped = line.strip()
        if stripped.startswith("##") or "📍" in stripped:
            skip_header = False
        if not skip_header:
            cleaned.append(line)
    return cleaned


def _promote_short_bold_lines_to_subheading(lines: list[str]) -> list[str]:
    """Một dòng CHỈ chứa `**...**` (dưới 100 ký tự) là nhãn nhiệm vụ
    (vd. `**Nhiệm vụ: Đo nhiệt độ (Matching Matrix)**`) — coi như tiêu đề
    cấp 3 để có kiểu chữ nghiêng nổi bật, thay vì chữ đậm nằm giữa đoạn văn."""
    result = []
    for line in lines:
        stripped = line.strip()
        match = _BOLD_ONLY_RE.match(stripped)
        if match and len(stripped) < 100:
            result.append(f"#### {match.group(1).strip()}")
        else:
            result.append(line)
    return result


def export_pht(md_path: str, output_path: str, lesson_label: str) -> None:
    md_file = Path(md_path)
    doc = Document()
    styles.setup_margins(doc, "pht")
    styles.set_default_style(doc)
    styles.add_page_number_footer(doc)

    p_header = doc.add_paragraph()
    p_header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_header.paragraph_format.space_after = Pt(2)
    run_sch = p_header.add_run("TRƯỜNG THPT: ................................................")
    run_sch.font.name = VietFonts.BODY
    run_sch.font.size = VietFonts.SIZE_INFO
    run_sch.italic = True
    run_sch.font.color.rgb = VietColors.GRAY

    styles.add_title(doc, "PHIẾU HỌC TẬP")

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(10)
    run_sub = p_sub.add_run("(Mẫu tài liệu học tập theo CTGDPT 2018)")
    run_sub.font.name = VietFonts.BODY
    run_sub.font.size = VietFonts.SIZE_TINY
    run_sub.font.color.rgb = VietColors.GRAY
    run_sub.italic = True

    _add_student_info_table(doc, lesson_label)

    if not md_file.exists():
        doc.add_paragraph(f"Lỗi: Không tìm thấy file Markdown tại {md_path}")
        doc.save(output_path)
        return

    raw_lines = md_file.read_text(encoding="utf-8").splitlines()
    lines = _strip_duplicated_header(raw_lines)
    lines = _promote_short_bold_lines_to_subheading(lines)

    layout = get_layout("pht")
    image_dirs = [md_file.parent.parent]
    blocks = parse_markdown_blocks(lines, image_search_dirs=image_dirs)
    render_blocks(doc, blocks, layout, image_width_cm=12.0)

    _add_self_evaluation_section(doc)

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
