"""Xuất Bài tập về nhà từ Markdown sang DOCX.

Phần thân bài (câu hỏi cho học sinh) có ngữ pháp riêng (trắc nghiệm ABCD,
Đúng/Sai 4 ý, trả lời ngắn) nên được parse thủ công ở đây. Phần "ĐÁP ÁN VÀ
HƯỚNG DẪN GIẢI CHI TIẾT" ở cuối dùng chung `markdown_ir`/`renderer` — bản cũ
tự viết 1 parser riêng yếu hơn cho phần này (không có nhánh dự phòng cho
đoạn văn thường trong Phần II/III, làm rớt nội dung không rõ ràng); dùng lại
bộ phân tích chung ở đây vừa đơn giản vừa không còn rớt nội dung.
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from gems.config.loader import Identity
from gems.docx_export import styles
from gems.docx_export.layout import get_layout
from gems.docx_export.markdown_ir import _extract_image, parse_markdown_blocks
from gems.docx_export.palette import VietColors, VietFonts
from gems.docx_export.renderer import embed_image, render_blocks
from gems.docx_export.run_formatter import add_formatted_runs

_QUESTION_STEM_RE = re.compile(r"^\*{0,2}(?:Câu\s+)?(\d+)(?:\s*\([^)]*\))?[\.:\)]\s*\*{0,2}\s*(.*)$")
_OPTION_RE = re.compile(r"^\*{0,2}([A-Da-d])[\.\)]\s*\*{0,2}\s*(.*)$")
_STATEMENT_RE = re.compile(r"^\*{0,2}([a-d])[\.\)]\s*\*{0,2}\s*(.*)$")
_LEAKED_ANSWER_RE = re.compile(r"\(?\s*Đ/S\s*:?[^)]*\)?\s*$", re.IGNORECASE)
_CHECKBOX_ONLY_RE = re.compile(r"^☐?\s*Đ/S", re.IGNORECASE)
_CHO_BIET_RE = re.compile(r"^\*Cho biết:\*\s*(.*)$")

_ANSWER_MARKERS = ("ĐÁP ÁN", "HƯỚNG DẪN GIẢI")
_PART_HEADING_RE = re.compile(r"^#{0,4}\s*PHẦN\s+(III|II|I)(?=\s*[:.])", re.IGNORECASE)


def _strip_leaked_answer(text: str) -> str:
    """Bản gốc từng ghi thẳng đáp án Đúng/Sai vào ngay câu dẫn Phần II khi
    sinh markdown (`(Đ/S: Đ)`) — lộ đáp án cho học sinh ngay trong đề. Loại
    bỏ mọi annotation kiểu này khỏi phần in cho học sinh; đáp án thật chỉ
    nằm ở mục "ĐÁP ÁN..." cuối file."""
    return _LEAKED_ANSWER_RE.sub("", text).strip()


def _add_question_stem(doc, number: int, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.keep_together = True
    run_num = p.add_run(f"Câu {number}. ")
    run_num.bold = True
    run_num.font.name = VietFonts.BODY
    run_num.font.size = VietFonts.SIZE_BODY
    run_num.font.color.rgb = VietColors.DARK
    add_formatted_runs(p, text)


def _add_mcq_options(doc, options: list[str], layout) -> None:
    max_len = max((len(opt) for opt in options), default=0)
    if max_len < 18:
        cols, rows_layout = 4, [options[:4]]
    elif max_len < 38:
        cols = 2
        rows_layout = [options[i:i + 2] for i in range(0, len(options), 2)]
    else:
        cols, rows_layout = 1, [[opt] for opt in options]

    labels = ["A", "B", "C", "D"]
    table = doc.add_table(rows=len(rows_layout), cols=cols)
    table.autofit = False
    styles.set_table_width_dxa(table, layout.content_width_cm)
    col_width = layout.content_width_cm / cols
    idx = 0
    for r_idx, row in enumerate(rows_layout):
        for c_idx in range(cols):
            cell = table.rows[r_idx].cells[c_idx]
            cell.width = Cm(col_width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            styles.remove_cell_borders(cell)
            styles.set_cell_margins(cell, top_pt=0.5, bottom_pt=0.5, left_pt=4.0, right_pt=4.0)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            if c_idx < len(row):
                run_label = p.add_run(f"{labels[idx]}. ")
                run_label.bold = True
                run_label.font.name = VietFonts.BODY
                run_label.font.size = VietFonts.SIZE_INFO
                add_formatted_runs(p, row[c_idx])
                idx += 1


def _add_tf_statement_table(doc, statements: list[str], layout) -> None:
    """Đề mẫu in bốn nhận định a)-d) dạng đoạn văn, không có bảng/ô chọn."""
    labels = ["a", "b", "c", "d"]
    for label, stmt in zip(labels, statements):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Cm(0.65)
        p.paragraph_format.first_line_indent = Cm(-0.45)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = 1.0
        run_label = p.add_run(f"{label}) ")
        run_label.bold = True
        run_label.font.name = VietFonts.BODY
        run_label.font.size = VietFonts.SIZE_INFO
        add_formatted_runs(p, stmt)


def _add_short_answer_box(doc) -> None:
    """Đề mẫu dành khoảng trống tự nhiên sau câu hỏi, không in hộp đáp số."""
    return None


def _add_header(doc, lesson_title: str, constants_text: str) -> None:
    """Header đơn giản như PHT/KHBD: tiêu đề bài + Họ tên học sinh/Lớp/Ngày.
    Trước đây mô phỏng hình thức đề thi tốt nghiệp THPT thật (Sở GD&ĐT,
    Trường THPT, Mã đề thi, Số báo danh) — bỏ vì đây là bài tập về nhà bình
    thường, không cần đóng giả hình thức đề thi; cấu trúc CÂU HỎI (18 MCQ +
    4 Đúng/Sai + 6 trả lời ngắn, phong cách bám đề tốt nghiệp) vẫn giữ
    nguyên, chỉ đổi phần vỏ hình thức."""
    styles.setup_margins(doc, "homework")
    styles.set_default_style(doc)
    styles.add_page_number_footer(doc)

    styles.add_title(doc, lesson_title)

    p_student = doc.add_paragraph()
    p_student.paragraph_format.space_after = Pt(0)
    r_name = p_student.add_run("Họ và tên học sinh: ..........................................................")
    r_name.font.name = VietFonts.BODY
    r_name.font.size = VietFonts.SIZE_INFO

    p_lop = doc.add_paragraph()
    p_lop.paragraph_format.space_after = Pt(4)
    r_lop = p_lop.add_run("Lớp: .................. Ngày: ......./......./.......")
    r_lop.font.name = VietFonts.BODY
    r_lop.font.size = VietFonts.SIZE_INFO

    if constants_text:
        p_const = doc.add_paragraph()
        p_const.paragraph_format.space_before = Pt(2)
        p_const.paragraph_format.space_after = Pt(0)
        run_const = p_const.add_run(f"+ Cho biết: {constants_text}")
        run_const.bold = True
        run_const.font.name = VietFonts.BODY
        run_const.font.size = VietFonts.SIZE_INFO
        run_const.font.color.rgb = VietColors.DARK

        p_round = doc.add_paragraph()
        p_round.paragraph_format.space_before = Pt(0)
        p_round.paragraph_format.space_after = Pt(2)
        run_round = p_round.add_run("+ Không làm tròn kết quả các phép tính trung gian.")
        run_round.bold = True
        run_round.font.name = VietFonts.BODY
        run_round.font.size = VietFonts.SIZE_INFO

    styles.add_separator(doc)


def _add_part_heading(doc, part: str) -> None:
    descriptions = {
        "I": "Thí sinh trả lời từ câu 1 đến câu 18. Mỗi câu hỏi thí sinh chỉ chọn một phương án.",
        "II": "Thí sinh trả lời từ câu 1 đến câu 4. Trong mỗi ý a), b), c), d) ở mỗi câu, thí sinh chọn đúng hoặc sai.",
        "III": "Thí sinh trả lời từ câu 1 đến câu 6.",
    }
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.0
    r_part = p.add_run(f"PHẦN {part}. ")
    r_part.bold = True
    r_part.font.name = VietFonts.BODY
    r_part.font.size = VietFonts.SIZE_INFO
    r_desc = p.add_run(descriptions[part])
    r_desc.font.name = VietFonts.BODY
    r_desc.font.size = VietFonts.SIZE_INFO


def _split_answer_section(lines: list[str]) -> tuple[list[str], list[str]]:
    for i, line in enumerate(lines):
        upper = line.strip().upper()
        if any(marker in upper for marker in _ANSWER_MARKERS):
            return lines[:i], lines[i:]
    return lines, []


def _render_student_body(doc, lines: list[str], layout, image_dirs: list[Path]) -> None:
    """Đề mẫu có ngữ pháp riêng (trắc nghiệm/Đúng-Sai/trả lời ngắn) nên KHÔNG
    dùng `parse_markdown_blocks` chung cho toàn bộ — nhưng vẫn hỗ trợ nhúng
    ảnh (`ready/hinh_anh/*.png`) qua đúng cùng cơ chế `markdown_ir._extract_image`
    dùng ở PHT/KHBD, áp dụng cho các dòng KHÔNG phải câu dẫn/phương án/mệnh đề
    (ví dụ dòng dữ kiện dùng chung `*Nội dung câu X và Y: ...*`). KHÔNG đặt
    ảnh xen giữa 1 câu dẫn Phần I và các phương án A-D, hoặc giữa 1 câu dẫn
    Phần II và 4 mệnh đề a-d — 2 nhánh đó quét tiếp các dòng ngay sau câu dẫn
    theo đúng mẫu phương án/mệnh đề, gặp dòng ảnh sẽ dừng quét sớm."""
    current_part = None
    i = 0
    n = len(lines)
    while i < n:
        stripped = lines[i].strip()
        if not stripped:
            i += 1
            continue

        upper = stripped.upper()
        part_match = _PART_HEADING_RE.match(stripped)
        matched_part = part_match.group(1).upper() if part_match else None
        if matched_part:
            current_part = matched_part
            _add_part_heading(doc, matched_part)
            i += 1
            continue

        stem_match = _QUESTION_STEM_RE.match(stripped) if current_part else None
        if stem_match and current_part == "I":
            number = int(stem_match.group(1))
            _add_question_stem(doc, number, stem_match.group(2))
            options: list[str] = []
            j = i + 1
            while j < n and len(options) < 4:
                opt_match = _OPTION_RE.match(lines[j].strip())
                if not opt_match:
                    break
                options.append(opt_match.group(2))
                j += 1
            if len(options) >= 3:
                _add_mcq_options(doc, options, layout)
                i = j
                continue
            i += 1
            continue

        if stem_match and current_part == "II":
            number = int(stem_match.group(1))
            _add_question_stem(doc, number, stem_match.group(2))
            statements: list[str] = []
            j = i + 1
            while j < n and len(statements) < 4:
                candidate = lines[j].strip()
                if _CHECKBOX_ONLY_RE.match(candidate):
                    j += 1
                    continue
                stmt_match = _STATEMENT_RE.match(candidate)
                if not stmt_match:
                    break
                statements.append(_strip_leaked_answer(stmt_match.group(2)))
                j += 1
            if statements:
                _add_tf_statement_table(doc, statements, layout)
                i = j
                continue
            i += 1
            continue

        if stem_match and current_part == "III":
            number = int(stem_match.group(1))
            _add_question_stem(doc, number, stem_match.group(2))
            _add_short_answer_box(doc)
            i += 1
            continue

        if current_part in ("I", "II", "III") and not _CHECKBOX_ONLY_RE.match(stripped):
            cleaned, image_path = _extract_image(stripped, image_dirs)
            if image_path is not None:
                embed_image(doc, image_path, 12.0)
            elif cleaned:
                styles.add_body(doc, cleaned, size=VietFonts.SIZE_INFO)
        i += 1


def _add_end_notice(doc) -> None:
    """Dòng phân cách '----- HẾT -----' đứng ngay trước phần đáp án (đáp án
    tách trang riêng). Không còn 2 dòng "Thí sinh không được sử dụng tài
    liệu/Giám thị không giải thích gì thêm" — đó là ngôn ngữ đóng vai coi
    thi, không phù hợp một bài tập về nhà bình thường."""
    p_end = doc.add_paragraph()
    p_end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_end.paragraph_format.space_before = Pt(12)
    run_end = p_end.add_run("----- HẾT -----")
    run_end.bold = True
    run_end.font.name = VietFonts.BODY
    run_end.font.size = VietFonts.SIZE_INFO
    run_end.font.color.rgb = VietColors.DARK


def export_homework(md_path: str, output_path: str, identity: Identity, lesson_label: str = "") -> None:
    md_file = Path(md_path)
    doc = Document()

    if not md_file.exists():
        styles.setup_margins(doc, "homework")
        doc.add_paragraph(f"Lỗi: Không tìm thấy file Markdown tại {md_path}")
        doc.save(output_path)
        return

    raw_lines = md_file.read_text(encoding="utf-8").splitlines()
    title_line = next((line.lstrip("#").strip() for line in raw_lines if line.strip().startswith("#")), lesson_label)
    constants_text = next((m.group(1) for line in raw_lines if (m := _CHO_BIET_RE.match(line.strip()))), "")

    _add_header(doc, title_line or f"BÀI TẬP VỀ NHÀ: {lesson_label}", constants_text)

    layout = get_layout("homework")
    image_dirs = [md_file.parent.parent]
    student_lines, answer_lines = _split_answer_section(raw_lines)
    _render_student_body(doc, student_lines, layout, image_dirs)
    _add_end_notice(doc)

    if answer_lines:
        doc.add_page_break()
        styles.add_heading(doc, "ĐÁP ÁN VÀ HƯỚNG DẪN GIẢI CHI TIẾT", level=1)
        blocks = parse_markdown_blocks(answer_lines[1:], image_search_dirs=image_dirs)
        render_blocks(doc, blocks, layout)

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
