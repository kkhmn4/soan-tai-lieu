"""Render danh sách Block (từ markdown_ir) ra tài liệu python-docx.

Dùng chung cho phần "80% bề mặt chung" của cả 3 loại tài liệu — phần ngữ
pháp riêng (câu hỏi trắc nghiệm, bảng GV/HS 4 pha CV5512...) do exporter
tương ứng tự xử lý bằng cách gọi thẳng các hàm trong `styles.py`.
"""
from __future__ import annotations

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm

from gems.docx_export import styles
from gems.docx_export.layout import DocLayout
from gems.docx_export.markdown_ir import Block


def embed_image(doc, image_path, width_cm: float) -> None:
    if image_path is None:
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))


def render_blocks(doc, blocks: list[Block], layout: DocLayout, *, image_width_cm: float = 12.0) -> None:
    for block in blocks:
        if block.kind == "heading":
            styles.add_heading(doc, block.text, level=block.level)
        elif block.kind == "bullet":
            styles.add_bullet(doc, block.text, level=block.level)
        elif block.kind == "blockquote":
            styles.add_body(doc, block.text, is_blockquote=True)
        elif block.kind == "dotline":
            styles.add_dot_line(doc, count=1)
        elif block.kind == "separator":
            styles.add_separator(doc)
        elif block.kind == "table":
            styles.make_table(doc, block.table_headers or [], block.table_rows, layout)
        else:  # "body"
            if block.text:
                styles.add_body(doc, block.text)

        embed_image(doc, block.image_path, image_width_cm)
