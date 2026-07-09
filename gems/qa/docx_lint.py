"""Kiểm định cấu trúc file .docx đã biên dịch — đối chiếu với `layout.py`
(không hardcode lại số liệu lề/độ rộng như bản gốc `check_khbd_standards.py`,
vốn cũng chỉ là 1 script rời không được gọi từ pipeline chính).
"""
from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path

from docx import Document

from gems.docx_export.layout import get_layout
from gems.docx_export.palette import VietFonts


@dataclass
class DocxLintReport:
    file_name: str
    issues: list[str] = field(default_factory=list)

    @property
    def passed(self) -> bool:
        return not self.issues


def _iter_runs(doc):
    for p in doc.paragraphs:
        yield from p.runs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield from p.runs


def check_margins(doc, doc_type: str) -> str | None:
    layout = get_layout(doc_type)
    section = doc.sections[0]
    actual = (round(section.top_margin.cm, 1), round(section.bottom_margin.cm, 1),
              round(section.left_margin.cm, 1), round(section.right_margin.cm, 1))
    expected = (layout.top_margin_cm, layout.bottom_margin_cm, layout.left_margin_cm, layout.right_margin_cm)
    if actual != expected:
        return f"Lề trang {actual} không khớp chuẩn {doc_type} {expected} (T/B/L/R, cm)."
    return None


def check_font_consistency(doc) -> str | None:
    fonts = {run.font.name for run in _iter_runs(doc) if run.font.name}
    if fonts and fonts != {VietFonts.BODY}:
        return f"Phát hiện font chữ khác {VietFonts.BODY}: {sorted(fonts)}."
    return None


def check_khbd_heading_structure(doc) -> str | None:
    headings_text = " ".join(p.text.upper() for p in doc.paragraphs if p.style.name.startswith("Heading"))
    missing = [label for label, needle in [
        ("I. YÊU CẦU CẦN ĐẠT", "YÊU CẦU CẦN ĐẠT"),
        ("II. THIẾT BỊ DẠY HỌC VÀ HỌC LIỆU", "THIẾT BỊ DẠY HỌC"),
        ("III. TIẾN TRÌNH DẠY HỌC", "TIẾN TRÌNH DẠY HỌC"),
    ] if needle not in headings_text]
    if missing:
        return f"KHBD thiếu các mục lớn bắt buộc: {', '.join(missing)}."
    return None


def check_khbd_gv_hs_table(doc) -> str | None:
    gv_hs_tables = []
    for table in doc.tables:
        if len(table.columns) != 2:
            continue
        headers = [c.text.strip().lower() for c in table.rows[0].cells]
        if any("giáo viên" in h or "gv" in h for h in headers) or any("học sinh" in h or "hs" in h for h in headers):
            gv_hs_tables.append(table)
    if not gv_hs_tables:
        return "Không tìm thấy bảng hoạt động Giáo viên/Học sinh 2 cột nào."
    if len(gv_hs_tables[0].rows) > 1:
        for run in _iter_runs_in_cell(gv_hs_tables[0].rows[1].cells[0]):
            if run.font.size and round(run.font.size.pt) > 11:
                return "Cỡ chữ trong bảng hoạt động GV/HS lớn hơn 11pt."
    return None


def _iter_runs_in_cell(cell):
    for p in cell.paragraphs:
        yield from p.runs


def check_khbd_reflection_and_signature(doc) -> str | None:
    has_reflection = any("ĐIỀU CHỈNH BÀI DẠY" in p.text.upper() or "RÚT KINH NGHIỆM" in p.text.upper()
                          for p in doc.paragraphs)
    has_signature = False
    if doc.tables:
        last_table = doc.tables[-1]
        if len(last_table.columns) == 3:
            header_texts = [c.text.strip().upper() for c in last_table.rows[0].cells]
            has_signature = any("TỔ TRƯỞNG" in t or "GIÁO VIÊN" in t for t in header_texts)
    problems = []
    if not has_reflection:
        problems.append("thiếu phần Điều chỉnh bài dạy/Rút kinh nghiệm")
    if not has_signature:
        problems.append("thiếu bảng chữ ký phê duyệt 3 bên ở cuối")
    return "; ".join(problems) if problems else None


def lint_khbd_docx(path: str | Path) -> DocxLintReport:
    doc = Document(str(path))
    report = DocxLintReport(file_name=Path(path).name)
    issues = [
        check_margins(doc, "khbd"),
        check_font_consistency(doc),
        check_khbd_heading_structure(doc),
        check_khbd_gv_hs_table(doc),
        check_khbd_reflection_and_signature(doc),
    ]
    report.issues = [issue for issue in issues if issue]
    return report


def lint_docx_margins(path: str | Path, doc_type: str) -> DocxLintReport:
    """Kiểm tra lề trang cho PHT/Bài tập (không có cấu trúc heading/bảng
    đặc thù như KHBD nên chỉ cần kiểm tra lề + font)."""
    doc = Document(str(path))
    report = DocxLintReport(file_name=Path(path).name)
    for issue in (check_margins(doc, doc_type), check_font_consistency(doc)):
        if issue:
            report.issues.append(issue)
    return report
