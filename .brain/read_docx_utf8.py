import docx
import sys
import io

# Set encoding to utf-8 for stdout
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

def extract_docx(docx_path, txt_path):
    doc = docx.Document(docx_path)
    with open(txt_path, 'w', encoding='utf-8') as f:
        f.write("=== PARAGRAPHS ===\n")
        for para in doc.paragraphs:
            if para.text.strip():
                f.write(para.text + "\n")
        
        f.write("\n=== TABLES ===\n")
        for ti, table in enumerate(doc.tables):
            f.write(f"\n--- TABLE {ti} ---\n")
            for ri, row in enumerate(table.rows):
                cells = []
                for cell in row.cells:
                    cell_text = ' '.join(p.text.strip() for p in cell.paragraphs if p.text.strip())
                    cells.append(cell_text)
                f.write(f"ROW {ri}: " + " || ".join(cells) + "\n")

if __name__ == '__main__':
    extract_docx(
        r"c:\Users\Admin\.antigravity-ide\soạn tài liệu\tai-lieu-goc\L12 Chương 1 - Tailieuchuan\Chủ đề 4. NHIỆT DUNG RIÊNG - GV.docx",
        r"c:\Users\Admin\.antigravity-ide\soạn tài liệu\output\nhiet_dung_rieng_raw.txt"
    )
    print("Done extracting!")
