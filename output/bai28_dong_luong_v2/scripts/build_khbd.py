"""Soạn Kế hoạch bài dạy (KHBD) V2 cho Bài 28 - Động lượng theo ĐÚNG mẫu
"KẾ HOẠCH GIẢNG DẠY.docx" (mẫu Trường THCS và THPT Đinh Thiện Lý).

BẢN CẬP NHẬT (sau vòng phản hồi 1): Bài 28 dạy trong 2 TIẾT, KHÔNG gộp 1 tiết
như bản nháp đầu — Tiết 1 = "I. Động lượng" (★ TIẾT DẠY MINH HỌA, trọng tâm
chất lượng), Tiết 2 = "II. Xung lượng của lực". Khởi động Tiết 1 dùng trò
chơi AR "Bảo vệ hành tinh xanh" (HS điều khiển tên lửa bắn thiên thạch đổi
hướng). Loại hình nhiệm vụ đa dạng hoá lại, ưu tiên loại hình gây hào hứng ở
Tiết 1 (game, meme, TikTok, Decision Tree) — xem `dac_ta_v2.md` mục 3.

Tài liệu one-off, không dùng `khbd_exporter.py`, không sửa file nào trong
`gems/`. Lượt này KHÔNG nhúng ảnh minh họa.

Chạy: python "output/bai28_dong_luong_v2/scripts/build_khbd.py"
"""
from __future__ import annotations

import sys
from pathlib import Path

sys.stdout.reconfigure(encoding="utf-8")

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Emu, Pt, RGBColor

REPO_ROOT = Path(__file__).resolve().parent.parent.parent.parent
OUT_PATH = REPO_ROOT / "output" / "bai28_dong_luong_v2" / "ready" / "bai28_dong_luong_v2_ke_hoach_bai_day.docx"

FONT = "Times New Roman"
PRIMARY = RGBColor(0x1F, 0x4E, 0x79)
RED = RGBColor(0xC0, 0x00, 0x00)
DARK = RGBColor(0x00, 0x00, 0x00)


def dxa(value: int) -> Emu:
    return Emu(value * 635)


def set_default_font(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(12)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), FONT)


def set_cell_borders(cell) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    borders = tcPr.makeelement(qn("w:tcBorders"), {})
    for edge in ("top", "left", "bottom", "right"):
        el = borders.makeelement(qn(f"w:{edge}"), {})
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        borders.append(el)
    tcPr.append(borders)


def add_run(p, text: str, *, bold=False, italic=False, size=12, color=DARK, highlight=None) -> None:
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.name = FONT
    r.font.size = Pt(size)
    r.font.color.rgb = color
    if highlight:
        rpr = r._element.get_or_add_rPr()
        hl = rpr.makeelement(qn("w:highlight"), {})
        hl.set(qn("w:val"), highlight)
        rpr.append(hl)


def cell_para(cell, text: str = "", *, bold=False, italic=False, size=12, color=DARK,
              space_after=2, first=False):
    p = cell.paragraphs[0] if first and not cell.paragraphs[0].runs else cell.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    if text:
        add_run(p, text, bold=bold, italic=italic, size=size, color=color)
    return p


def cell_multiline(cell, lines: list[tuple], *, first_clear=True):
    if first_clear:
        cell.paragraphs[0].text = ""
    for i, (text, kwargs) in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(kwargs.pop("space_after", 3))
        p.paragraph_format.space_before = Pt(0)
        add_run(p, text, **kwargs)


# ============================================================
# DỰNG TÀI LIỆU
# ============================================================
doc = Document()
set_default_font(doc)

section = doc.sections[0]
section.orientation = WD_ORIENT.LANDSCAPE
section.page_width = dxa(15840)
section.page_height = dxa(12240)
section.top_margin = dxa(720)
section.bottom_margin = dxa(720)
section.left_margin = dxa(720)
section.right_margin = dxa(720)
section.header_distance = dxa(720)
section.footer_distance = dxa(300)

# --- Tiêu đề ---
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p_title, "TÊN BÀI/CHỦ ĐỀ: ", bold=True, size=16)
add_run(p_title, "BÀI 28 – ĐỘNG LƯỢNG (V2, 2 TIẾT)", bold=True, size=16, color=PRIMARY)
doc.add_paragraph()

# --- Bảng thông tin chung ---
info_table = doc.add_table(rows=0, cols=2)
info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
info_table.autofit = False
col_w = [3847, 10267]


def add_info_row(label: str, content_lines: list[tuple], *, label_color=DARK):
    row = info_table.add_row()
    for cell, w in zip(row.cells, col_w):
        cell.width = dxa(w)
        set_cell_borders(cell)
    cell_para(info_table.rows[-1].cells[0], label, bold=True, color=label_color, first=True)
    cell_multiline(info_table.rows[-1].cells[1], content_lines)
    return row


add_info_row("Tuần", [("......... (Giáo viên điền theo phân phối chương trình thực tế)", {"italic": True})])
add_info_row("Lớp", [("10", {})])
add_info_row("Sách (tài liệu)", [
    ("Vật lí 10 – Bộ sách Kết nối tri thức với cuộc sống, NXB Giáo dục Việt Nam.", {}),
    ("Bài 28 – Động lượng (Chương V – Động lượng), trang 110–112.", {}),
])
add_info_row("Mục tiêu", [
    ("1. Kiến thức", {"bold": True}),
    ("– Phát biểu được định nghĩa của động lượng và nêu được ý nghĩa vật lí của đại lượng đó.", {}),
    ("– Phát biểu và viết được công thức liên hệ giữa lực tác dụng lên vật và tốc độ biến thiên "
     "của động lượng (thường được gọi là dạng thứ 2 của định luật II Newton).", {}),
    ("2. Phát triển năng lực", {"bold": True}),
    ("a) Năng lực chung:", {"bold": True, "size": 11}),
    ("– Tự học: chủ động tích cực thực hiện công việc bản thân, đóng góp ý tưởng, đặt câu hỏi, "
     "trả lời yêu cầu.", {}),
    ("– Giao tiếp và hợp tác: khiêm tốn tiếp thu góp ý, nhiệt tình chia sẻ/hỗ trợ thành viên "
     "nhóm.", {}),
    ("b) Năng lực vật lí:", {"bold": True, "size": 11}),
    ("– Nêu được ý nghĩa vật lí và định nghĩa động lượng.", {}),
    ("– Phát biểu được mối liên hệ giữa lực tác dụng lên vật và tốc độ biến thiên của động "
     "lượng.", {}),
    ("– Viết được công thức liên hệ giữa lực tác dụng lên vật và tốc độ biến thiên của động "
     "lượng.", {}),
    ("(Nguồn: nguyên văn Mục tiêu do giáo viên phụ trách cung cấp — xem dac_ta_v2.md mục 1.)",
     {"italic": True, "size": 10, "color": RGBColor(0x59, 0x59, 0x59)}),
])
add_info_row("Tiến trình", [
    ("★ TIẾT 1 — I. ĐỘNG LƯỢNG (TIẾT DẠY MINH HỌA)", {"bold": True, "color": RED}),
    ("Hoạt động 1 – Khởi động: trò chơi AR \"Bảo vệ hành tinh xanh\" (10 phút)", {}),
    ("Hoạt động 2 – Hình thành kiến thức mới: Động lượng (20 phút)", {}),
    ("Hoạt động 3 – Luyện tập (10 phút)", {}),
    ("Hoạt động 4 – Vận dụng, củng cố và dặn dò (5 phút)", {}),
    ("TIẾT 2 — II. XUNG LƯỢNG CỦA LỰC", {"bold": True}),
    ("Hoạt động 1 – Khởi động (5 phút)", {}),
    ("Hoạt động 2 – Hình thành kiến thức mới: Xung lượng của lực (20 phút)", {}),
    ("Hoạt động 3 – Luyện tập (12 phút)", {}),
    ("Hoạt động 4 – Vận dụng, củng cố và dặn dò (8 phút)", {}),
    ("(Mỗi tiết 45 phút, đủ 4 bước CV5512: Khởi động → Hình thành kiến thức mới → Luyện tập → "
     "Vận dụng.)", {"italic": True, "size": 11}),
])
add_info_row("Chuẩn bị của giáo viên", [
    ("– SGK Vật lí 10 (Kết nối tri thức), Kế hoạch bài dạy, Phiếu học tập (PHT) Bài 28 V2.", {}),
    ("– Tiết 1: thiết bị trình chiếu/máy tính bảng chạy trò chơi AR \"Bảo vệ hành tinh xanh\" "
     "(hoặc bản mô phỏng tương tác tương đương nếu lớp không đủ thiết bị mỗi nhóm — GV có thể "
     "cho HS lần lượt lên chơi trên 1 thiết bị chung, chiếu màn hình cho cả lớp theo dõi).", {}),
    ("– Bộ dụng cụ thí nghiệm Hình 28.1: máng trượt (có thể dùng ống nhựa cắt dọc) + 3 viên bi "
     "A, B, C (bi B nặng hơn bi A) — mỗi nhóm 1 bộ nếu đủ điều kiện; nếu không đủ dụng cụ, GV "
     "làm mẫu 1 bộ, các nhóm quan sát trực tiếp và tự ghi số liệu vào PHT.", {}),
    ("– Tiết 1: 1 hình chế (meme) \"xe tăng vs xe đạp\" hoặc tương tự để khai thác NV Giải mã "
     "Meme Vật lý; 1 đoạn clip/ảnh chụp màn hình mô phỏng phát biểu sai lan truyền mạng xã hội "
     "về động lượng để khai thác NV Bóc phốt TikTok/Shorts.", {}),
    ("– Tiết 2: đồ thị lực – thời gian (F không đổi trong khoảng Δt), video/ảnh thủ môn bắt "
     "bóng và túi khí ô tô bung ra.", {}),
    ("– Phiếu quan sát/đánh giá hoạt động nhóm.", {}),
])
add_info_row("Chuẩn bị của học sinh", [
    ("– SGK Vật lí 10, vở ghi, dụng cụ học tập.", {}),
    ("– Ôn lại định luật II Newton (F = m.a) và khái niệm gia tốc đã học ở Bài 15 (cần cho "
     "Tiết 2).", {}),
    ("– Đọc trước nội dung Bài 28 – Động lượng ở nhà.", {}),
])
add_info_row("Tài liệu hỗ trợ giảng dạy, học tập.", [
    ("– Phiếu học tập Bài 28 V2 (bai28_dong_luong_v2_phieu_hoc_tap.docx).", {}),
    ("– Bài tập về nhà Bài 28 V2 (bai28_dong_luong_v2_bai_tap_ve_nha.docx).", {}),
    ("– Slide Giáo viên và Slide Phiếu học tập: soạn ở lượt sau, cần xác nhận riêng trước khi "
     "gọi NotebookLM (chưa thực hiện trong lượt này).", {}),
])
demo_row = info_table.add_row()
for cell, w in zip(demo_row.cells, col_w):
    cell.width = dxa(w)
    set_cell_borders(cell)
cell_para(demo_row.cells[0], "", first=True)
p_demo_label = demo_row.cells[0].paragraphs[0]
add_run(p_demo_label, "*** ", bold=True, color=RED)
add_run(p_demo_label, "TIẾT CHỌN DẠY MINH HỌA:", bold=True, color=RED, highlight="yellow")
cell_multiline(demo_row.cells[1], [
    ("TIẾT 1 — I. Động lượng (toàn bộ 4 hoạt động, 45 phút).", {"bold": True}),
    ("Lí do: mở đầu bằng trò chơi AR gây hào hứng cao, có hoạt động nhóm thực hành thí nghiệm "
     "thật (tư duy bậc cao qua Model Builder), nhiệm vụ đa dạng gần gũi học sinh (Meme, "
     "TikTok/Shorts, Decision Tree), có phương thức đánh giá rõ ràng sau mỗi hoạt động, và Vận "
     "dụng khép vòng trực tiếp với bối cảnh Khởi động — thể hiện rõ sự liên kết chặt chẽ giữa "
     "các hoạt động giảng dạy.", {}),
])

doc.add_paragraph()

# --- TIẾN TRÌNH DẠY HỌC ---
p_ttdh = doc.add_paragraph()
add_run(p_ttdh, "  \tTIẾN TRÌNH DẠY HỌC", bold=True, size=14, color=PRIMARY)
doc.add_paragraph()

process_cols = [1565, 1920, 4697, 5348, 955]
headers = ["Hoạt động", "Mục tiêu", "Chi tiết hoạt động", "Nội dung kiến thức", "Ghi chú"]

proc_table = doc.add_table(rows=1, cols=5)
proc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
proc_table.autofit = False
for cell, w, header in zip(proc_table.rows[0].cells, process_cols, headers):
    cell.width = dxa(w)
    set_cell_borders(cell)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, header, bold=True, size=12)


def add_activity_row(hoat_dong: str, muc_tieu: list[tuple], chi_tiet: list[tuple],
                      noi_dung: list[tuple], ghi_chu: list[tuple]) -> None:
    row = proc_table.add_row()
    for cell, w in zip(row.cells, process_cols):
        cell.width = dxa(w)
        set_cell_borders(cell)
    p0 = row.cells[0].paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p0, hoat_dong, bold=True, size=11)
    cell_multiline(row.cells[1], [(t, {**k, "size": k.get("size", 11)}) for t, k in muc_tieu])
    cell_multiline(row.cells[2], [(t, {**k, "size": k.get("size", 11)}) for t, k in chi_tiet])
    cell_multiline(row.cells[3], [(t, {**k, "size": k.get("size", 11)}) for t, k in noi_dung])
    cell_multiline(row.cells[4], [(t, {**k, "size": k.get("size", 10)}) for t, k in ghi_chu])
    return row


def add_section_marker(title: str) -> None:
    row = proc_table.add_row()
    for cell, w in zip(row.cells, process_cols):
        cell.width = dxa(w)
        set_cell_borders(cell)
    row.cells[0].merge(row.cells[-1])
    p = row.cells[0].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, title, bold=True, size=13, color=RGBColor(0xFF, 0xFF, 0xFF), highlight="darkBlue")


# ============================================================
# TIẾT 1 — I. ĐỘNG LƯỢNG (★ DEMO)
# ============================================================
add_section_marker("★ TIẾT 1 — I. ĐỘNG LƯỢNG (TIẾT DẠY MINH HỌA)")

# HĐ1 — Khởi động: Trò chơi AR "Bảo vệ hành tinh xanh"
add_activity_row(
    "Hoạt động 1\nKhởi động\n(10 phút)",
    [("Tạo hứng thú qua trải nghiệm trò chơi, huy động hiểu biết ban đầu về mối liên hệ giữa "
      "khối lượng, tốc độ và khả năng \"làm lệch hướng/truyền tương tác\" của một vật, kích "
      "thích tò mò dẫn vào khái niệm động lượng.", {})],
    [
        ("Chuyển giao:", {"bold": True}),
        ("GV tổ chức trò chơi AR \"Bảo vệ hành tinh xanh\": học sinh dùng tay điều khiển tên "
         "lửa bắn vào thiên thạch đang lao về phía Trái Đất để làm nó đổi hướng, tránh va "
         "chạm. Loại hình: trò chơi mô phỏng tương tác. Hình thức: cả lớp lần lượt trải nghiệm "
         "(2–3 lượt tiêu biểu, có thể chia đội thi đua); thời gian: 6 phút chơi + 4 phút thảo "
         "luận; tài liệu: thiết bị/màn hình chạy trò chơi AR + PHT mục 1.", {}),
        ("Thực hiện:", {"bold": True}),
        ("HS lần lượt chơi, cả lớp quan sát và ghi nhanh nhận xét vào PHT: với cùng 1 tên lửa, "
         "khi nào thiên thạch bị lệch hướng nhiều hơn — khi tên lửa nặng hơn hay khi tên lửa "
         "bắn nhanh hơn?", {}),
        ("Báo cáo:", {"bold": True}),
        ("Kĩ thuật \"Hỏi nhanh – Đáp nhanh\": GV mời 2–3 HS chia sẻ quan sát và dự đoán của "
         "mình.", {}),
        ("Kết luận:", {"bold": True}),
        ("GV chốt vấn đề: \"Cả khối lượng lẫn tốc độ của tên lửa đều ảnh hưởng đến khả năng làm "
         "lệch hướng thiên thạch. Đại lượng nào đặc trưng cho điều này? Chúng ta cùng khám phá "
         "qua thí nghiệm thực tế trong bài học hôm nay — Động lượng.\"", {}),
    ],
    [("Nêu vấn đề: khả năng truyền tương tác (làm lệch hướng/truyền chuyển động) phụ thuộc cả "
      "khối lượng và tốc độ của vật.", {})],
    [("Đánh giá định tính qua quan sát: mức độ hào hứng, số lượng học sinh xung phong chơi/trả "
      "lời — không chấm điểm, dùng để dẫn dắt.", {})],
)

# HĐ2 — Hình thành kiến thức: Động lượng (Model Builder + Meme Analyzer)
add_activity_row(
    "Hoạt động 2\nHình thành\nkiến thức mới:\nĐộng lượng\n(20 phút)",
    [("Thông qua thí nghiệm trực tiếp, rút ra được các yếu tố ảnh hưởng đến khả năng truyền "
      "chuyển động; phát biểu được định nghĩa động lượng, viết được công thức p = m.v, nêu "
      "được ý nghĩa vật lí và tính chất vectơ của động lượng (đáp ứng yêu cầu cần đạt 1).", {})],
    [
        ("Chuyển giao NV1:", {"bold": True}),
        ("GV giao NV1 \"Thí nghiệm ba viên bi A, B, C\" (PHT mục 2.1.a), yêu cầu HS đối chiếu "
         "ngược lại với trò chơi AR vừa chơi (viên bi ↔ tên lửa). Loại hình: Model Builder (tự "
         "làm thí nghiệm, rút ra công thức). Hình thức: nhóm 4–5 học sinh; thời gian: 8 phút; "
         "tài liệu: bộ dụng cụ thí nghiệm (máng trượt + 3 viên bi) + SGK trang 110–111 + PHT.", {}),
        ("Thực hiện NV1:", {"bold": True}),
        ("Học sinh TỰ TAY làm Thí nghiệm 1 (thả bi A rồi bi B từ cùng độ cao, so sánh quãng "
         "đường bi C lăn) và Thí nghiệm 2 (chỉ thả bi A, đổi độ dốc), ghi số liệu vào PHT, thảo "
         "luận nhóm trả lời 3 câu hỏi rút ra kết luận.", {}),
        ("Chuyển giao NV2:", {"bold": True}),
        ("GV chiếu 1 hình chế (meme) \"xe tăng vs xe đạp\" kèm câu chú thích gây tranh cãi. "
         "Loại hình: Giải mã Meme Vật lý (Meme Analyzer). Hình thức: nhóm đôi; thời gian: 5 "
         "phút; tài liệu: hình chế đã chuẩn bị + PHT mục 2.1.b.", {}),
        ("Thực hiện NV2:", {"bold": True}),
        ("HS phân tích hình chế: chú thích trong meme đúng hay sai về mặt vật lí, vì sao, dựa "
         "trên kết luận thí nghiệm vừa rút ra.", {}),
        ("Báo cáo:", {"bold": True}),
        ("Đại diện 1–2 nhóm trình bày số liệu/kết luận NV1; 1–2 nhóm đôi trình bày phân tích "
         "meme NV2, các nhóm khác nhận xét, bổ sung.", {}),
        ("Kết luận:", {"bold": True}),
        ("GV chốt: khả năng truyền chuyển động phụ thuộc CẢ khối lượng lẫn tốc độ của vật — đây "
         "chính là động lượng; chốt định nghĩa p = m.v, tính vectơ, đơn vị kg.m/s, ý nghĩa vật "
         "lí; yêu cầu HS hoàn thành phần điền khuyết \"Kiến thức trọng tâm\" trong PHT.", {}),
    ],
    [
        ("Động lượng p = m.v là đại lượng vectơ, cùng hướng với vận tốc, đơn vị kg.m/s.", {}),
        ("Động lượng đặc trưng cho khả năng truyền tương tác (truyền chuyển động) giữa các vật "
         "khi va chạm — rút ra trực tiếp từ kết quả thí nghiệm (bi nặng hơn hoặc bi nhanh hơn "
         "đều làm bi C lăn xa hơn), khớp với quan sát ở trò chơi AR mở đầu.", {}),
    ],
    [("Đánh giá qua bảng số liệu thí nghiệm + câu trả lời thảo luận (NV1) và sản phẩm phân tích "
      "meme (NV2) trong PHT — GV quan sát nhóm làm thí nghiệm trực tiếp, quan sát nhanh 2–3 "
      "phiếu đại diện. Tiêu chí: đo/ghi số liệu hợp lý, kết luận đúng dựa trên cả m và v, có "
      "tinh thần hợp tác khi cùng làm thí nghiệm.", {})],
)

# HĐ3 — Luyện tập (Fact Check Influencer)
add_activity_row(
    "Hoạt động 3\nLuyện tập\n(10 phút)",
    [("Vận dụng công thức p = m.v để giải 1 câu hỏi luyện tập theo đúng 1 trong 3 dạng đề thi "
      "tốt nghiệp THPT, đồng thời rèn kĩ năng phản biện thông tin mạng xã hội.", {})],
    [
        ("Chuyển giao:", {"bold": True}),
        ("GV chiếu 1 đoạn mô phỏng phát biểu sai lan truyền trên mạng xã hội về động lượng (PHT "
         "mục 3.1). Loại hình: Bóc phốt TikTok/Shorts (Fact Check Influencer) kết hợp tính "
         "toán. Hình thức: cá nhân; thời gian: 6 phút; tài liệu: PHT mục 3.1.", {}),
        ("Thực hiện:", {"bold": True}),
        ("Học sinh chỉ ra chỗ sai trong phát biểu, giải thích bằng lập luận vật lí, và giải 1 "
         "câu tính toán minh hoạ đi kèm.", {}),
        ("Báo cáo:", {"bold": True}),
        ("Kĩ thuật gọi tên ngẫu nhiên: GV gọi ngẫu nhiên 2 học sinh lên trình bày.", {}),
        ("Kết luận:", {"bold": True}),
        ("GV nhận xét, chốt đáp án và cách trình bày chuẩn (lưu ý đổi đơn vị km/h → m/s).", {}),
    ],
    [("Vận dụng p = m.v để phản biện 1 phát biểu sai phổ biến, khắc sâu vai trò ngang nhau của "
      "khối lượng và tốc độ.", {})],
    [("Chấm nhanh tại lớp, công bố đáp án ngay — đánh giá qua bài làm trong PHT.", {})],
)

# HĐ4 — Vận dụng: Decision Tree, khép vòng với Khởi động
add_activity_row(
    "Hoạt động 4\nVận dụng, củng cố\nvà dặn dò\n(5 phút)",
    [("Vận dụng khái niệm động lượng để ra quyết định trong đúng bối cảnh trò chơi mở đầu — "
      "khép vòng Khởi động↔Vận dụng; củng cố toàn bộ nội dung Tiết 1 và giao nhiệm vụ về nhà.", {})],
    [
        ("Chuyển giao:", {"bold": True}),
        ("GV quay lại đúng bối cảnh trò chơi AR mở đầu, giao NV \"Bản đồ lựa chọn sinh tử\" (PHT "
         "mục 4.1): chỉ có 1 tên lửa duy nhất, chọn phương án khối lượng/vận tốc nào để bảo vệ "
         "Trái Đất hiệu quả nhất. Loại hình: Decision Tree. Hình thức: nhóm đôi; thời gian: 3 "
         "phút; tài liệu: PHT mục 4.1 (sơ đồ cây lựa chọn).", {}),
        ("Thực hiện:", {"bold": True}),
        ("HS thảo luận theo cặp, đi theo sơ đồ cây lựa chọn, chọn nhánh và giải thích lí do dựa "
         "trên p = m.v.", {}),
        ("Báo cáo:", {"bold": True}),
        ("1–2 cặp chia sẻ nhanh lựa chọn trước lớp, biểu quyết nhanh cả lớp xem đa số chọn "
         "nhánh nào.", {}),
        ("Kết luận:", {"bold": True}),
        ("GV chốt kiến thức trọng tâm Tiết 1 (định nghĩa, công thức, ý nghĩa vật lí động "
         "lượng); giao Bài tập về nhà liên quan và dặn dò chuẩn bị Tiết 2 (Xung lượng của lực — "
         "ôn lại định luật II Newton F = m.a).", {}),
    ],
    [("Vận dụng p = m.v để ra quyết định trong tình huống thực tế/giả lập; tổng kết Tiết 1.", {})],
    [("Câu hỏi vận dụng miệng + biểu quyết nhanh, không chấm điểm. Dặn dò: ôn lại F = m.a, "
      "chuẩn bị Tiết 2.", {})],
)

# ============================================================
# TIẾT 2 — II. XUNG LƯỢNG CỦA LỰC
# ============================================================
add_section_marker("TIẾT 2 — II. XUNG LƯỢNG CỦA LỰC")

# HĐ1 — Khởi động
add_activity_row(
    "Hoạt động 1\nKhởi động\n(5 phút)",
    [("Ôn nhanh động lượng, huy động hiểu biết ban đầu để dẫn vào khái niệm xung lượng của "
      "lực.", {})],
    [
        ("Chuyển giao:", {"bold": True}),
        ("GV chiếu tình huống mở bài SGK trang 110 (thủ môn bắt bóng sút phạt 11 m) và ảnh túi "
         "khí ô tô bung ra. Hình thức: cá nhân; thời gian: 3 phút; tài liệu: PHT mục 2.2 (mở "
         "đầu).", {}),
        ("Thực hiện:", {"bold": True}),
        ("HS quan sát, dự đoán nhanh: thủ môn khó bắt bóng hơn khi bóng bay tới nhanh hay chậm? "
         "Vì sao — liên hệ lại động lượng đã học Tiết 1.", {}),
        ("Báo cáo:", {"bold": True}),
        ("GV mời 2–3 HS nêu dự đoán.", {}),
        ("Kết luận:", {"bold": True}),
        ("GV dẫn dắt: \"Vậy khi động lượng biến đổi, đại lượng nào liên quan đến lực và thời "
         "gian tác dụng? Hôm nay ta tìm hiểu Xung lượng của lực.\"", {}),
    ],
    [("Ôn động lượng p = m.v, dẫn vào vấn đề lực tác dụng trong thời gian ngắn.", {})],
    [("GV quan sát câu trả lời dự đoán, không chấm điểm.", {})],
)

# HĐ2 — Hình thành kiến thức: Xung lượng của lực (Algorithmic Ordering + Bug Buster)
add_activity_row(
    "Hoạt động 2\nHình thành\nkiến thức mới:\nXung lượng\ncủa lực\n(20 phút)",
    [("Nêu được định nghĩa xung lượng của lực, thiết lập được công thức liên hệ giữa xung "
      "lượng của lực và độ biến thiên động lượng, viết được dạng tổng quát của định luật II "
      "Newton (đáp ứng yêu cầu cần đạt 2); kế thừa trực tiếp p = m.v đã học ở Tiết 1.", {})],
    [
        ("Chuyển giao:", {"bold": True}),
        ("GV giao NV1 \"Sắp xếp tiến trình suy luận F.Δt = Δp\" (Algorithmic Ordering — HS sắp "
         "xếp lại đúng thứ tự các bước biến đổi từ a = Δv/Δt đến F.Δt = Δp) và NV2 \"Bug Buster "
         "– tìm lỗi trong lời giải mẫu\" (PHT mục 2.2.a). Hình thức: nhóm 4 học sinh (kĩ thuật "
         "khăn trải bàn); thời gian: 10 phút; tài liệu: công thức định luật II Newton đã học ở "
         "Bài 15 + SGK trang 111–112 + PHT.", {}),
        ("Thực hiện:", {"bold": True}),
        ("Nhóm thảo luận sắp xếp đúng thứ tự 4 bước suy luận (đã bị xáo trộn trong PHT) từ "
         "a = (v2-v1)/Δt → F = m.a → F.Δt = m.v2 - m.v1 → F.Δt = Δp; sau đó đọc 1 lời giải mẫu "
         "cố ý có lỗi sai (Bug Buster) và chỉ ra chỗ sai, sửa lại cho đúng.", {}),
        ("Báo cáo:", {"bold": True}),
        ("Đại diện nhóm trình bày thứ tự sắp xếp đúng trên bảng; 1 nhóm khác trình bày lỗi sai "
         "tìm được trong bài Bug Buster, cả lớp nhận xét từng bước biến đổi.", {}),
        ("Kết luận:", {"bold": True}),
        ("GV chốt: xung lượng của lực F.Δt (đơn vị N.s); công thức liên hệ F.Δt = Δp = p2 – p1; "
         "dạng tổng quát định luật II Newton F = Δp/Δt. Liên hệ ngay thực tế thủ môn bắt bóng, "
         "túi khí ô tô.", {}),
    ],
    [
        ("Xung lượng của lực F.Δt (N.s).", {}),
        ("Liên hệ xung lượng của lực và độ biến thiên động lượng: F.Δt = Δp = p2 – p1.", {}),
        ("Dạng tổng quát của định luật II Newton: F = Δp/Δt.", {}),
    ],
    [("Đánh giá qua sản phẩm PHT (sắp xếp đúng thứ tự 4 bước + phát hiện đúng lỗi sai) và câu "
      "trả lời khi báo cáo — chấm chéo giữa các nhóm ngay tại lớp.", {})],
)

# HĐ3 — Luyện tập
add_activity_row(
    "Hoạt động 3\nLuyện tập\n(12 phút)",
    [("Vận dụng công thức p = m.v và F = Δp/Δt để giải câu hỏi Đúng/Sai có biện giải, gộp cả 2 "
      "tiết, theo đúng 1 trong 3 dạng đề thi tốt nghiệp THPT.", {})],
    [
        ("Chuyển giao:", {"bold": True}),
        ("GV giao bài luyện tập Đúng/Sai (PHT mục 3.2). Hình thức: cá nhân; thời gian: 8 phút; "
         "tài liệu: PHT mục 3.2.", {}),
        ("Thực hiện:", {"bold": True}),
        ("Học sinh làm bài cá nhân vào PHT, biện giải rõ Đúng/Sai từng ý.", {}),
        ("Báo cáo:", {"bold": True}),
        ("Kĩ thuật gọi tên ngẫu nhiên: GV gọi ngẫu nhiên 2 học sinh lên trình bày lời giải.", {}),
        ("Kết luận:", {"bold": True}),
        ("GV nhận xét, chốt đáp án ngay tại lớp và cách trình bày chuẩn.", {}),
    ],
    [("Vận dụng p = m.v và F = Δp/Δt trong 1 tình huống tổng hợp cả hai đơn vị kiến thức.", {})],
    [("Chấm nhanh tại lớp, công bố đáp án ngay — đánh giá qua bài làm trong PHT.", {})],
)

# HĐ4 — Vận dụng: Engineering Debugger + Infographic
add_activity_row(
    "Hoạt động 4\nVận dụng, củng cố\nvà dặn dò\n(8 phút)",
    [("Liên hệ thực tiễn phù hợp (an toàn giao thông/thể thao) không gượng ép; củng cố toàn bài "
      "và giao nhiệm vụ về nhà.", {})],
    [
        ("Chuyển giao:", {"bold": True}),
        ("GV giao NV \"Gỡ lỗi thiết kế mũ bảo hiểm lót thép\" (Engineering Debugger, PHT mục "
         "4.2.a) và đoạn thông tin về túi khí/dây an toàn (Khai thác Infographic, PHT mục "
         "4.2.b). Hình thức: nhóm đôi; thời gian: 5 phút; tài liệu: PHT mục 4.1–4.2 + công thức "
         "F = Δp/Δt.", {}),
        ("Thực hiện:", {"bold": True}),
        ("Học sinh chỉ ra điểm bất hợp lí trong đề xuất kĩ thuật sai (lót thép cứng), đề xuất "
         "cách khắc phục đúng nguyên lí; đọc thông tin túi khí/dây an toàn, thảo luận nhanh "
         "theo cặp.", {}),
        ("Báo cáo:", {"bold": True}),
        ("1–2 cặp chia sẻ nhanh câu trả lời trước lớp (câu hỏi vận dụng miệng, không chấm "
         "điểm).", {}),
        ("Kết luận:", {"bold": True}),
        ("GV chốt kiến thức trọng tâm toàn bài bằng sơ đồ tư duy ngắn gọn (động lượng – xung "
         "lượng – liên hệ – định luật II Newton dạng tổng quát); giao Bài tập về nhà (18 trắc "
         "nghiệm + 4 Đúng/Sai + 6 trả lời ngắn) và phần \"Mở rộng\" trong PHT.", {}),
    ],
    [("Vận dụng liên hệ xung lượng – độ biến thiên động lượng để phân tích/gỡ lỗi thiết kế kĩ "
      "thuật và nguyên lí an toàn giao thông/thể thao; tổng kết toàn bài.", {})],
    [("Câu hỏi vận dụng miệng + sản phẩm gỡ lỗi, không chấm điểm. Dặn dò: hoàn thành Bài tập về "
      "nhà; chuẩn bị Bài 29 – Định luật bảo toàn động lượng.", {})],
)

doc.add_paragraph()

# --- PHẢN HỒI – ĐIỀU CHỈNH ---
p_ph = doc.add_paragraph()
add_run(p_ph, "PHẢN HỒI – ĐIỀU CHỈNH", bold=True, size=14, color=PRIMARY)

for label in ("Ưu điểm", "Hạn chế", "Hướng điều chỉnh"):
    p = doc.add_paragraph()
    add_run(p, f"{label}: ", bold=True)
    add_run(p, "..." * 30, color=RGBColor(0xBF, 0xBF, 0xBF))

OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
doc.save(str(OUT_PATH))
print("OK - da xuat KHBD V2 (2 tiet)")
