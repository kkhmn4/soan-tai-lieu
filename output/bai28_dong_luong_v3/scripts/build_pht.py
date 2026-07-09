"""Soạn Phiếu học tập (PHT) V3 cho Bài 28 - Động lượng: CHỈ 1 TIẾT, chỉ mục
"I. Động lượng" — KHÔNG có phần Xung lượng của lực (khác V2). Cấu trúc còn 4
mục lớn nhưng không đánh số con 2.x/3.x/4.x vì chỉ có 1 đơn vị kiến thức:

    1. KHỞI ĐỘNG — Trò chơi AR "Bảo vệ hành tinh xanh"
    2. HÌNH THÀNH KIẾN THỨC MỚI — Thí nghiệm 3 viên bi (Model Builder) + Giải
       mã Meme Vật lý (Meme Analyzer)
    3. LUYỆN TẬP — Bóc phốt TikTok/Shorts (Fact Check Influencer)
    4. VẬN DỤNG - MỞ RỘNG — Bản đồ lựa chọn sinh tử (Decision Tree, khép
       vòng với mục 1) + đọc mở rộng

Nội dung khớp 1-1 với KHBD V3 (build_khbd.py). Không nhúng ảnh minh họa
(ưu tiên nội dung chữ trước). Tài liệu one-off, dựng bằng python-docx, KHÔNG
dùng `gems/docx_export/pht_exporter.py`, KHÔNG sửa file nào trong `gems/`.

Chạy: python "output/bai28_dong_luong_v3/scripts/build_pht.py"
"""
from __future__ import annotations

import sys
from pathlib import Path

sys.stdout.reconfigure(encoding="utf-8")

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

REPO_ROOT = Path(__file__).resolve().parent.parent.parent.parent
OUT_DIR = REPO_ROOT / "output" / "bai28_dong_luong_v3" / "ready"
OUT_PATH = OUT_DIR / "bai28_dong_luong_v3_phieu_hoc_tap.docx"

FONT = "Times New Roman"
PRIMARY = RGBColor(0x1F, 0x4E, 0x79)
DARK = RGBColor(0x00, 0x00, 0x00)
GRAY = RGBColor(0x59, 0x59, 0x59)


# ============================================================
# HẠ TẦNG DỰNG DOCX
# ============================================================
def set_default_font(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(13)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), FONT)


def add_run(p, text: str, *, bold=False, italic=False, size=13, color=DARK) -> None:
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.name = FONT
    r.font.size = Pt(size)
    r.font.color.rgb = color


def add_para(doc, text: str = "", *, bold=False, italic=False, size=13, color=DARK,
             space_before=0, space_after=6, align=None, indent_cm=None) -> object:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if align is not None:
        p.alignment = align
    if indent_cm is not None:
        p.paragraph_format.left_indent = Cm(indent_cm)
    if text:
        add_run(p, text, bold=bold, italic=italic, size=size, color=color)
    return p


def section_heading(doc, number: str, title: str) -> None:
    p = add_para(doc, space_before=14, space_after=6)
    add_run(p, f"{number}. {title}", bold=True, size=15, color=PRIMARY)


def sub_item_label(doc, letter: str, label: str, *, indent_cm=0.5) -> object:
    p = add_para(doc, space_before=4, space_after=2, indent_cm=indent_cm)
    add_run(p, f"{letter}. {label}", bold=True, size=13)
    return p


def body(doc, text: str, *, indent_cm=0.9, size=13, italic=False, bold=False, space_after=4):
    return add_para(doc, text, size=size, italic=italic, bold=bold, indent_cm=indent_cm,
                     space_after=space_after)


def instructions(doc, text: str, *, indent_cm=0.9):
    return body(doc, f"Hướng dẫn thực hiện: {text}", italic=True, size=12, indent_cm=indent_cm)


def bullet(doc, text: str, *, indent_cm=1.3):
    p = add_para(doc, space_after=3, indent_cm=indent_cm)
    add_run(p, "– ", bold=True)
    add_run(p, text)
    return p


def _set_bottom_border(p, *, sz=6, color="000000") -> None:
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(sz))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)
    pPr.append(pbdr)


def ruled_lines(doc, n: int, *, indent_cm=0.9) -> None:
    for _ in range(n):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(indent_cm)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(11)
        p.paragraph_format.line_spacing = 1.0
        r = p.add_run(" ")
        r.font.name = FONT
        r.font.size = Pt(13)
        _set_bottom_border(p)


def separator(doc) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    _set_bottom_border(p, sz=4, color="BFBFBF")


def _set_cell_borders(cell) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "BFBFBF")
        borders.append(el)
    tcPr.append(borders)


def data_table(doc, headers: list[str], rows: list[list[str]], *, indent_cm=1.1) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.autofit = True
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        _set_cell_borders(hdr[i])
        add_run(hdr[i].paragraphs[0], h, bold=True, size=12)
    for r_idx, row in enumerate(rows, start=1):
        cells = table.rows[r_idx].cells
        for c_idx, val in enumerate(row):
            _set_cell_borders(cells[c_idx])
            add_run(cells[c_idx].paragraphs[0], val or "…………………………", size=12)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent_cm)
    p.paragraph_format.space_after = Pt(8)


# ============================================================
# DỰNG TÀI LIỆU
# ============================================================
doc = Document()
set_default_font(doc)

section = doc.sections[0]
section.top_margin = Cm(2.0)
section.bottom_margin = Cm(2.0)
section.left_margin = Cm(2.0)
section.right_margin = Cm(2.0)

# --- Header ---
add_para(doc, "TRƯỜNG: ................................................",
         italic=True, size=12, color=GRAY, space_after=2)

p_title = add_para(doc, space_after=2, align=WD_ALIGN_PARAGRAPH.CENTER)
add_run(p_title, "PHIẾU HỌC TẬP", bold=True, size=20, color=PRIMARY)
add_para(doc, "BÀI 28 – ĐỘNG LƯỢNG (V3 — 1 TIẾT)", bold=True, size=14, color=DARK,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)

info_table = doc.add_table(rows=2, cols=2)
info_table.autofit = True
cells = info_table.rows[0].cells
add_run(cells[0].paragraphs[0], "Họ và tên học sinh: ..........................................", size=12)
add_run(cells[1].paragraphs[0], "Lớp: .............. Nhóm: ..............", size=12)
cells2 = info_table.rows[1].cells
add_run(cells2[0].paragraphs[0], "Môn học: Vật lí 10 (CTGDPT 2018)", size=12)
add_run(cells2[1].paragraphs[0], "Tiết học: ......... Ngày: ......./......./.......", size=12)
for row in info_table.rows:
    for cell in row.cells:
        for p in cell.paragraphs:
            p.paragraph_format.space_after = Pt(3)
doc.add_paragraph()
separator(doc)

# ============================================================
# 1. KHỞI ĐỘNG
# ============================================================
section_heading(doc, "1", "KHỞI ĐỘNG — Trò chơi AR \"Bảo vệ hành tinh xanh\"")
instructions(doc, "cả lớp lần lượt trải nghiệm (2–3 lượt tiêu biểu), 6 phút chơi + 4 phút thảo luận, "
                   "dùng thiết bị/màn hình chạy trò chơi AR do giáo viên chuẩn bị.", indent_cm=0.5)

sub_item_label(doc, "a", "Luật chơi")
body(doc, "Em dùng tay điều khiển một tên lửa, bắn vào thiên thạch đang lao về phía Trái Đất để làm "
          "nó ĐỔI HƯỚNG, tránh va chạm. Mỗi lượt chơi, độ lớn khối lượng và tốc độ của tên lửa có thể "
          "thay đổi khác nhau.")

sub_item_label(doc, "b", "Quan sát")
body(doc, "Trong các lượt chơi mà em quan sát được, với CÙNG một thiên thạch, khi nào tên lửa làm "
          "thiên thạch lệch hướng NHIỀU hơn: khi tên lửa nặng hơn, hay khi tên lửa bắn nhanh hơn, hay "
          "cả hai đều có ảnh hưởng? Ghi lại ví dụ cụ thể em quan sát được.")
ruled_lines(doc, 3)

sub_item_label(doc, "c", "Dự đoán")
body(doc, "Từ trò chơi trên, đại lượng nào đặc trưng cho \"khả năng làm lệch hướng/truyền tương tác\" "
          "của tên lửa lên thiên thạch, và đại lượng đó phụ thuộc vào những yếu tố nào của tên lửa?")
ruled_lines(doc, 2)
separator(doc)

# ============================================================
# 2. HÌNH THÀNH KIẾN THỨC MỚI
# ============================================================
section_heading(doc, "2", "HÌNH THÀNH KIẾN THỨC MỚI — Động lượng")

sub_item_label(doc, "a", "Nhiệm vụ 1 — Thí nghiệm ba viên bi A, B, C (Model Builder)")
instructions(doc, "hoạt động nhóm 4–5 học sinh, 10 phút, dùng bộ dụng cụ thí nghiệm (máng trượt + 3 "
                   "viên bi A, B, C — bi B nặng hơn bi A) đã chuẩn bị sẵn, đọc mục I. Động lượng SGK "
                   "trang 110–111. Hãy đối chiếu lại với trò chơi AR ở mục 1: viên bi ở đây tương tự "
                   "vai trò tên lửa trong trò chơi.")
body(doc, "Bố trí thí nghiệm (Hình 28.1 SGK): đặt máng trượt nghiêng, viên bi C đặt ngay dưới chân "
          "máng trượt, đứng yên chờ va chạm.", indent_cm=1.3)

body(doc, "Thí nghiệm 1: Lần lượt thả bi A rồi bi B từ CÙNG một độ cao trên máng trượt (cùng tốc độ "
          "khi tới chân dốc). Quan sát và đo quãng đường bi C dịch chuyển sau mỗi lần va chạm.",
     indent_cm=1.3)
data_table(doc, ["Lần thả", "Quãng đường bi C dịch chuyển"], [["Bi A", ""], ["Bi B", ""]])

body(doc, "Thí nghiệm 2: Chỉ thả bi A, nhưng tăng độ dốc của máng trượt. Quan sát và đo quãng đường "
          "bi C dịch chuyển sau mỗi lần va chạm.", indent_cm=1.3)
data_table(doc, ["Độ dốc", "Quãng đường bi C dịch chuyển"], [["Thấp", ""], ["Cao", ""]])

body(doc, "Thảo luận nhóm:", bold=True, indent_cm=1.3, space_after=2)
for label, text, n_lines in [
    ("a)", "Trong thí nghiệm 1, vận tốc của bi A và bi B khi đến chân dốc có giống nhau không? Viên "
           "bi nào làm bi C lăn xa hơn? Tại sao?", 2),
    ("b)", "Trong thí nghiệm 2, ứng với độ dốc nào thì bi A có vận tốc lớn hơn khi va chạm với bi C? "
           "Ở trường hợp nào bi C lăn xa hơn? Tại sao?", 2),
    ("c)", "Từ hai thí nghiệm trên, hãy rút ra kết luận: khả năng truyền chuyển động (sau này gọi là "
           "động lượng) của một vật phụ thuộc vào những yếu tố nào?", 2),
]:
    body(doc, f"{label} {text}", indent_cm=1.3, space_after=2)
    ruled_lines(doc, n_lines, indent_cm=1.3)

sub_item_label(doc, "b", "Nhiệm vụ 2 — Giải mã Meme Vật lý (Meme Analyzer)")
instructions(doc, "hoạt động nhóm đôi, 6 phút, quan sát hình chế do giáo viên chiếu, dựa vào kết luận "
                   "thí nghiệm vừa rút ra.", indent_cm=1.1)
body(doc, "Hình chế: một chiếc xe tăng đang chạy CHẬM và một chiếc xe đạp đang chạy RẤT NHANH, với "
          "chú thích đùa: \"Xe tăng luôn luôn có động lượng lớn hơn xe đạp — vì nó nặng hơn nhiều "
          "mà!\"", indent_cm=1.3, italic=True)
for label, text, n_lines in [
    ("a)", "Chú thích trong hình chế đúng hay sai về mặt vật lí? Vì sao?", 2),
    ("b)", "Hãy đặt ra 1 bộ số liệu cụ thể (khối lượng, tốc độ) cho xe tăng và xe đạp trong hình để "
           "chứng minh rõ ý kiến của nhóm em.", 2),
]:
    body(doc, f"{label} {text}", indent_cm=1.3, space_after=2)
    ruled_lines(doc, n_lines, indent_cm=1.3)

sub_item_label(doc, "c", "Kiến thức trọng tâm")
for text in [
    "Định nghĩa: động lượng của vật khối lượng m chuyển động với vận tốc v là p = (1)….........",
    "Đặc điểm: động lượng là đại lượng (2)…........., luôn (3)…......... với vectơ vận tốc của vật.",
    "Đơn vị: (4)…...........",
    "Ý nghĩa vật lí: động lượng đặc trưng cho (5)…........................................ giữa các "
    "vật khi chúng tương tác (va chạm) với nhau.",
]:
    bullet(doc, text)
body(doc, "Gợi ý từ khóa: p = m.v (dạng vectơ) · vectơ · cùng phương, cùng chiều · kg.m/s · khả năng "
          "truyền tương tác/truyền chuyển động.", italic=True, size=11, indent_cm=1.3)
separator(doc)

# ============================================================
# 3. LUYỆN TẬP
# ============================================================
section_heading(doc, "3", "LUYỆN TẬP — Bóc phốt TikTok/Shorts")
instructions(doc, "hoạt động cá nhân, 5 phút, đọc đoạn mô phỏng phát biểu sai lan truyền trên mạng "
                   "xã hội.", indent_cm=0.5)
body(doc, "Một đoạn video ngắn trên mạng xã hội có nội dung: \"Xe máy luôn luôn khó dừng lại hơn xe "
          "đạp vì xe máy nặng hơn — động lượng chỉ phụ thuộc khối lượng thôi!\"", indent_cm=1.1,
     italic=True)
for label, text, n_lines in [
    ("a)", "Phát biểu trên đúng hay sai? Chỉ ra chỗ sai (nếu có) và giải thích bằng kiến thức động "
           "lượng vừa học.", 2),
    ("b)", "Tính toán minh hoạ: Một xe máy khối lượng 120 kg (kể cả người lái) đang chuyển động với "
           "tốc độ 54 km/h. Tính độ lớn động lượng của xe máy (làm tròn hàng đơn vị, đơn vị kg.m/s).", 3),
]:
    body(doc, f"{label} {text}", indent_cm=1.1, space_after=2)
    ruled_lines(doc, n_lines, indent_cm=1.1)
separator(doc)

# ============================================================
# 4. VẬN DỤNG - MỞ RỘNG
# ============================================================
section_heading(doc, "4", "VẬN DỤNG - MỞ RỘNG — Bản đồ lựa chọn sinh tử")
instructions(doc, "hoạt động nhóm đôi, 3 phút, quay lại đúng bối cảnh trò chơi AR ở mục 1.",
             indent_cm=0.5)
body(doc, "Tình huống: Trạm quan sát vũ trụ chỉ còn ĐÚNG 1 TÊN LỬA để bắn chặn thiên thạch trước khi "
          "nó va chạm Trái Đất. Em được chọn 1 trong 2 phương án chế tạo tên lửa:", indent_cm=1.1)
body(doc, "Phương án A: tên lửa khối lượng RẤT LỚN, tốc độ bắn CHẬM.\n"
          "Phương án B: tên lửa khối lượng NHỎ, tốc độ bắn RẤT NHANH.", indent_cm=1.1)
for label, text, n_lines in [
    ("a)", "Dựa vào công thức p = m.v, phương án nào có khả năng làm thiên thạch lệch hướng nhiều "
           "hơn nếu 2 phương án có động lượng bằng nhau? (Gợi ý: xét thêm các yếu tố thực tế khác "
           "ngoài động lượng, ví dụ độ chính xác khi bắn tốc độ rất nhanh.)", 2),
    ("b)", "Nhóm em chọn phương án nào? Giải thích lí do dựa trên kiến thức động lượng vừa học.", 2),
]:
    body(doc, f"{label} {text}", indent_cm=1.1, space_after=2)
    ruled_lines(doc, n_lines, indent_cm=1.1)

body(doc, "Đọc mở rộng: Khái niệm động lượng không chỉ dùng để so sánh \"vật nào khó dừng lại hơn\" "
          "mà còn là nền tảng của nguyên lí phóng tên lửa: tên lửa đẩy khí nóng phụt ra phía sau với "
          "tốc độ rất lớn để tự thân nó thu được động lượng theo chiều ngược lại (bay lên phía "
          "trước). Đây chính là ý tưởng mở đầu cho định luật bảo toàn động lượng mà em sẽ học ở bài "
          "tiếp theo. (Nguồn tham khảo: SGK Vật lí 10 – Kết nối tri thức với cuộc sống, mục \"Em có "
          "thể\", trang 112.)", indent_cm=1.1, italic=True)

doc.add_paragraph()
p_end = add_para(doc, align=WD_ALIGN_PARAGRAPH.RIGHT, space_before=10)
add_run(p_end, "Ngày ........ tháng ........ năm 202...\n", italic=True, size=12)
add_run(p_end, "Học sinh ký tên\n\n\n..........................................", bold=True, size=12)

OUT_DIR.mkdir(parents=True, exist_ok=True)
doc.save(str(OUT_PATH))
print("OK - da xuat PHT V3 (1 tiet, chi Dong luong)")
