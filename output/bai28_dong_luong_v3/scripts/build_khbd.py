"""Soạn Kế hoạch bài dạy (KHBD) V3 cho Bài 28 - Động lượng: CHỈ 1 TIẾT, chỉ
dạy phần "I. Động lượng" (KHÔNG có "II. Xung lượng của lực" — khác V2 vốn
chia 2 tiết). Theo đúng mẫu "KẾ HOẠCH GIẢNG DẠY.docx" (mẫu Trường THCS và
THPT Đinh Thiện Lý).

Kế thừa thiết kế Tiết 1 của V2 (đã duyệt làm tiết demo): Khởi động = trò chơi
AR "Bảo vệ hành tinh xanh", Hình thành kiến thức = thí nghiệm 3 viên bi +
Giải mã Meme Vật lý, Luyện tập = Bóc phốt TikTok/Shorts, Vận dụng = Bản đồ
lựa chọn sinh tử (khép vòng với Khởi động). Xem `dac_ta_v3.md`.

Tài liệu one-off, không dùng `khbd_exporter.py`, không sửa file nào trong
`gems/`. Lượt này KHÔNG nhúng ảnh minh họa.

Chạy: python "output/bai28_dong_luong_v3/scripts/build_khbd.py"
"""
from __future__ import annotations

import sys
from pathlib import Path

sys.stdout.reconfigure(encoding="utf-8")

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Emu, Pt, RGBColor

REPO_ROOT = Path(__file__).resolve().parent.parent.parent.parent
OUT_PATH = REPO_ROOT / "output" / "bai28_dong_luong_v3" / "ready" / "bai28_dong_luong_v3_ke_hoach_bai_day.docx"

FONT = "Times New Roman"
PRIMARY = RGBColor(0x1F, 0x4E, 0x79)
RED = RGBColor(0xC0, 0x00, 0x00)
DARK = RGBColor(0x00, 0x00, 0x00)


def dxa(value: int) -> Emu:
    return Emu(value * 635)


def set_default_font(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(12)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), FONT)


def set_cell_borders(cell) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    borders = tcPr.makeelement(qn("w:tcBorders"), {})
    for edge in ("top", "left", "bottom", "right"):
        el = borders.makeelement(qn(f"w:{edge}"), {})
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        borders.append(el)
    tcPr.append(borders)


def add_run(p, text: str, *, bold=False, italic=False, size=12, color=DARK, highlight=None) -> None:
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.name = FONT
    r.font.size = Pt(size)
    r.font.color.rgb = color
    if highlight:
        rpr = r._element.get_or_add_rPr()
        hl = rpr.makeelement(qn("w:highlight"), {})
        hl.set(qn("w:val"), highlight)
        rpr.append(hl)


def cell_para(cell, text: str = "", *, bold=False, italic=False, size=12, color=DARK,
              space_after=2, first=False):
    p = cell.paragraphs[0] if first and not cell.paragraphs[0].runs else cell.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    if text:
        add_run(p, text, bold=bold, italic=italic, size=size, color=color)
    return p


def cell_multiline(cell, lines: list[tuple], *, first_clear=True):
    if first_clear:
        cell.paragraphs[0].text = ""
    for i, (text, kwargs) in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(kwargs.pop("space_after", 3))
        p.paragraph_format.space_before = Pt(0)
        add_run(p, text, **kwargs)


# ============================================================
# DỰNG TÀI LIỆU
# ============================================================
doc = Document()
set_default_font(doc)

section = doc.sections[0]
section.orientation = WD_ORIENT.LANDSCAPE
section.page_width = dxa(15840)
section.page_height = dxa(12240)
section.top_margin = dxa(720)
section.bottom_margin = dxa(720)
section.left_margin = dxa(720)
section.right_margin = dxa(720)
section.header_distance = dxa(720)
section.footer_distance = dxa(300)

# --- Tiêu đề ---
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p_title, "TÊN BÀI/CHỦ ĐỀ: ", bold=True, size=16)
add_run(p_title, "BÀI 28 – ĐỘNG LƯỢNG (V3, RIÊNG PHẦN I. ĐỘNG LƯỢNG — 1 TIẾT)", bold=True, size=16,
        color=PRIMARY)
doc.add_paragraph()

# --- Bảng thông tin chung ---
info_table = doc.add_table(rows=0, cols=2)
info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
info_table.autofit = False
col_w = [3847, 10267]


def add_info_row(label: str, content_lines: list[tuple], *, label_color=DARK):
    row = info_table.add_row()
    for cell, w in zip(row.cells, col_w):
        cell.width = dxa(w)
        set_cell_borders(cell)
    cell_para(info_table.rows[-1].cells[0], label, bold=True, color=label_color, first=True)
    cell_multiline(info_table.rows[-1].cells[1], content_lines)
    return row


add_info_row("Tuần", [("......... (Giáo viên điền theo phân phối chương trình thực tế)", {"italic": True})])
add_info_row("Lớp", [("10", {})])
add_info_row("Sách (tài liệu)", [
    ("Vật lí 10 – Bộ sách Kết nối tri thức với cuộc sống, NXB Giáo dục Việt Nam.", {}),
    ("Bài 28 – Động lượng (Chương V – Động lượng), mục I. Động lượng, trang 110–111.", {}),
    ("Lưu ý: tiết học này CHỈ dạy mục I. Động lượng, KHÔNG dạy mục II. Xung lượng của lực (soạn "
     "riêng ở buổi khác nếu cần).", {"italic": True, "size": 10, "color": RGBColor(0x59, 0x59, 0x59)}),
])
add_info_row("Mục tiêu", [
    ("1. Kiến thức", {"bold": True}),
    ("– Phát biểu được định nghĩa của động lượng và nêu được ý nghĩa vật lí của đại lượng đó.", {}),
    ("2. Phát triển năng lực", {"bold": True}),
    ("a) Năng lực chung:", {"bold": True, "size": 11}),
    ("– Tự học: chủ động tích cực thực hiện công việc bản thân, đóng góp ý tưởng, đặt câu hỏi, "
     "trả lời yêu cầu.", {}),
    ("– Giao tiếp và hợp tác: khiêm tốn tiếp thu góp ý, nhiệt tình chia sẻ/hỗ trợ thành viên "
     "nhóm.", {}),
    ("b) Năng lực vật lí:", {"bold": True, "size": 11}),
    ("– Nêu được ý nghĩa vật lí và định nghĩa động lượng.", {}),
    ("– Xác định được động lượng của một vật qua công thức p = m.v trong các tình huống cụ "
     "thể.", {}),
    ("(Nguồn: thu gọn từ Mục tiêu do giáo viên phụ trách cung cấp, chỉ giữ phần liên quan Động "
     "lượng — xem dac_ta_v3.md mục 1.)",
     {"italic": True, "size": 10, "color": RGBColor(0x59, 0x59, 0x59)}),
])
add_info_row("Tiến trình", [
    ("Hoạt động 1 – Khởi động: trò chơi AR \"Bảo vệ hành tinh xanh\" (10 phút)", {"bold": True}),
    ("Hoạt động 2 – Hình thành kiến thức mới: Động lượng (22 phút)", {"bold": True}),
    ("Hoạt động 3 – Luyện tập (8 phút)", {"bold": True}),
    ("Hoạt động 4 – Vận dụng, củng cố và dặn dò (5 phút)", {"bold": True}),
    ("(Tổng thời lượng: 1 tiết – 45 phút, trọn vẹn yêu cầu cần đạt \"I. Động lượng\".)",
     {"italic": True, "size": 11}),
])
add_info_row("Chuẩn bị của giáo viên", [
    ("– SGK Vật lí 10 (Kết nối tri thức), Kế hoạch bài dạy, Phiếu học tập (PHT) Bài 28 V3.", {}),
    ("– Thiết bị/máy tính bảng chạy trò chơi AR \"Bảo vệ hành tinh xanh\" (hoặc bản mô phỏng "
     "tương tác tương đương nếu lớp không đủ thiết bị mỗi nhóm — GV cho HS lần lượt lên chơi "
     "trên 1 thiết bị chung, chiếu màn hình cho cả lớp theo dõi).", {}),
    ("– Bộ dụng cụ thí nghiệm Hình 28.1: máng trượt (có thể dùng ống nhựa cắt dọc) + 3 viên bi "
     "A, B, C (bi B nặng hơn bi A) — mỗi nhóm 1 bộ nếu đủ điều kiện; nếu không đủ dụng cụ, GV "
     "làm mẫu 1 bộ, các nhóm quan sát trực tiếp và tự ghi số liệu vào PHT.", {}),
    ("– 1 hình chế (meme) \"xe tăng vs xe đạp\" để khai thác NV Giải mã Meme Vật lý; 1 đoạn "
     "clip/ảnh chụp màn hình mô phỏng phát biểu sai lan truyền mạng xã hội về động lượng để "
     "khai thác NV Bóc phốt TikTok/Shorts.", {}),
    ("– Phiếu quan sát/đánh giá hoạt động nhóm.", {}),
])
add_info_row("Chuẩn bị của học sinh", [
    ("– SGK Vật lí 10, vở ghi, dụng cụ học tập.", {}),
    ("– Đọc trước mục I. Động lượng, Bài 28, SGK trang 110–111.", {}),
])
add_info_row("Tài liệu hỗ trợ giảng dạy, học tập.", [
    ("– Phiếu học tập Bài 28 V3 (bai28_dong_luong_v3_phieu_hoc_tap.docx).", {}),
    ("– Bài tập về nhà Bài 28 V3 (bai28_dong_luong_v3_bai_tap_ve_nha.docx).", {}),
    ("– Slide Giáo viên và Slide Phiếu học tập: soạn ở lượt sau, cần xác nhận riêng trước khi "
     "gọi NotebookLM (chưa thực hiện trong lượt này).", {}),
])
demo_row = info_table.add_row()
for cell, w in zip(demo_row.cells, col_w):
    cell.width = dxa(w)
    set_cell_borders(cell)
cell_para(demo_row.cells[0], "", first=True)
p_demo_label = demo_row.cells[0].paragraphs[0]
add_run(p_demo_label, "*** ", bold=True, color=RED)
add_run(p_demo_label, "GHI CHÚ:", bold=True, color=RED, highlight="yellow")
cell_multiline(demo_row.cells[1], [
    ("Toàn bộ tiết học này có thể dùng làm tiết dạy minh họa: mở đầu bằng trò chơi AR gây hào "
     "hứng cao, có hoạt động nhóm thực hành thí nghiệm thật (tư duy bậc cao qua Model Builder), "
     "nhiệm vụ đa dạng gần gũi học sinh (Meme, TikTok/Shorts, Decision Tree), có phương thức "
     "đánh giá rõ ràng sau mỗi hoạt động, và Vận dụng khép vòng trực tiếp với bối cảnh Khởi "
     "động — thể hiện rõ sự liên kết chặt chẽ giữa các hoạt động giảng dạy.", {}),
])

doc.add_paragraph()

# --- TIẾN TRÌNH DẠY HỌC ---
p_ttdh = doc.add_paragraph()
add_run(p_ttdh, "  \tTIẾN TRÌNH DẠY HỌC", bold=True, size=14, color=PRIMARY)
doc.add_paragraph()

process_cols = [1565, 1920, 4697, 5348, 955]
headers = ["Hoạt động", "Mục tiêu", "Chi tiết hoạt động", "Nội dung kiến thức", "Ghi chú"]

proc_table = doc.add_table(rows=1, cols=5)
proc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
proc_table.autofit = False
for cell, w, header in zip(proc_table.rows[0].cells, process_cols, headers):
    cell.width = dxa(w)
    set_cell_borders(cell)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, header, bold=True, size=12)


def add_activity_row(hoat_dong: str, muc_tieu: list[tuple], chi_tiet: list[tuple],
                      noi_dung: list[tuple], ghi_chu: list[tuple]) -> None:
    row = proc_table.add_row()
    for cell, w in zip(row.cells, process_cols):
        cell.width = dxa(w)
        set_cell_borders(cell)
    p0 = row.cells[0].paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p0, hoat_dong, bold=True, size=11)
    cell_multiline(row.cells[1], [(t, {**k, "size": k.get("size", 11)}) for t, k in muc_tieu])
    cell_multiline(row.cells[2], [(t, {**k, "size": k.get("size", 11)}) for t, k in chi_tiet])
    cell_multiline(row.cells[3], [(t, {**k, "size": k.get("size", 11)}) for t, k in noi_dung])
    cell_multiline(row.cells[4], [(t, {**k, "size": k.get("size", 10)}) for t, k in ghi_chu])
    return row


# HĐ1 — Khởi động: Trò chơi AR "Bảo vệ hành tinh xanh"
add_activity_row(
    "Hoạt động 1\nKhởi động\n(10 phút)",
    [("Tạo hứng thú qua trải nghiệm trò chơi, huy động hiểu biết ban đầu về mối liên hệ giữa "
      "khối lượng, tốc độ và khả năng \"làm lệch hướng/truyền tương tác\" của một vật, kích "
      "thích tò mò dẫn vào khái niệm động lượng.", {})],
    [
        ("Chuyển giao:", {"bold": True}),
        ("GV tổ chức trò chơi AR \"Bảo vệ hành tinh xanh\": học sinh dùng tay điều khiển tên "
         "lửa bắn vào thiên thạch đang lao về phía Trái Đất để làm nó đổi hướng, tránh va "
         "chạm. Loại hình: trò chơi mô phỏng tương tác. Hình thức: cả lớp lần lượt trải nghiệm "
         "(2–3 lượt tiêu biểu, có thể chia đội thi đua); thời gian: 6 phút chơi + 4 phút thảo "
         "luận; tài liệu: thiết bị/màn hình chạy trò chơi AR + PHT mục 1.", {}),
        ("Thực hiện:", {"bold": True}),
        ("HS lần lượt chơi, cả lớp quan sát và ghi nhanh nhận xét vào PHT: với cùng 1 tên lửa, "
         "khi nào thiên thạch bị lệch hướng nhiều hơn — khi tên lửa nặng hơn hay khi tên lửa "
         "bắn nhanh hơn?", {}),
        ("Báo cáo:", {"bold": True}),
        ("Kĩ thuật \"Hỏi nhanh – Đáp nhanh\": GV mời 2–3 HS chia sẻ quan sát và dự đoán của "
         "mình.", {}),
        ("Kết luận:", {"bold": True}),
        ("GV chốt vấn đề: \"Cả khối lượng lẫn tốc độ của tên lửa đều ảnh hưởng đến khả năng làm "
         "lệch hướng thiên thạch. Đại lượng nào đặc trưng cho điều này? Chúng ta cùng khám phá "
         "qua thí nghiệm thực tế trong bài học hôm nay — Động lượng.\"", {}),
    ],
    [("Nêu vấn đề: khả năng truyền tương tác (làm lệch hướng/truyền chuyển động) phụ thuộc cả "
      "khối lượng và tốc độ của vật.", {})],
    [("Đánh giá định tính qua quan sát: mức độ hào hứng, số lượng học sinh xung phong chơi/trả "
      "lời — không chấm điểm, dùng để dẫn dắt.", {})],
)

# HĐ2 — Hình thành kiến thức: Động lượng (Model Builder + Meme Analyzer), có thêm thời gian
add_activity_row(
    "Hoạt động 2\nHình thành\nkiến thức mới:\nĐộng lượng\n(22 phút)",
    [("Thông qua thí nghiệm trực tiếp, rút ra được các yếu tố ảnh hưởng đến khả năng truyền "
      "chuyển động; phát biểu được định nghĩa động lượng, viết được công thức p = m.v, nêu "
      "được ý nghĩa vật lí và tính chất vectơ của động lượng.", {})],
    [
        ("Chuyển giao NV1:", {"bold": True}),
        ("GV giao NV1 \"Thí nghiệm ba viên bi A, B, C\" (PHT mục 2.a), yêu cầu HS đối chiếu "
         "ngược lại với trò chơi AR vừa chơi (viên bi ↔ tên lửa). Loại hình: Model Builder (tự "
         "làm thí nghiệm, rút ra công thức). Hình thức: nhóm 4–5 học sinh; thời gian: 10 phút; "
         "tài liệu: bộ dụng cụ thí nghiệm (máng trượt + 3 viên bi) + SGK trang 110–111 + PHT.", {}),
        ("Thực hiện NV1:", {"bold": True}),
        ("Học sinh TỰ TAY làm Thí nghiệm 1 (thả bi A rồi bi B từ cùng độ cao, so sánh quãng "
         "đường bi C lăn) và Thí nghiệm 2 (chỉ thả bi A, đổi độ dốc), ghi số liệu vào PHT, thảo "
         "luận nhóm trả lời 3 câu hỏi rút ra kết luận. Có thêm thời gian để GV mời 2 nhóm làm "
         "lại thí nghiệm trước lớp cho cả lớp cùng quan sát trực tiếp.", {}),
        ("Chuyển giao NV2:", {"bold": True}),
        ("GV chiếu 1 hình chế (meme) \"xe tăng vs xe đạp\" kèm câu chú thích gây tranh cãi. "
         "Loại hình: Giải mã Meme Vật lý (Meme Analyzer). Hình thức: nhóm đôi; thời gian: 6 "
         "phút; tài liệu: hình chế đã chuẩn bị + PHT mục 2.b.", {}),
        ("Thực hiện NV2:", {"bold": True}),
        ("HS phân tích hình chế: chú thích trong meme đúng hay sai về mặt vật lí, vì sao, dựa "
         "trên kết luận thí nghiệm vừa rút ra; tự đặt 1 bộ số liệu cụ thể để minh chứng.", {}),
        ("Báo cáo:", {"bold": True}),
        ("Đại diện 1–2 nhóm trình bày số liệu/kết luận NV1; 1–2 nhóm đôi trình bày phân tích "
         "meme NV2, các nhóm khác nhận xét, bổ sung.", {}),
        ("Kết luận:", {"bold": True}),
        ("GV chốt: khả năng truyền chuyển động phụ thuộc CẢ khối lượng lẫn tốc độ của vật — đây "
         "chính là động lượng; chốt định nghĩa p = m.v, tính vectơ, đơn vị kg.m/s, ý nghĩa vật "
         "lí; yêu cầu HS hoàn thành phần điền khuyết \"Kiến thức trọng tâm\" trong PHT.", {}),
    ],
    [
        ("Động lượng p = m.v là đại lượng vectơ, cùng hướng với vận tốc, đơn vị kg.m/s.", {}),
        ("Động lượng đặc trưng cho khả năng truyền tương tác (truyền chuyển động) giữa các vật "
         "khi va chạm — rút ra trực tiếp từ kết quả thí nghiệm (bi nặng hơn hoặc bi nhanh hơn "
         "đều làm bi C lăn xa hơn), khớp với quan sát ở trò chơi AR mở đầu.", {}),
    ],
    [("Đánh giá qua bảng số liệu thí nghiệm + câu trả lời thảo luận (NV1) và sản phẩm phân tích "
      "meme (NV2) trong PHT — GV quan sát nhóm làm thí nghiệm trực tiếp, quan sát nhanh 2–3 "
      "phiếu đại diện. Tiêu chí: đo/ghi số liệu hợp lý, kết luận đúng dựa trên cả m và v, có "
      "tinh thần hợp tác khi cùng làm thí nghiệm.", {})],
)

# HĐ3 — Luyện tập (Fact Check Influencer)
add_activity_row(
    "Hoạt động 3\nLuyện tập\n(8 phút)",
    [("Vận dụng công thức p = m.v để giải 1 câu hỏi luyện tập theo đúng 1 trong 3 dạng đề thi "
      "tốt nghiệp THPT, đồng thời rèn kĩ năng phản biện thông tin mạng xã hội.", {})],
    [
        ("Chuyển giao:", {"bold": True}),
        ("GV chiếu 1 đoạn mô phỏng phát biểu sai lan truyền trên mạng xã hội về động lượng (PHT "
         "mục 3). Loại hình: Bóc phốt TikTok/Shorts (Fact Check Influencer) kết hợp tính toán. "
         "Hình thức: cá nhân; thời gian: 5 phút; tài liệu: PHT mục 3.", {}),
        ("Thực hiện:", {"bold": True}),
        ("Học sinh chỉ ra chỗ sai trong phát biểu, giải thích bằng lập luận vật lí, và giải 1 "
         "câu tính toán minh hoạ đi kèm.", {}),
        ("Báo cáo:", {"bold": True}),
        ("Kĩ thuật gọi tên ngẫu nhiên: GV gọi ngẫu nhiên 2 học sinh lên trình bày.", {}),
        ("Kết luận:", {"bold": True}),
        ("GV nhận xét, chốt đáp án và cách trình bày chuẩn (lưu ý đổi đơn vị km/h → m/s).", {}),
    ],
    [("Vận dụng p = m.v để phản biện 1 phát biểu sai phổ biến, khắc sâu vai trò ngang nhau của "
      "khối lượng và tốc độ.", {})],
    [("Chấm nhanh tại lớp, công bố đáp án ngay — đánh giá qua bài làm trong PHT.", {})],
)

# HĐ4 — Vận dụng: Decision Tree, khép vòng với Khởi động
add_activity_row(
    "Hoạt động 4\nVận dụng, củng cố\nvà dặn dò\n(5 phút)",
    [("Vận dụng khái niệm động lượng để ra quyết định trong đúng bối cảnh trò chơi mở đầu — "
      "khép vòng Khởi động↔Vận dụng; củng cố toàn bộ nội dung tiết học và giao nhiệm vụ về "
      "nhà.", {})],
    [
        ("Chuyển giao:", {"bold": True}),
        ("GV quay lại đúng bối cảnh trò chơi AR mở đầu, giao NV \"Bản đồ lựa chọn sinh tử\" (PHT "
         "mục 4): chỉ có 1 tên lửa duy nhất, chọn phương án khối lượng/vận tốc nào để bảo vệ "
         "Trái Đất hiệu quả nhất. Loại hình: Decision Tree. Hình thức: nhóm đôi; thời gian: 3 "
         "phút; tài liệu: PHT mục 4 (sơ đồ cây lựa chọn).", {}),
        ("Thực hiện:", {"bold": True}),
        ("HS thảo luận theo cặp, đi theo sơ đồ cây lựa chọn, chọn nhánh và giải thích lí do dựa "
         "trên p = m.v.", {}),
        ("Báo cáo:", {"bold": True}),
        ("1–2 cặp chia sẻ nhanh lựa chọn trước lớp, biểu quyết nhanh cả lớp xem đa số chọn "
         "nhánh nào.", {}),
        ("Kết luận:", {"bold": True}),
        ("GV chốt kiến thức trọng tâm cả bài (định nghĩa, công thức, ý nghĩa vật lí động "
         "lượng); giao Bài tập về nhà và dặn dò đọc trước mục II. Xung lượng của lực (nếu có "
         "buổi học tiếp theo).", {}),
    ],
    [("Vận dụng p = m.v để ra quyết định trong tình huống thực tế/giả lập; tổng kết bài học.", {})],
    [("Câu hỏi vận dụng miệng + biểu quyết nhanh, không chấm điểm. Dặn dò: hoàn thành Bài tập "
      "về nhà.", {})],
)

doc.add_paragraph()

# --- PHẢN HỒI – ĐIỀU CHỈNH ---
p_ph = doc.add_paragraph()
add_run(p_ph, "PHẢN HỒI – ĐIỀU CHỈNH", bold=True, size=14, color=PRIMARY)

for label in ("Ưu điểm", "Hạn chế", "Hướng điều chỉnh"):
    p = doc.add_paragraph()
    add_run(p, f"{label}: ", bold=True)
    add_run(p, "..." * 30, color=RGBColor(0xBF, 0xBF, 0xBF))

OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
doc.save(str(OUT_PATH))
print("OK - da xuat KHBD V3 (1 tiet, chi Dong luong)")
