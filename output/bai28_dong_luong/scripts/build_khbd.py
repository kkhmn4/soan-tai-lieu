"""Soạn Kế hoạch bài dạy (KHBD) cho Bài 28 - Động lượng theo ĐÚNG mẫu
"KẾ HOẠCH GIẢNG DẠY.docx" (mẫu Trường THCS và THPT Đinh Thiện Lý) do người
dùng cung cấp — KHÁC với mẫu Công văn 5512/Phụ lục IV mà `gems/docx_export/
khbd_exporter.py` xuất ra. Đây là tài liệu một lần (one-off), không dùng
`khbd_exporter.py`, không sửa bất kỳ file nào trong `gems/`.

Cấu trúc bám sát nguyên bản mẫu (đọc trực tiếp từ file .docx gốc):
  - Tiêu đề "TÊN BÀI/CHỦ ĐỀ"
  - Bảng thông tin: Tuần / Lớp / Sách (tài liệu) / Mục tiêu / Tiến trình /
    Chuẩn bị của giáo viên / Chuẩn bị của học sinh / Tài liệu hỗ trợ giảng
    dạy, học tập / *** HOẠT ĐỘNG CHỌN DEMO
  - "TIẾN TRÌNH DẠY HỌC": bảng 5 cột Hoạt động | Mục tiêu | Chi tiết hoạt
    động | Nội dung kiến thức | Ghi chú
  - "PHẢN HỒI – ĐIỀU CHỈNH"
Khổ trang/lề/độ rộng cột lấy đúng số liệu (đơn vị dxa) từ file mẫu gốc.

Nội dung được soạn để bám sát Tiêu chí đánh giá giáo án + hoạt động giảng
dạy minh họa (Trường Đinh Thiện Lý): mục tiêu cụ thể + mở rộng thực tiễn phù
hợp, hoạt động liên kết chặt chẽ + kĩ thuật dạy học tích cực + phát triển tư
duy bậc cao/hợp tác nhóm, có phương thức đánh giá sau mỗi hoạt động.

Chạy: python "output/bai28_dong_luong/scripts/build_khbd.py"
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Emu, Pt, RGBColor

REPO_ROOT = Path(__file__).resolve().parent.parent.parent.parent
OUT_PATH = REPO_ROOT / "output" / "bai28_dong_luong" / "ready" / "bai28_dong_luong_ke_hoach_bai_day.docx"
IMG_DIR = REPO_ROOT / "output" / "bai28_dong_luong" / "ready" / "hinh_anh"

FONT = "Times New Roman"
PRIMARY = RGBColor(0x1F, 0x4E, 0x79)
RED = RGBColor(0xC0, 0x00, 0x00)
DARK = RGBColor(0x00, 0x00, 0x00)

TWIP = Emu(635)  # 1 dxa (twip) = 635 EMU


def dxa(value: int) -> Emu:
    return Emu(value * 635)


def set_default_font(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(12)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), FONT)


def set_cell_borders(cell) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    borders = tcPr.makeelement(qn("w:tcBorders"), {})
    for edge in ("top", "left", "bottom", "right"):
        el = borders.makeelement(qn(f"w:{edge}"), {})
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        borders.append(el)
    tcPr.append(borders)


def add_run(p, text: str, *, bold=False, italic=False, size=12, color=DARK, highlight=None) -> None:
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.name = FONT
    r.font.size = Pt(size)
    r.font.color.rgb = color
    if highlight:
        rpr = r._element.get_or_add_rPr()
        hl = rpr.makeelement(qn("w:highlight"), {})
        hl.set(qn("w:val"), highlight)
        rpr.append(hl)


def cell_para(cell, text: str = "", *, bold=False, italic=False, size=12, color=DARK,
              space_after=2, first=False):
    p = cell.paragraphs[0] if first and not cell.paragraphs[0].runs else cell.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    if text:
        add_run(p, text, bold=bold, italic=italic, size=size, color=color)
    return p


def add_cell_image(cell, image_path: Path, width_cm: float) -> None:
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image_path), width=Cm(width_cm))


def cell_multiline(cell, lines: list[tuple], *, first_clear=True):
    """`lines` là list các tuple (text, kwargs) để ghi nhiều đoạn trong 1 ô."""
    if first_clear:
        cell.paragraphs[0].text = ""
    for i, (text, kwargs) in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(kwargs.pop("space_after", 3))
        p.paragraph_format.space_before = Pt(0)
        add_run(p, text, **kwargs)


# ============================================================
# DỰNG TÀI LIỆU
# ============================================================
doc = Document()
set_default_font(doc)

section = doc.sections[0]
section.orientation = WD_ORIENT.LANDSCAPE
section.page_width = dxa(15840)
section.page_height = dxa(12240)
section.top_margin = dxa(720)
section.bottom_margin = dxa(720)
section.left_margin = dxa(720)
section.right_margin = dxa(720)
section.header_distance = dxa(720)
section.footer_distance = dxa(300)

# --- Tiêu đề ---
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p_title, "TÊN BÀI/CHỦ ĐỀ: ", bold=True, size=16)
add_run(p_title, "BÀI 28 – ĐỘNG LƯỢNG", bold=True, size=16, color=PRIMARY)
doc.add_paragraph()

# --- Bảng thông tin chung ---
info_table = doc.add_table(rows=0, cols=2)
info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
info_table.autofit = False
col_w = [3847, 10267]


def add_info_row(label: str, content_lines: list[tuple], *, label_color=DARK):
    row = info_table.add_row()
    row.cells[0].width = dxa(col_w[0])
    row.cells[1].width = dxa(col_w[1])
    for cell, w in zip(row.cells, col_w):
        cell.width = dxa(w)
        set_cell_borders(cell)
    cell_para(info_table.rows[-1].cells[0], label, bold=True, color=label_color, first=True)
    cell_multiline(info_table.rows[-1].cells[1], content_lines)
    return row


add_info_row("Tuần", [("......... (Giáo viên điền theo phân phối chương trình thực tế)", {"italic": True})])
add_info_row("Lớp", [("10", {})])
add_info_row("Sách (tài liệu)", [
    ("Vật lí 10 – Bộ sách Kết nối tri thức với cuộc sống, NXB Giáo dục Việt Nam.", {}),
    ("Bài 28 – Động lượng (Chương V – Động lượng), trang 110–112.", {}),
])
add_info_row("Mục tiêu", [
    ("1. Kiến thức", {"bold": True}),
    ("– Phát biểu được định nghĩa của động lượng và nêu được ý nghĩa vật lí của đại lượng động lượng.", {}),
    ("– Phát biểu và viết được công thức liên hệ giữa lực tác dụng lên vật và tốc độ biến thiên động "
     "lượng (dạng thứ hai của định luật II Newton).", {}),
    ("2. Năng lực", {"bold": True}),
    ("a) Năng lực chung: Tự học (chủ động đóng góp ý tưởng, đặt câu hỏi và trả lời yêu cầu); Giao tiếp "
     "và hợp tác (khiêm tốn tiếp thu góp ý, nhiệt tình chia sẻ, hỗ trợ thành viên trong nhóm).", {}),
    ("b) Năng lực vật lí: Nêu được ý nghĩa vật lí và định nghĩa động lượng; phát biểu được mối liên hệ "
     "giữa lực tác dụng lên vật và tốc độ biến thiên động lượng; viết được công thức liên hệ.", {}),
    ("3. Phẩm chất", {"bold": True}),
    ("Trách nhiệm (hoàn thành nhiệm vụ cá nhân/nhóm được giao); Trung thực (báo cáo đúng kết quả thảo "
     "luận, không sao chép kết quả nhóm khác).", {}),
    ("4. Mở rộng, liên hệ thực tiễn", {"bold": True}),
    ("An toàn giao thông (nguyên lí kéo dài thời gian va chạm của túi khí ô tô, dây an toàn); kĩ thuật "
     "thể thao (thủ môn bắt bóng, cầu thủ sút bóng, vợt bóng bàn/tennis); an toàn lao động và thiết bị "
     "bảo hộ (mũ bảo hiểm); kĩ thuật đóng cọc bằng búa máy.", {}),
])
add_info_row("Tiến trình", [
    ("Hoạt động 1 – Khởi động (5 phút)", {"bold": True}),
    ("Hoạt động 2 – Hình thành kiến thức mới: 2.1. Động lượng (12 phút)", {"bold": True}),
    ("Hoạt động 3 – Hình thành kiến thức mới: 2.2. Xung lượng của lực (15 phút)", {"bold": True}),
    ("Hoạt động 4 – Luyện tập (8 phút)", {"bold": True}),
    ("Hoạt động 5 – Vận dụng, củng cố và dặn dò (5 phút)", {"bold": True}),
    ("(Tổng thời lượng: 1 tiết – 45 phút, thực hiện trọn vẹn cả 2 đơn vị kiến thức của Bài 28 theo đúng "
     "2 yêu cầu cần đạt.)", {"italic": True, "size": 11}),
])
add_info_row("Chuẩn bị của giáo viên", [
    ("– SGK Vật lí 10 (Kết nối tri thức), Kế hoạch bài dạy, Phiếu học tập (PHT) Bài 28.", {}),
    ("– Slide GIÁO VIÊN trình chiếu cả lớp (bai28_dong_luong_slide_giao_vien.pptx) và Slide PHIẾU HỌC "
     "TẬP chiếu song song để học sinh đối chiếu (bai28_dong_luong_slide_phieu_hoc_tap.pptx).", {}),
    ("– Bộ dụng cụ thí nghiệm Hình 28.1: máng trượt (có thể dùng ống nhựa cắt dọc) + 3 viên bi A, B, C "
     "(bi B nặng hơn bi A) — chuẩn bị đủ số bộ cho các nhóm (khuyến khích mỗi nhóm 1 bộ để tự làm thí "
     "nghiệm trực tiếp; nếu không đủ dụng cụ, GV làm mẫu 1 bộ, các nhóm quan sát trực tiếp và tự ghi "
     "số liệu vào PHT).", {}),
    ("– Sơ đồ thí nghiệm bi A/B/C, vectơ động lượng và đồ thị lực – thời gian (GV tự thiết kế, xem thư "
     "mục ready/hinh_anh/).", {}),
    ("– Quả bóng nhỏ (nếu có điều kiện) để minh họa trực tiếp trên lớp; phiếu quan sát/đánh giá hoạt "
     "động nhóm.", {}),
])
add_info_row("Chuẩn bị của học sinh", [
    ("– SGK Vật lí 10, vở ghi, dụng cụ học tập.", {}),
    ("– Ôn lại định luật II Newton (F = m.a) và khái niệm gia tốc đã học ở Bài 15.", {}),
    ("– Đọc trước nội dung Bài 28 – Động lượng ở nhà.", {}),
])
add_info_row("Tài liệu hỗ trợ giảng dạy, học tập.", [
    ("– Phiếu học tập Bài 28 – Động lượng (bai28_dong_luong_phieu_hoc_tap.docx).", {}),
    ("– Bài tập về nhà Bài 28 – Động lượng (bai28_dong_luong_bai_tap_ve_nha.docx).", {}),
    ("– Slide Giáo viên (bai28_dong_luong_slide_giao_vien.pptx) và Slide Phiếu học tập "
     "(bai28_dong_luong_slide_phieu_hoc_tap.pptx) — cùng bộ ảnh minh họa, khác tông màu.", {}),
    ("– Hình ảnh minh họa: vectơ động lượng, đồ thị lực – thời gian, sơ đồ thí nghiệm bi A/B/C, sơ đồ "
     "va chạm bóng – tường (thư mục ready/hinh_anh/).", {}),
])
demo_row = info_table.add_row()
for cell, w in zip(demo_row.cells, col_w):
    cell.width = dxa(w)
    set_cell_borders(cell)
cell_para(demo_row.cells[0], "", first=True)
p_demo_label = demo_row.cells[0].paragraphs[0]
add_run(p_demo_label, "*** ", bold=True, color=RED)
add_run(p_demo_label, "HOẠT ĐỘNG CHỌN DEMO:", bold=True, color=RED, highlight="yellow")
cell_multiline(demo_row.cells[1], [
    ("Đề xuất Hoạt động 3 – Hình thành kiến thức mới: 2.2. Xung lượng của lực (15 phút) làm hoạt "
     "động dạy minh họa.", {"bold": True}),
    ("Lí do: hoạt động có đủ các yếu tố được tiêu chí đánh giá chú trọng — làm việc nhóm (khăn trải "
     "bàn), phát triển tư duy bậc cao (đọc/phân tích đồ thị, tự suy luận xây dựng công thức bằng kĩ "
     "thuật Model Builder), có phương thức đánh giá rõ ràng qua sản phẩm PHT, và liên hệ thực tiễn tự "
     "nhiên (thủ môn bắt bóng). Giáo viên có thể điều chỉnh lựa chọn khác tùy điều kiện lớp học thực "
     "tế.", {}),
])

doc.add_paragraph()

# --- TIẾN TRÌNH DẠY HỌC ---
p_ttdh = doc.add_paragraph()
add_run(p_ttdh, "  \tTIẾN TRÌNH DẠY HỌC", bold=True, size=14, color=PRIMARY)
doc.add_paragraph()

process_cols = [1565, 1920, 4697, 5348, 955]
headers = ["Hoạt động", "Mục tiêu", "Chi tiết hoạt động", "Nội dung kiến thức", "Ghi chú"]

proc_table = doc.add_table(rows=1, cols=5)
proc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
proc_table.autofit = False
for cell, w, header in zip(proc_table.rows[0].cells, process_cols, headers):
    cell.width = dxa(w)
    set_cell_borders(cell)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, header, bold=True, size=12)


def add_activity_row(hoat_dong: str, muc_tieu: list[tuple], chi_tiet: list[tuple],
                      noi_dung: list[tuple], ghi_chu: list[tuple]) -> None:
    row = proc_table.add_row()
    for cell, w in zip(row.cells, process_cols):
        cell.width = dxa(w)
        set_cell_borders(cell)
    p0 = row.cells[0].paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p0, hoat_dong, bold=True, size=11)
    cell_multiline(row.cells[1], [(t, {**k, "size": k.get("size", 11)}) for t, k in muc_tieu])
    cell_multiline(row.cells[2], [(t, {**k, "size": k.get("size", 11)}) for t, k in chi_tiet])
    cell_multiline(row.cells[3], [(t, {**k, "size": k.get("size", 11)}) for t, k in noi_dung])
    cell_multiline(row.cells[4], [(t, {**k, "size": k.get("size", 10)}) for t, k in ghi_chu])
    return row


# Hoạt động 1 — Khởi động
add_activity_row(
    "Hoạt động 1\nKhởi động\n(5 phút)",
    [("Tạo tình huống có vấn đề, huy động hiểu biết ban đầu của học sinh về \"khả năng gây khó khăn "
      "khi làm vật dừng lại/đổi hướng\", kích thích tò mò dẫn vào khái niệm động lượng.", {})],
    [
        ("Chuyển giao:", {"bold": True}),
        ("GV chiếu 2 tình huống mở bài SGK trang 110 (xe tải – xe con cùng tốc độ khi phanh gấp; cầu "
         "thủ sút phạt 11 m). Hình thức: cá nhân; thời gian: 3 phút; tài liệu: hình ảnh GV trình chiếu "
         "(NV1 – PHT mục 1).", {}),
        ("Thực hiện:", {"bold": True}),
        ("HS quan sát, suy nghĩ cá nhân, ghi nhanh dự đoán ra nháp.", {}),
        ("Báo cáo:", {"bold": True}),
        ("Kĩ thuật \"Hỏi nhanh – Đáp nhanh\": GV mời 2–3 HS xung phong nêu dự đoán và giải thích ngắn "
         "gọn.", {}),
        ("Kết luận:", {"bold": True}),
        ("GV chưa chốt đúng/sai, dẫn dắt: \"Đại lượng nào đặc trưng cho điều này? Chúng ta cùng tìm "
         "hiểu qua bài học hôm nay — Động lượng.\"", {}),
    ],
    [("Nêu vấn đề: khả năng truyền chuyển động khi tương tác phụ thuộc cả khối lượng và tốc độ của "
      "vật.", {})],
    [("Đánh giá định tính qua quan sát: mức độ hào hứng, số lượng học sinh xung phong trả lời.", {})],
)

# Hoạt động 2 — 2.1 Động lượng
row2 = add_activity_row(
    "Hoạt động 2\nHình thành\nkiến thức mới\n2.1. Động lượng\n(12 phút)",
    [("Thông qua thí nghiệm trực tiếp, rút ra được các yếu tố ảnh hưởng đến khả năng truyền chuyển "
      "động; phát biểu được định nghĩa động lượng, viết được công thức p = m.v, nêu được ý nghĩa vật "
      "lí và tính chất vectơ của động lượng (đáp ứng yêu cầu cần đạt 1).", {})],
    [
        ("Chuyển giao:", {"bold": True}),
        ("GV giao NV2 \"Thí nghiệm ba viên bi A, B, C\" (PHT mục 2.1.a). Hình thức: nhóm 4–5 học sinh; "
         "thời gian: 7 phút; tài liệu: bộ dụng cụ thí nghiệm (máng trượt + 3 viên bi) + SGK trang "
         "110–111 + PHT.", {}),
        ("Thực hiện:", {"bold": True}),
        ("Học sinh TỰ TAY làm Thí nghiệm 1 (thả bi A rồi bi B từ cùng độ cao, so sánh quãng đường bi C "
         "lăn) và Thí nghiệm 2 (chỉ thả bi A, đổi độ dốc), ghi số liệu quan sát vào bảng kết quả trong "
         "PHT, thảo luận nhóm để trả lời 3 câu hỏi rút ra kết luận.", {}),
        ("Báo cáo:", {"bold": True}),
        ("Đại diện 1–2 nhóm trình bày số liệu và kết luận rút ra, các nhóm khác đối chiếu kết quả, "
         "nhận xét, bổ sung.", {}),
        ("Kết luận:", {"bold": True}),
        ("GV chốt: khả năng truyền chuyển động phụ thuộc CẢ khối lượng lẫn tốc độ của vật — đây chính "
         "là động lượng; chốt định nghĩa p = m.v, tính vectơ, đơn vị kg.m/s, ý nghĩa vật lí; chiếu sơ "
         "đồ vectơ động lượng minh họa; yêu cầu HS hoàn thành phần điền khuyết \"Kiến thức trọng tâm\" "
         "trong PHT.", {}),
    ],
    [
        ("Động lượng p = m.v là đại lượng vectơ, cùng hướng với vận tốc, đơn vị kg.m/s.", {}),
        ("Động lượng đặc trưng cho khả năng truyền tương tác (truyền chuyển động) giữa các vật khi va "
         "chạm — rút ra trực tiếp từ kết quả thí nghiệm (bi nặng hơn hoặc bi nhanh hơn đều làm bi C "
         "lăn xa hơn).", {}),
    ],
    [("Đánh giá qua bảng số liệu thí nghiệm + câu trả lời thảo luận trong PHT (GV quan sát nhóm làm "
      "thí nghiệm trực tiếp, quan sát nhanh 2–3 phiếu đại diện). Tiêu chí: đo/ghi số liệu hợp lý, kết "
      "luận đúng dựa trên cả m và v, có tinh thần hợp tác khi cùng làm thí nghiệm.", {})],
)
add_cell_image(row2.cells[3], IMG_DIR / "so_do_thi_nghiem_bi_abc.png", 7.5)

# Hoạt động 3 — 2.2 Xung lượng của lực (đề xuất demo)
row3 = add_activity_row(
    "Hoạt động 3\nHình thành\nkiến thức mới\n2.2. Xung lượng\ncủa lực\n(15 phút)\n★ Đề xuất demo",
    [("Nêu được định nghĩa xung lượng của lực, thiết lập được công thức liên hệ giữa xung lượng của "
      "lực và độ biến thiên động lượng, viết được dạng tổng quát của định luật II Newton (đáp ứng yêu "
      "cầu cần đạt 2).", {})],
    [
        ("Chuyển giao:", {"bold": True}),
        ("GV giao đồng thời NV3 \"Đọc đồ thị lực – thời gian\" (nhóm 4 học sinh, 7 phút, tài liệu: đồ "
         "thị GV chiếu + SGK trang 111) và NV4 \"Suy luận công thức liên hệ\" (cá nhân 3 phút rồi thảo "
         "luận chung lớp 2 phút, tài liệu: công thức định luật II Newton đã học ở Bài 15) — PHT mục "
         "2.2.a.", {}),
        ("Thực hiện:", {"bold": True}),
        ("Kĩ thuật khăn trải bàn: nhóm 4 HS thảo luận đọc đồ thị. Kĩ thuật đặt câu hỏi dẫn dắt "
         "(scaffolding từng bước a→b→c→d trong PHT): cá nhân tự suy luận biến đổi F = m.a thành F.Δt = "
         "Δp.", {}),
        ("Báo cáo:", {"bold": True}),
        ("Đại diện nhóm trình bày kết quả đọc đồ thị; 1 học sinh trình bày bảng chuỗi suy luận công "
         "thức, cả lớp nhận xét từng bước biến đổi.", {}),
        ("Kết luận:", {"bold": True}),
        ("GV chốt: xung lượng của lực F.Δt (đơn vị N.s); công thức liên hệ F.Δt = Δp = p2 – p1; dạng "
         "tổng quát định luật II Newton F = Δp/Δt. Liên hệ ngay thực tế thủ môn bắt bóng (từ ý c, "
         "NV3).", {}),
    ],
    [
        ("Xung lượng của lực F.Δt (N.s).", {}),
        ("Liên hệ xung lượng của lực và độ biến thiên động lượng: F.Δt = Δp = p2 – p1.", {}),
        ("Dạng tổng quát của định luật II Newton: F = Δp/Δt.", {}),
    ],
    [("Đánh giá qua sản phẩm PHT (đọc đúng đồ thị + chuỗi suy luận 4 bước a–d) và câu trả lời khi báo "
      "cáo. Hoạt động có nhiều tương tác nhóm và tư duy bậc cao (phân tích đồ thị, suy luận xây dựng "
      "công thức) — phù hợp chọn làm hoạt động dạy minh họa.", {})],
)
add_cell_image(row3.cells[3], IMG_DIR / "do_thi_luc_thoi_gian.png", 7.5)

# Hoạt động 4 — Luyện tập
add_activity_row(
    "Hoạt động 4\nLuyện tập\n(8 phút)",
    [("Vận dụng công thức p = m.v và F = Δp/Δt để giải các câu hỏi luyện tập bám sát cấu trúc đề thi "
      "tốt nghiệp THPT.", {})],
    [
        ("Chuyển giao:", {"bold": True}),
        ("GV giao 2 bài luyện tập (PHT mục 3.1: câu trắc nghiệm 4 phương án; PHT mục 3.2: câu trả lời "
         "ngắn). Hình thức: cá nhân; thời gian: 5 phút; tài liệu: PHT mục 3.1 và 3.2.", {}),
        ("Thực hiện:", {"bold": True}),
        ("Học sinh làm bài cá nhân vào PHT.", {}),
        ("Báo cáo:", {"bold": True}),
        ("Kĩ thuật gọi tên ngẫu nhiên: GV gọi ngẫu nhiên 2 học sinh lên trình bày lời giải trên bảng.", {}),
        ("Kết luận:", {"bold": True}),
        ("GV nhận xét, chốt đáp án và cách trình bày chuẩn (lưu ý đổi đơn vị km/h → m/s).", {}),
    ],
    [("Vận dụng p = m.v tính động lượng xe máy; vận dụng F = Δp/Δt tính lực trung bình vợt tác dụng "
      "lên bóng bàn.", {})],
    [("Đánh giá qua bài làm trong PHT (chấm nhanh đúng/sai đáp số), kết hợp phần trình bày trên bảng "
      "của 2 học sinh được gọi.", {})],
)

# Hoạt động 5 — Vận dụng
add_activity_row(
    "Hoạt động 5\nVận dụng, củng cố\nvà dặn dò\n(5 phút)",
    [("Vận dụng kiến thức để phân tích, đánh giá một tình huống kĩ thuật thực tế (phát triển tư duy "
      "bậc cao); củng cố toàn bài và giao nhiệm vụ về nhà.", {})],
    [
        ("Chuyển giao:", {"bold": True}),
        ("GV giao NV5 \"Bắt lỗi thiết kế mũ bảo hiểm lót thép\" (PHT mục 4.2.a). Hình thức: nhóm đôi; "
         "thời gian: 3 phút; tài liệu: PHT mục 4.2.a + công thức F = Δp/Δt.", {}),
        ("Thực hiện:", {"bold": True}),
        ("Học sinh thảo luận nhanh theo cặp, ghi câu trả lời vào PHT.", {}),
        ("Báo cáo:", {"bold": True}),
        ("1–2 cặp chia sẻ nhanh câu trả lời trước lớp.", {}),
        ("Kết luận:", {"bold": True}),
        ("GV chốt kiến thức trọng tâm toàn bài bằng sơ đồ tư duy ngắn gọn (động lượng – xung lượng – "
         "liên hệ – định luật II Newton dạng tổng quát); giao Bài tập về nhà (18 trắc nghiệm + 4 Đúng/"
         "Sai + 6 trả lời ngắn) và phần \"Mở rộng\" trong PHT (đọc thêm ở nhà nếu chưa đủ thời gian).", {}),
    ],
    [("Vận dụng liên hệ xung lượng – độ biến thiên động lượng để phân tích ưu/nhược điểm một thiết kế "
      "bảo hộ; tổng kết toàn bài.", {})],
    [("Đánh giá qua chất lượng phân tích của học sinh (có chỉ ra đúng điểm bất hợp lí liên quan đến "
      "Δt hay không). Dặn dò: hoàn thành Bài tập về nhà; chuẩn bị Bài 29 – Định luật bảo toàn động "
      "lượng.", {})],
)

doc.add_paragraph()

# --- PHẢN HỒI – ĐIỀU CHỈNH ---
p_ph = doc.add_paragraph()
add_run(p_ph, "PHẢN HỒI – ĐIỀU CHỈNH", bold=True, size=14, color=PRIMARY)

for label in ("Ưu điểm", "Hạn chế", "Hướng điều chỉnh"):
    p = doc.add_paragraph()
    add_run(p, f"{label}: ", bold=True)
    add_run(p, "..." * 30, color=RGBColor(0xBF, 0xBF, 0xBF))

OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
doc.save(str(OUT_PATH))
print("OK - da xuat KHBD")
