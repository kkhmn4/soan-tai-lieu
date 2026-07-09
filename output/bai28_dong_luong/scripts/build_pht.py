"""Soạn Phiếu học tập (PHT) cho Bài 28 - Động lượng theo cấu trúc MỚI do
người dùng yêu cầu (khác chuẩn PHT phẳng 3 mục của `gems` skill v9.4.0):

    1. KHỞI ĐỘNG
    2. HÌNH THÀNH KIẾN THỨC MỚI
       2.1 <tên đơn vị kiến thức 1>          (không còn chữ "ĐVKT")
           a. Nhiệm vụ
           b. Kiến thức trọng tâm            (trình bày dạng gạch đầu dòng
                                               ngắn gọn, không viết đoạn văn dài)
       2.2 <tên đơn vị kiến thức 2>
    3. LUYỆN TẬP
       3.1 / 3.2 theo từng đơn vị kiến thức
    4. VẬN DỤNG - MỞ RỘNG
       4.1 / 4.2 theo từng đơn vị kiến thức

Khoảng trống để học sinh làm bài dùng DÒNG KẺ LIỀN (như vở kẻ ngang), số dòng
tùy độ dài câu trả lời kỳ vọng — không dùng dòng chấm `[DOT_LINE_90]` của
chuẩn PHT thường (chỉ giữ dấu chấm NGẮN cho chỗ trống điền 1 từ/cụm từ trong
câu — 2 việc khác nhau).

Đây là tài liệu một lần (one-off, Vật Lý 10, mẫu khác chuẩn GEMS Vật Lý 12),
dựng trực tiếp bằng python-docx, KHÔNG dùng `gems/docx_export/pht_exporter.py`,
KHÔNG sửa file nào trong `gems/`.

Chạy: python "output/bai28_dong_luong/scripts/build_pht.py"
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

REPO_ROOT = Path(__file__).resolve().parent.parent.parent.parent
OUT_DIR = REPO_ROOT / "output" / "bai28_dong_luong" / "ready"
IMG_DIR = OUT_DIR / "hinh_anh"
OUT_PATH = OUT_DIR / "bai28_dong_luong_phieu_hoc_tap.docx"

FONT = "Times New Roman"
PRIMARY = RGBColor(0x1F, 0x4E, 0x79)
DARK = RGBColor(0x00, 0x00, 0x00)
GRAY = RGBColor(0x59, 0x59, 0x59)


# ============================================================
# HẠ TẦNG DỰNG DOCX
# ============================================================
def set_default_font(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(13)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), FONT)


def add_run(p, text: str, *, bold=False, italic=False, size=13, color=DARK) -> None:
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.name = FONT
    r.font.size = Pt(size)
    r.font.color.rgb = color


def add_para(doc, text: str = "", *, bold=False, italic=False, size=13, color=DARK,
             space_before=0, space_after=6, align=None, indent_cm=None) -> object:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if align is not None:
        p.alignment = align
    if indent_cm is not None:
        p.paragraph_format.left_indent = Cm(indent_cm)
    if text:
        add_run(p, text, bold=bold, italic=italic, size=size, color=color)
    return p


def section_heading(doc, number: str, title: str) -> None:
    p = add_para(doc, space_before=14, space_after=6)
    add_run(p, f"{number}. {title}", bold=True, size=15, color=PRIMARY)


def subsection_heading(doc, number: str, title: str) -> None:
    p = add_para(doc, space_before=8, space_after=4)
    add_run(p, f"{number} {title}", bold=True, size=14, color=DARK)


def sub_item_label(doc, letter: str, label: str, *, indent_cm=0.5) -> object:
    p = add_para(doc, space_before=4, space_after=2, indent_cm=indent_cm)
    add_run(p, f"{letter}. {label}", bold=True, size=13)
    return p


def body(doc, text: str, *, indent_cm=0.9, size=13, italic=False, bold=False, space_after=4):
    return add_para(doc, text, size=size, italic=italic, bold=bold, indent_cm=indent_cm,
                     space_after=space_after)


def instructions(doc, text: str, *, indent_cm=0.9):
    return body(doc, f"Hướng dẫn thực hiện: {text}", italic=True, size=12, indent_cm=indent_cm)


def bullet(doc, text: str, *, indent_cm=1.3):
    p = add_para(doc, space_after=3, indent_cm=indent_cm)
    add_run(p, "– ", bold=True)
    add_run(p, text)
    return p


def _set_bottom_border(p, *, sz=6, color="000000") -> None:
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(sz))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)
    pPr.append(pbdr)


def ruled_lines(doc, n: int, *, indent_cm=0.9) -> None:
    """`n` dòng kẻ liền (như vở kẻ ngang) để học sinh viết câu trả lời — số
    dòng tùy độ dài đáp án kỳ vọng, KHÔNG dùng dòng chấm."""
    for _ in range(n):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(indent_cm)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(11)
        p.paragraph_format.line_spacing = 1.0
        r = p.add_run(" ")
        r.font.name = FONT
        r.font.size = Pt(13)
        _set_bottom_border(p)


def embed_image(doc, image_path: Path, width_cm: float) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    p.add_run().add_picture(str(image_path), width=Cm(width_cm))


def separator(doc) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    _set_bottom_border(p, sz=4, color="BFBFBF")


def _set_cell_borders(cell) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "BFBFBF")
        borders.append(el)
    tcPr.append(borders)


def data_table(doc, headers: list[str], rows: list[list[str]], *, indent_cm=1.1) -> None:
    """Bảng nhỏ để học sinh ghi số liệu thí nghiệm (không phải bảng câu hỏi) — ô trống để viết
    có viền rõ ràng và đường chấm gợi ý chỗ điền."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.autofit = True
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        _set_cell_borders(hdr[i])
        add_run(hdr[i].paragraphs[0], h, bold=True, size=12)
    for r_idx, row in enumerate(rows, start=1):
        cells = table.rows[r_idx].cells
        for c_idx, val in enumerate(row):
            _set_cell_borders(cells[c_idx])
            add_run(cells[c_idx].paragraphs[0], val or "…………………………", size=12)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent_cm)
    p.paragraph_format.space_after = Pt(8)


# ============================================================
# DỰNG TÀI LIỆU
# ============================================================
doc = Document()
set_default_font(doc)

section = doc.sections[0]
section.top_margin = Cm(2.0)
section.bottom_margin = Cm(2.0)
section.left_margin = Cm(2.0)
section.right_margin = Cm(2.0)

# --- Header ---
p_school = add_para(doc, "TRƯỜNG: ................................................",
                     italic=True, size=12, color=GRAY, space_after=2)

p_title = add_para(doc, space_after=2, align=WD_ALIGN_PARAGRAPH.CENTER)
add_run(p_title, "PHIẾU HỌC TẬP", bold=True, size=20, color=PRIMARY)
p_sub = add_para(doc, "BÀI 28 – ĐỘNG LƯỢNG", bold=True, size=14, color=DARK,
                  align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)

info_table = doc.add_table(rows=2, cols=2)
info_table.autofit = True
cells = info_table.rows[0].cells
add_run(cells[0].paragraphs[0], "Họ và tên học sinh: ..........................................", size=12)
add_run(cells[1].paragraphs[0], "Lớp: .............. Nhóm: ..............", size=12)
cells2 = info_table.rows[1].cells
add_run(cells2[0].paragraphs[0], "Môn học: Vật lí 10 (CTGDPT 2018)", size=12)
add_run(cells2[1].paragraphs[0], "Tiết học: ......... Ngày: ......./......./.......", size=12)
for row in info_table.rows:
    for cell in row.cells:
        for p in cell.paragraphs:
            p.paragraph_format.space_after = Pt(3)
doc.add_paragraph()
separator(doc)

# ============================================================
# 1. KHỞI ĐỘNG
# ============================================================
section_heading(doc, "1", "KHỞI ĐỘNG")
instructions(doc, "hoạt động cá nhân, 3 phút, quan sát 2 hình ảnh tình huống mở bài do giáo viên "
                   "trình chiếu và đọc phần mở bài SGK trang 110.", indent_cm=0.5)

sub_item_label(doc, "a", "Tình huống 1")
body(doc, "Xe tải và xe con đang chạy cạnh nhau với CÙNG một tốc độ. Khi đèn tín hiệu chuyển đỏ, xe "
          "nào cần một lực hãm (phanh) lớn hơn để dừng lại trong cùng một khoảng thời gian? Vì sao?")
ruled_lines(doc, 2)

sub_item_label(doc, "b", "Tình huống 2")
body(doc, "Trong pha sút phạt 11 m, thủ môn khó bắt bóng hơn khi bóng bay tới với tốc độ lớn hay tốc "
          "độ nhỏ? Vì sao?")
ruled_lines(doc, 2)

sub_item_label(doc, "c", "Dự đoán")
body(doc, "Từ hai tình huống trên, đại lượng nào đặc trưng cho \"khả năng gây khó khăn khi muốn làm "
          "vật dừng lại/đổi hướng\" và phụ thuộc vào những yếu tố nào của vật đang chuyển động?")
ruled_lines(doc, 2)
separator(doc)

# ============================================================
# 2. HÌNH THÀNH KIẾN THỨC MỚI
# ============================================================
section_heading(doc, "2", "HÌNH THÀNH KIẾN THỨC MỚI")

# --- 2.1 Động lượng ---
subsection_heading(doc, "2.1", "Động lượng")

sub_item_label(doc, "a", "Nhiệm vụ — Thí nghiệm ba viên bi A, B, C")
instructions(doc, "hoạt động nhóm 4–5 học sinh, 7 phút, dùng bộ dụng cụ thí nghiệm (máng trượt + 3 "
                   "viên bi A, B, C — bi B nặng hơn bi A) đã chuẩn bị sẵn, đọc mục I. Động lượng SGK "
                   "trang 110–111.")
embed_image(doc, IMG_DIR / "so_do_thi_nghiem_bi_abc.png", 9.0)
body(doc, "Chuẩn bị: đặt viên bi C ngay dưới chân máng trượt như hình vẽ.", indent_cm=1.3)

body(doc, "Thí nghiệm 1: Lần lượt thả bi A rồi bi B từ CÙNG một độ cao trên máng trượt (cùng tốc độ "
          "khi tới chân dốc). Quan sát và đo quãng đường bi C dịch chuyển sau mỗi lần va chạm.",
     indent_cm=1.3)
data_table(doc, ["Lần thả", "Quãng đường bi C dịch chuyển"], [["Bi A", ""], ["Bi B", ""]])

body(doc, "Thí nghiệm 2: Chỉ thả bi A, nhưng tăng độ dốc của máng trượt. Quan sát và đo quãng đường "
          "bi C dịch chuyển sau mỗi lần va chạm.", indent_cm=1.3)
data_table(doc, ["Độ dốc", "Quãng đường bi C dịch chuyển"], [["Thấp", ""], ["Cao", ""]])

body(doc, "Thảo luận:", bold=True, indent_cm=1.3, space_after=2)
for label, text, n_lines in [
    ("a)", "Trong thí nghiệm 1, vận tốc của bi A và bi B khi đến chân dốc có giống nhau không? Viên "
           "bi nào làm bi C lăn xa hơn? Tại sao?", 2),
    ("b)", "Trong thí nghiệm 2, ứng với độ dốc nào thì bi A có vận tốc lớn hơn khi va chạm với bi C? "
           "Ở trường hợp nào bi C lăn xa hơn? Tại sao?", 2),
    ("c)", "Từ hai thí nghiệm trên, hãy rút ra kết luận: khả năng truyền chuyển động (sau này gọi là "
           "động lượng) của một vật phụ thuộc vào những yếu tố nào?", 2),
]:
    body(doc, f"{label} {text}", indent_cm=1.3, space_after=2)
    ruled_lines(doc, n_lines, indent_cm=1.3)

sub_item_label(doc, "b", "Kiến thức trọng tâm")
embed_image(doc, IMG_DIR / "vecto_dong_luong.png", 9.5)
for text in [
    "Định nghĩa: động lượng của vật khối lượng m chuyển động với vận tốc v là p = (1)….........",
    "Đặc điểm: động lượng là đại lượng (2)…........., luôn (3)…......... với vectơ vận tốc của vật.",
    "Đơn vị: (4)…...........",
    "Ý nghĩa vật lí: động lượng đặc trưng cho (5)…........................................ giữa các "
    "vật khi chúng tương tác (va chạm) với nhau.",
]:
    bullet(doc, text)
body(doc, "Gợi ý từ khóa: p = m.v (dạng vectơ) · vectơ · cùng phương, cùng chiều · kg.m/s · khả năng "
          "truyền tương tác/truyền chuyển động.", italic=True, size=11, indent_cm=1.3)
separator(doc)

# --- 2.2 Xung lượng của lực ---
subsection_heading(doc, "2.2", "Xung lượng của lực")

sub_item_label(doc, "a", "Nhiệm vụ")

body(doc, "Nhiệm vụ 1 – Đọc đồ thị lực – thời gian", bold=True, indent_cm=1.1, space_after=2)
instructions(doc, "hoạt động nhóm 4 học sinh, 7 phút, quan sát đồ thị lực – thời gian bên và đọc mục "
                   "II.1 Xung lượng, SGK trang 111.", indent_cm=1.3)
embed_image(doc, IMG_DIR / "do_thi_luc_thoi_gian.png", 9.0)
body(doc, "Đồ thị mô tả độ lớn F của một lực không đổi tác dụng lên một vật trong khoảng thời gian từ "
          "t₁ đến t₂.", indent_cm=1.3)
for label, text, n_lines in [
    ("a)", "Diện tích phần tô màu trên đồ thị được tính bằng biểu thức nào theo F và Δt = t₂ – t₁?", 1),
    ("b)", "Đại lượng đó được gọi là gì trong Vật lí? Nêu đơn vị của đại lượng đó.", 1),
    ("c)", "Liên hệ thực tế: khi thủ môn bắt bóng, tại sao động tác co tay lại và lùi người theo hướng "
           "bóng bay giúp đỡ đau tay hơn? (Gợi ý: hành động đó làm thay đổi đại lượng nào trên đồ thị — "
           "F hay Δt?)", 2),
]:
    body(doc, f"{label} {text}", indent_cm=1.3, space_after=2)
    ruled_lines(doc, n_lines, indent_cm=1.3)

body(doc, "Nhiệm vụ 2 – Suy luận công thức liên hệ", bold=True, indent_cm=1.1, space_after=2)
instructions(doc, "hoạt động cá nhân 3 phút, sau đó thảo luận chung cả lớp 2 phút, dựa vào công thức "
                   "định luật II Newton F = m.a đã học ở Bài 15 và định nghĩa gia tốc a = (v₂ – v₁)/Δt.",
             indent_cm=1.3)
body(doc, "Một vật khối lượng m đang chuyển động với vận tốc v₁. Một lực F không đổi tác dụng lên vật "
          "trong khoảng thời gian Δt làm vận tốc của vật biến đổi thành v₂.", indent_cm=1.3)
for label, text, n_lines in [
    ("a)", "Viết biểu thức gia tốc a của vật theo v₁, v₂, Δt.", 1),
    ("b)", "Thay biểu thức gia tốc vào định luật II Newton F = m.a, hãy biến đổi để đưa F.Δt về dạng "
           "chứa m.v₂ và m.v₁.", 1),
    ("c)", "Nhận xét: m.v₂ và m.v₁ chính là động lượng của vật ở hai thời điểm nào? Từ đó hãy phát "
           "biểu mối liên hệ giữa xung lượng của lực F.Δt và độ biến thiên động lượng Δp = p₂ – p₁.", 2),
    ("d)", "Viết lại định luật II Newton dưới dạng tổng quát theo Δp và Δt.", 1),
]:
    body(doc, f"{label} {text}", indent_cm=1.3, space_after=2)
    ruled_lines(doc, n_lines, indent_cm=1.3)

sub_item_label(doc, "b", "Kiến thức trọng tâm")
for text in [
    "Xung lượng của lực: khi lực F (không đổi) tác dụng lên vật trong khoảng thời gian ngắn Δt thì "
    "tích (1)…........ được gọi là xung lượng của lực F. Đơn vị: (2)…...........",
    "Liên hệ với độ biến thiên động lượng: F.Δt = (3)…...........",
    "Dạng tổng quát của định luật II Newton: (4)…...........",
]:
    bullet(doc, text)
body(doc, "Gợi ý từ khóa: F.Δt · N.s · độ biến thiên động lượng (p₂ – p₁ = Δp) · F = Δp/Δt.",
     italic=True, size=11, indent_cm=1.3)
separator(doc)

# ============================================================
# 3. LUYỆN TẬP
# ============================================================
section_heading(doc, "3", "LUYỆN TẬP")

subsection_heading(doc, "3.1", "Động lượng")
body(doc, "Câu hỏi trắc nghiệm 4 phương án: Một xe máy khối lượng 120 kg (kể cả người lái) đang chuyển "
          "động với tốc độ 54 km/h. Độ lớn động lượng của xe máy là", indent_cm=1.1)
body(doc, "A. 1800 kg.m/s.        B. 6480 kg.m/s.        C. 900 kg.m/s.        D. 15 kg.m/s.",
     indent_cm=1.3)
ruled_lines(doc, 3, indent_cm=1.1)

subsection_heading(doc, "3.2", "Xung lượng của lực")
body(doc, "Câu hỏi trả lời ngắn: Một quả bóng bàn khối lượng 2,7 g bay tới đập vuông góc vào mặt vợt "
          "với tốc độ 15 m/s rồi bật ngược trở lại theo phương cũ với tốc độ 12 m/s. Biết thời gian va "
          "chạm giữa bóng và vợt là 2.10⁻³ s. Tính độ lớn lực trung bình mà vợt tác dụng lên bóng (làm "
          "tròn kết quả đến hàng đơn vị, đơn vị N).", indent_cm=1.1)
ruled_lines(doc, 4, indent_cm=1.1)
separator(doc)

# ============================================================
# 4. VẬN DỤNG - MỞ RỘNG
# ============================================================
section_heading(doc, "4", "VẬN DỤNG - MỞ RỘNG")

subsection_heading(doc, "4.1", "Động lượng")
body(doc, "Mở rộng: Khái niệm động lượng không chỉ dùng để so sánh \"vật nào khó dừng lại hơn\" mà "
          "còn là nền tảng của nguyên lí phóng tên lửa: tên lửa đẩy khí nóng phụt ra phía sau với tốc "
          "độ rất lớn để tự thân nó thu được động lượng theo chiều ngược lại (bay lên phía trước). Đây "
          "chính là ý tưởng mở đầu cho định luật bảo toàn động lượng mà em sẽ học ở bài tiếp theo. "
          "(Nguồn tham khảo: SGK Vật lí 10 – Kết nối tri thức với cuộc sống, mục \"Em có thể\", trang "
          "112.)", indent_cm=1.1, italic=True)

subsection_heading(doc, "4.2", "Xung lượng của lực")
sub_item_label(doc, "a", "Nhiệm vụ vận dụng")
instructions(doc, "hoạt động nhóm đôi, 7 phút, dựa vào công thức liên hệ xung lượng – độ biến thiên "
                   "động lượng vừa học (mục 2.2) và hiểu biết thực tế về mũ bảo hiểm.")
body(doc, "Một kĩ sư đề xuất thiết kế mũ bảo hiểm gắn thêm một lớp lót bằng THÉP CỨNG sát da đầu (thay "
          "cho lớp xốp EPS êm như mũ bảo hiểm thật), với lí do \"thép cứng chịu lực tốt hơn xốp nên sẽ "
          "bảo vệ đầu tốt hơn\".", indent_cm=1.3)
for label, text, n_lines in [
    ("a)", "Dựa vào công thức F = Δp/Δt, hãy chỉ ra điểm bất hợp lí trong lí do mà kĩ sư đưa ra.", 2),
    ("b)", "Đề xuất cách khắc phục đúng nguyên lí vật lí (gợi ý: vai trò của lớp xốp êm là gì trong Δt "
           "lúc va chạm?).", 2),
]:
    body(doc, f"{label} {text}", indent_cm=1.3, space_after=2)
    ruled_lines(doc, n_lines, indent_cm=1.3)

sub_item_label(doc, "b", "Mở rộng")
body(doc, "Nguyên lí \"kéo dài Δt để giảm lực F khi Δp không đổi\" được ứng dụng rộng rãi trong an "
          "toàn giao thông: túi khí ô tô bung ra và xẹp từ từ trong khoảng 30–50 mili giây khi va chạm, "
          "kéo dài đáng kể thời gian tương tác giữa người và vô-lăng/bảng táp-lô so với va chạm trực "
          "tiếp, nhờ đó giảm mạnh lực tác dụng lên cơ thể người ngồi trong xe. Dây an toàn co giãn nhẹ "
          "cũng hoạt động theo cùng nguyên lí này.", indent_cm=1.3, italic=True)

doc.add_paragraph()
p_end = add_para(doc, align=WD_ALIGN_PARAGRAPH.RIGHT, space_before=10)
add_run(p_end, "Ngày ........ tháng ........ năm 202...\n", italic=True, size=12)
add_run(p_end, "Học sinh ký tên\n\n\n..........................................", bold=True, size=12)

OUT_DIR.mkdir(parents=True, exist_ok=True)
doc.save(str(OUT_PATH))
print("OK - da xuat PHT cau truc moi")
